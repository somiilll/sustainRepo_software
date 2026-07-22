"""Phase B8: Reports router.

Five endpoints lifted verbatim from server.py:
  - GET  /reports/facility/{facility_id}      (DOCX/PDF facility report)
  - POST /reports/combined                    (multi-facility combined report)
  - POST /reports/ghg-inventory               (ISO 14064-1 GHG Inventory)
  - GET  /reports/download/{download_token}   (token-based download cache)
  - POST /reports/ai-summary                  (Anthropic AI summary)

Behaviour byte-identical: route bodies preserved; only the decorator
target changed from `api_router` to the modular `router`. Heavy imports
(reportlab, mammoth, playwright, report_generator) remain lazy/inline
inside route handlers — same as legacy.
"""
import asyncio
import base64
import io
import json
import logging
import secrets
import string
import uuid
from datetime import datetime, timezone, timedelta
import calendar
from io import BytesIO
from typing import Any, Dict, List, Optional

import anthropic
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel

from app.config.env import ANTHROPIC_API_KEY
from modules.auth.dependencies import get_current_user
from shared.cache.downloads import pending_downloads
from shared.database.mongo import db

import re

logger = logging.getLogger(__name__)
router = APIRouter()


def parse_period_to_date(period_str: str, is_end: bool = False) -> Optional[str]:
    """
    Parse FY/CY/YYYY-MM period strings to YYYY-MM-DD date format.
    For start dates, returns first day. For end dates, returns last day.
    """
    if not period_str:
        return None
    period_str = str(period_str).strip()
    upper_period = period_str.upper()
    
    # Handle "CY 2024" format
    if upper_period.startswith("CY"):
        try:
            year = int(upper_period.replace("CY", "").strip())
            return f"{year}-12-31" if is_end else f"{year}-01-01"
        except ValueError:
            pass
    
    # Handle "FY" formats (April-March fiscal year)
    if upper_period.startswith("FY"):
        try:
            fy_str = upper_period.replace("FY", "").strip()
            if "-" in fy_str:
                start_yr_str = fy_str.split("-")[0].strip()
                start_year = int(start_yr_str) if len(start_yr_str) == 4 else int("20" + start_yr_str[-2:])
                return f"{start_year + 1}-03-31" if is_end else f"{start_year}-04-01"
            else:
                end_year = int(fy_str) if len(fy_str) == 4 else int("20" + fy_str[-2:])
                return f"{end_year}-03-31" if is_end else f"{end_year - 1}-04-01"
        except ValueError:
            pass
    
    # Handle "YYYY-MM" format
    if re.match(r'^\d{4}-\d{2}$', period_str):
        year, month = period_str.split('-')
        if is_end:
            last_day = calendar.monthrange(int(year), int(month))[1]
            return f"{period_str}-{last_day:02d}"
        return f"{period_str}-01"
    
    # Handle "YYYY" format
    if re.match(r'^\d{4}$', period_str):
        return f"{period_str}-12-31" if is_end else f"{period_str}-01-01"
    
    return None


# Report generation endpoint with year-wise breakdown
@router.get("/reports/facility/{facility_id}")
async def generate_facility_report(
    facility_id: str,
    start_period: Optional[str] = None,
    end_period: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    facility = await db.facilities.find_one({"id": facility_id}, {"_id": 0})
    if not facility:
        raise HTTPException(status_code=404, detail="Facility not found")
    
    # Check access - Reports are admin-only
    if current_user["role"] == "user":
        raise HTTPException(status_code=403, detail="Reports are only accessible to admins")
    if current_user["role"] == "admin" and facility["organization_id"] != current_user.get("organization_id"):
        raise HTTPException(status_code=403, detail="Not authorized")
    
    query = {"facility_id": facility_id}
    if start_period and end_period:
        query["reporting_period"] = {"$gte": start_period, "$lte": end_period}
    
    emissions = await db.emission_records.find(query, {"_id": 0}).to_list(10000)
    
    doc = Document()
    
    # Title
    title = doc.add_heading('GHG Emissions Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add period info
    if start_period and end_period:
        period_para = doc.add_paragraph(f'Reporting Period: {start_period} to {end_period}')
        period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        period_para.runs[0].font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Facility details
    doc.add_heading('Facility Information', 1)
    doc.add_paragraph(f"Name: {facility['name']}")
    doc.add_paragraph(f"Address: {facility['address']}")
    if facility.get('sector'):
        doc.add_paragraph(f"Sector: {facility['sector']}")
    if facility.get('responsible_person'):
        doc.add_paragraph(f"Responsible Person: {facility['responsible_person']}")
    
    doc.add_paragraph()
    
    # Overall Summary
    doc.add_heading('Overall Emissions Summary', 1)
    total_emissions = sum(e.get("total_emissions", 0) or 0 for e in emissions)
    scope1_total = sum(e.get("total_emissions", 0) or 0 for e in emissions if e["scope"] == "scope1")
    scope2_total = sum(e.get("total_emissions", 0) or 0 for e in emissions if e["scope"] == "scope2")
    biogenic_total = sum(e.get("total_emissions", 0) or 0 for e in emissions if e["scope"] == "biogenic")
    
    doc.add_paragraph(f"Total Emissions: {round(total_emissions, 2)} kg CO2e")
    doc.add_paragraph(f"Scope 1 Emissions: {round(scope1_total, 2)} kg CO2e ({round(scope1_total/total_emissions*100 if total_emissions > 0 else 0, 1)}%)")
    doc.add_paragraph(f"Scope 2 Emissions: {round(scope2_total, 2)} kg CO2e ({round(scope2_total/total_emissions*100 if total_emissions > 0 else 0, 1)}%)")
    doc.add_paragraph(f"Biogenic Emissions: {round(biogenic_total, 2)} kg CO2e ({round(biogenic_total/total_emissions*100 if total_emissions > 0 else 0, 1)}%)")
    
    # Chart
    if emissions:
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
        
        # Pie chart
        labels = ['Scope 1', 'Scope 2', 'Biogenic']
        sizes = [scope1_total, scope2_total, biogenic_total]
        colors = ['#1A4D2E', '#4F6F52', '#E85C0D']
        non_zero = [(label, size, color) for label, size, color in zip(labels, sizes, colors) if size > 0]
        if non_zero:
            labels_nz, sizes_nz, colors_nz = zip(*non_zero)
            ax1.pie(sizes_nz, labels=labels_nz, colors=colors_nz, autopct='%1.1f%%', startangle=90)
        ax1.set_title('Overall Emissions by Scope')
        
        # Bar chart by period
        period_map = {}
        for e in emissions:
            period = e["reporting_period"]
            if period not in period_map:
                period_map[period] = {"scope1": 0, "scope2": 0, "biogenic": 0}
            if e["scope"] == "scope1":
                period_map[period]["scope1"] += e.get("total_emissions", 0) or 0
            elif e["scope"] == "scope2":
                period_map[period]["scope2"] += e.get("total_emissions", 0) or 0
            else:
                period_map[period]["biogenic"] += e.get("total_emissions", 0) or 0
        
        periods = sorted(period_map.keys())
        scope1_data = [period_map[p]["scope1"] for p in periods]
        scope2_data = [period_map[p]["scope2"] for p in periods]
        biogenic_data = [period_map[p]["biogenic"] for p in periods]
        
        x = range(len(periods))
        width = 0.25
        ax2.bar([i - width for i in x], scope1_data, width, label='Scope 1', color='#1A4D2E')
        ax2.bar(x, scope2_data, width, label='Scope 2', color='#4F6F52')
        ax2.bar([i + width for i in x], biogenic_data, width, label='Biogenic', color='#E85C0D')
        ax2.set_xlabel('Reporting Period')
        ax2.set_ylabel('Emissions (kg CO2e)')
        ax2.set_title('Emissions Trend')
        ax2.set_xticks(x)
        ax2.set_xticklabels(periods, rotation=45, ha='right')
        ax2.legend()
        
        plt.tight_layout()
        
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=100, bbox_inches='tight')
        img_buffer.seek(0)
        plt.close()
        
        doc.add_picture(img_buffer, width=Inches(6))
    
    doc.add_page_break()
    
    # Year-wise breakdown
    doc.add_heading('Year-wise Emissions Breakdown', 1)
    
    # Group emissions by year
    year_emissions = {}
    for emission in emissions:
        year = emission.get("reporting_period", "").split('-')[0]
        if year not in year_emissions:
            year_emissions[year] = []
        year_emissions[year].append(emission)
    
    # Sort years in descending order (most recent first)
    for year in sorted(year_emissions.keys(), reverse=True):
        year_data = year_emissions[year]
        
        # Year heading
        doc.add_heading(f'Calendar Year {year}', 2)
        
        # Year summary
        year_total = sum(e.get("total_emissions", 0) or 0 for e in year_data)
        year_scope1 = sum(e.get("total_emissions", 0) or 0 for e in year_data if e["scope"] == "scope1")
        year_scope2 = sum(e.get("total_emissions", 0) or 0 for e in year_data if e["scope"] == "scope2")
        year_biogenic = sum(e.get("total_emissions", 0) or 0 for e in year_data if e["scope"] == "biogenic")
        
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Year {year} Total: ").bold = True
        summary_para.add_run(f"{round(year_total, 2)} kg CO2e\n")
        summary_para.add_run(f"  • Scope 1: {round(year_scope1, 2)} kg CO2e\n")
        summary_para.add_run(f"  • Scope 2: {round(year_scope2, 2)} kg CO2e\n")
        summary_para.add_run(f"  • Biogenic: {round(year_biogenic, 2)} kg CO2e")
        
        doc.add_paragraph()
        
        # Year table
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Period'
        hdr_cells[1].text = 'Scope'
        hdr_cells[2].text = 'Category'
        hdr_cells[3].text = 'Sub-category'
        hdr_cells[4].text = 'Quantity'
        hdr_cells[5].text = 'Factor'
        hdr_cells[6].text = 'Total (kg CO2e)'
        
        for emission in sorted(year_data, key=lambda x: x["reporting_period"]):
            row_cells = table.add_row().cells
            row_cells[0].text = emission.get("reporting_period", "")
            row_cells[1].text = emission["scope"].upper().replace("SCOPE", "Scope ").replace("BIOGENIC", "Biogenic")
            row_cells[2].text = emission["category"]
            row_cells[3].text = emission["sub_category"]
            row_cells[4].text = str(emission["quantity"])
            row_cells[5].text = str(emission["emission_factor"])
            row_cells[6].text = str(round(emission.get("total_emissions", 0) or 0, 2))
        
        doc.add_paragraph()
    
    # Save to buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    # Generate download token and store report
    download_token = str(uuid.uuid4())
    filename = f"GHG_Report_{facility['name'].replace(' ', '_')}_{start_period or 'all'}_{end_period or 'all'}.docx"
    
    # Clean up old downloads (older than 5 minutes)
    current_time = datetime.now(timezone.utc)
    expired_tokens = []
    
    for token, data in pending_downloads.items():
        created_at = data.get("created_at")
        # If created_at is missing, or if it's older than 300 seconds, expire it
        if not created_at or (current_time - created_at).total_seconds() > 300:
            expired_tokens.append(token)

    for token in expired_tokens:
        del pending_downloads[token]
    
    pending_downloads[download_token] = {
        "buffer": doc_buffer.read(),
        "filename": filename,
        "created_at": current_time
    }
    
    return {"download_token": download_token, "filename": filename}

# Combined Report for multiple facilities
@router.post("/reports/combined")
async def generate_combined_report(
    facility_ids: List[str],
    start_period: Optional[str] = None,
    end_period: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    if not facility_ids:
        raise HTTPException(status_code=400, detail="No facilities selected")
    
    # Get organization details
    organization = None
    if current_user.get("organization_id"):
        organization = await db.organizations.find_one(
            {"id": current_user["organization_id"]}, 
            {"_id": 0}
        )
    
    # Get all selected facilities
    facilities_data = []
    # Reports are admin-only
    if current_user["role"] == "user":
        raise HTTPException(status_code=403, detail="Reports are only accessible to admins")
    
    for fid in facility_ids:
        facility = await db.facilities.find_one({"id": fid}, {"_id": 0})
        if facility:
            # Check access for admin
            if current_user["role"] == "admin" and facility.get("organization_id") != current_user.get("organization_id"):
                continue
            
            query = {"facility_id": fid}
            if start_period and end_period:
                query["reporting_period"] = {"$gte": start_period, "$lte": end_period}
            
            emissions = await db.emission_records.find(query, {"_id": 0}).to_list(10000)
            facilities_data.append({"facility": facility, "emissions": emissions})
    
    if not facilities_data:
        raise HTTPException(status_code=404, detail="No accessible facilities found")
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Combined GHG Emissions Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if start_period and end_period:
        period_para = doc.add_paragraph(f'Reporting Period: {start_period} to {end_period}')
        period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Organization Details (if available)
    if organization:
        doc.add_heading('Organization Information', 1)
        doc.add_paragraph(f"Name: {organization.get('name', 'N/A')}")
        
        address_parts = [organization.get('corporate_address', '')]
        if organization.get('city'):
            address_parts.append(organization['city'])
        if organization.get('state'):
            address_parts.append(organization['state'])
        if organization.get('country'):
            address_parts.append(organization['country'])
        if organization.get('pincode'):
            address_parts.append(f"({organization['pincode']})")
        
        doc.add_paragraph(f"Address: {', '.join(filter(None, address_parts))}")
        
        if organization.get('general_description'):
            doc.add_paragraph(f"Description: {organization['general_description']}")
        if organization.get('mission'):
            doc.add_paragraph(f"Mission: {organization['mission']}")
        if organization.get('vision'):
            doc.add_paragraph(f"Vision: {organization['vision']}")
        
        doc.add_paragraph()
    
    # Overall Summary across all facilities
    doc.add_heading('Overall Summary', 1)
    all_emissions = []
    for fd in facilities_data:
        all_emissions.extend(fd["emissions"])
    
    total_emissions = sum(e.get("total_emissions", 0) or 0 for e in all_emissions)
    scope1_total = sum(e.get("total_emissions", 0) or 0 for e in all_emissions if e["scope"] == "scope1")
    scope2_total = sum(e.get("total_emissions", 0) or 0 for e in all_emissions if e["scope"] == "scope2")
    biogenic_total = sum(e.get("total_emissions", 0) or 0 for e in all_emissions if e["scope"] == "biogenic")
    
    doc.add_paragraph(f"Total Facilities Included: {len(facilities_data)}")
    doc.add_paragraph(f"Total Emissions: {round(total_emissions, 2)} kg CO2e")
    doc.add_paragraph(f"Scope 1 Emissions: {round(scope1_total, 2)} kg CO2e ({round(scope1_total/total_emissions*100 if total_emissions > 0 else 0, 1)}%)")
    doc.add_paragraph(f"Scope 2 Emissions: {round(scope2_total, 2)} kg CO2e ({round(scope2_total/total_emissions*100 if total_emissions > 0 else 0, 1)}%)")
    doc.add_paragraph(f"Biogenic Emissions: {round(biogenic_total, 2)} kg CO2e ({round(biogenic_total/total_emissions*100 if total_emissions > 0 else 0, 1)}%)")
    
    doc.add_page_break()
    
    # Year-wise breakdown across all facilities
    doc.add_heading('Year-wise Emissions Breakdown', 1)
    
    year_emissions = {}
    for emission in all_emissions:
        year = emission.get("reporting_period", "").split('-')[0]
        if year not in year_emissions:
            year_emissions[year] = {"emissions": [], "by_facility": {}}
        year_emissions[year]["emissions"].append(emission)
        
        fac_id = emission["facility_id"]
        if fac_id not in year_emissions[year]["by_facility"]:
            year_emissions[year]["by_facility"][fac_id] = []
        year_emissions[year]["by_facility"][fac_id].append(emission)
    
    for year in sorted(year_emissions.keys(), reverse=True):
        year_data = year_emissions[year]["emissions"]
        
        doc.add_heading(f'Calendar Year {year}', 2)
        
        year_total = sum(e.get("total_emissions", 0) or 0 for e in year_data)
        year_scope1 = sum(e.get("total_emissions", 0) or 0 for e in year_data if e["scope"] == "scope1")
        year_scope2 = sum(e.get("total_emissions", 0) or 0 for e in year_data if e["scope"] == "scope2")
        year_biogenic = sum(e.get("total_emissions", 0) or 0 for e in year_data if e["scope"] == "biogenic")
        
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Year {year} Total: ").bold = True
        summary_para.add_run(f"{round(year_total, 2)} kg CO2e\n")
        summary_para.add_run(f"  • Scope 1: {round(year_scope1, 2)} kg CO2e\n")
        summary_para.add_run(f"  • Scope 2: {round(year_scope2, 2)} kg CO2e\n")
        summary_para.add_run(f"  • Biogenic: {round(year_biogenic, 2)} kg CO2e")
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # Details for each facility
    for idx, fd in enumerate(facilities_data):
        facility = fd["facility"]
        emissions = fd["emissions"]
        
        doc.add_heading(f'Facility {idx + 1}: {facility["name"]}', 1)
        
        # Facility info
        doc.add_paragraph(f"Address: {facility.get('address', 'N/A')}")
        if facility.get('city') or facility.get('state') or facility.get('country'):
            location_parts = [facility.get('city'), facility.get('state'), facility.get('country')]
            doc.add_paragraph(f"Location: {', '.join(filter(None, location_parts))}")
        if facility.get('sector'):
            doc.add_paragraph(f"Sector: {facility['sector']}")
        if facility.get('responsible_person'):
            doc.add_paragraph(f"Responsible Person: {facility['responsible_person']}")
        
        # Facility emissions summary
        fac_total = sum(e.get("total_emissions", 0) or 0 for e in emissions)
        fac_scope1 = sum(e.get("total_emissions", 0) or 0 for e in emissions if e["scope"] == "scope1")
        fac_scope2 = sum(e.get("total_emissions", 0) or 0 for e in emissions if e["scope"] == "scope2")
        fac_biogenic = sum(e.get("total_emissions", 0) or 0 for e in emissions if e["scope"] == "biogenic")
        
        doc.add_paragraph()
        doc.add_paragraph(f"Total Emissions: {round(fac_total, 2)} kg CO2e")
        doc.add_paragraph(f"Scope 1: {round(fac_scope1, 2)} kg CO2e | Scope 2: {round(fac_scope2, 2)} kg CO2e | Biogenic: {round(fac_biogenic, 2)} kg CO2e")
        
        # Emission records table
        if emissions:
            doc.add_paragraph()
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'Period'
            hdr[1].text = 'Scope'
            hdr[2].text = 'Category'
            hdr[3].text = 'Quantity'
            hdr[4].text = 'Factor'
            hdr[5].text = 'Emissions (kg)'
            
            for em in sorted(emissions, key=lambda x: x["reporting_period"], reverse=True):
                row = table.add_row().cells
                row[0].text = em["reporting_period"]
                row[1].text = em["scope"].replace("scope", "Scope ")
                row[2].text = em.get("category", "")
                row[3].text = str(em["quantity"])
                row[4].text = str(em["emission_factor"])
                row[5].text = str(round(em["total_emissions"], 2))
        
        if idx < len(facilities_data) - 1:
            doc.add_page_break()
    
    # Save to buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    # Generate download token and store report
    download_token = str(uuid.uuid4())
    filename = f"Combined_GHG_Report_{start_period or 'all'}_{end_period or 'all'}.docx"
    
    # Clean up old downloads (older than 5 minutes)
    current_time = datetime.now(timezone.utc)
    expired_tokens = [
        token for token, data in pending_downloads.items()
        if (current_time - data["created_at"]).total_seconds() > 300
    ]
    for token in expired_tokens:
        del pending_downloads[token]
    
    pending_downloads[download_token] = {
        "buffer": doc_buffer.read(),
        "filename": filename,
        "created_at": current_time
    }
    
    return {"download_token": download_token, "filename": filename}

# GHG Inventory Report Generation
class GHGReportRequest(BaseModel):
    facility_ids: List[str]
    reporting_period_start: str  # Format: YYYY-MM
    reporting_period_end: str    # Format: YYYY-MM
    include_previous_years: bool = False
    organization_id: Optional[str] = None  # For SuperAdmin to specify organization
    output_format: str = "docx"  # "docx" or "pdf"
    report_type: str = "scope_1_2"  # "scope_1_2" or "scope_1_2_3"
    is_complete_organization: bool = True  # Whether all org facilities are included (for org-level sections)

@router.post("/reports/ghg-inventory")
async def generate_ghg_inventory_report(
    request: GHGReportRequest,
    current_user: dict = Depends(get_current_user)
):
    """Generate GHG Inventory Report based on template"""
    from report_generator import GHGReportGenerator
    
    if not request.facility_ids:
        raise HTTPException(status_code=400, detail="No facilities selected")
    
    # Get organization details - handle SuperAdmin case
    organization = None
    org_id = current_user.get("organization_id")
    
    # SuperAdmin can specify organization_id, or we get it from the first facility
    if current_user.get("role") == "super_admin":
        if request.organization_id:
            org_id = request.organization_id
        else:
            # Get organization from first facility
            first_facility = await db.facilities.find_one({"id": request.facility_ids[0]}, {"_id": 0})
            if first_facility:
                org_id = first_facility.get("organization_id")
    
    if org_id:
        organization = await db.organizations.find_one({"id": org_id}, {"_id": 0})
    
    # If still no organization, create a default one
    if not organization:
        organization = {
            "name": "Organization",
            "address": "Not Available",
            "city": "Not Available",
            "state": "Not Available",
            "country": "Not Available",
            "description": "Not Available"
        }
    
    # Get all selected facilities
    facilities_data = []
    for fid in request.facility_ids:
        facility = await db.facilities.find_one({"id": fid}, {"_id": 0})
        if facility:
            # Check access based on role - Reports are admin-only
            if current_user.get("role") == "super_admin":
                facilities_data.append(facility)
            elif current_user.get("role") == "user":
                continue  # Users cannot access reports
            elif current_user.get("role") == "admin" and facility.get("organization_id") != current_user.get("organization_id"):
                continue
            else:
                facilities_data.append(facility)
    
    if not facilities_data:
        raise HTTPException(status_code=404, detail="No accessible facilities found")
    
    # Base year validation based on is_complete_organization flag
    # If is_complete_organization = true: Org-level base year for Scope 1,2 must exist
    # If is_complete_organization = false: Each selected facility must have base year for Scope 1,2
    
    if request.is_complete_organization:
        # Check for org-level base year data
        org_base_year_record = await db.base_year_emissions.find_one(
            {"organization_id": org_id, "facility_id": None},
            {"_id": 0, "id": 1}
        )
        if not org_base_year_record:
            raise HTTPException(
                status_code=400, 
                detail="Organization-level Base Year Emissions data is required when including all facilities. Please configure base year emissions at the organization level first."
            )
    else:
        # Check individual facility base year data
        missing_base_year = []
        for facility in facilities_data:
            base_year_record = await db.base_year_emissions.find_one(
                {"facility_id": facility["id"]}, 
                {"_id": 0, "id": 1}
            )
            if not base_year_record:
                missing_base_year.append(facility.get("name", facility["id"]))
        
        if missing_base_year:
            raise HTTPException(
                status_code=400, 
                detail=f"Base year emissions data is required before generating the report. Missing for: {', '.join(missing_base_year)}"
            )
    
    # Get emissions within reporting period
    emissions_data = []
    for facility in facilities_data:
        # Fetch monthly records with date range filter
        monthly_query = {
            "facility_id": facility["id"],
            "frequency_type": {"$ne": "yearly"},
            "reporting_period": {
                "$gte": request.reporting_period_start,
                "$lte": request.reporting_period_end
            }
        }
        cursor = db.emission_records.find(monthly_query, {"_id": 0})
        monthly_emissions = await cursor.to_list(length=1000)
        emissions_data.extend(monthly_emissions)
        
        # Fetch yearly records separately (CY/FY format doesn't work with string comparison)
        # These will be filtered by _filter_emissions_by_period in the report generator
        yearly_query = {
            "facility_id": facility["id"],
            "frequency_type": "yearly"
        }
        cursor = db.emission_records.find(yearly_query, {"_id": 0})
        yearly_emissions = await cursor.to_list(length=1000)
        emissions_data.extend(yearly_emissions)
    
    # Get previous years data if requested
    previous_years_data = []
    if request.include_previous_years:
        for facility in facilities_data:
            # Get ONLY emissions BEFORE the reporting period start (not within the period)
            # This prevents double-counting emissions that are already in emissions_data
            query = {
                "facility_id": facility["id"],
                "reporting_period": {"$lt": request.reporting_period_start}
            }
            cursor = db.emission_records.find(query, {"_id": 0})
            prev_facility_emissions = await cursor.to_list(length=1000)
            previous_years_data.extend(prev_facility_emissions)
        # Add ONLY previous years emissions to emissions_data
        emissions_data.extend(previous_years_data)
    
    # Get sinks data within reporting period
    sinks_data = []
    # Parse period strings to actual dates for sinks query
    sinks_start_date = parse_period_to_date(request.reporting_period_start, is_end=False)
    sinks_end_date = parse_period_to_date(request.reporting_period_end, is_end=True)
    
    for facility in facilities_data:
        # Filter sinks that OVERLAP with the reporting period
        # Overlap condition: sink.start_date <= report_end AND sink.end_date >= report_start
        sinks_query = {"facility_id": facility["id"]}
        if sinks_start_date and sinks_end_date:
            sinks_query["start_date"] = {"$lte": sinks_end_date}
            sinks_query["end_date"] = {"$gte": sinks_start_date}
        cursor = db.sinks.find(sinks_query, {"_id": 0})
        facility_sinks = await cursor.to_list(length=1000)
        sinks_data.extend(facility_sinks)
    
    # Calculate proportional sinks for this period
    def calculate_sink_proportion(sink, report_start, report_end):
        """Calculate the proportion of sink that falls within the report period."""
        try:
            from datetime import datetime
            sink_start = datetime.strptime(sink.get('start_date', ''), '%Y-%m-%d')
            sink_end = datetime.strptime(sink.get('end_date', ''), '%Y-%m-%d')
            rep_start = datetime.strptime(report_start, '%Y-%m-%d')
            rep_end = datetime.strptime(report_end, '%Y-%m-%d')
            
            # Calculate overlap
            overlap_start = max(sink_start, rep_start)
            overlap_end = min(sink_end, rep_end)
            
            if overlap_start > overlap_end:
                return 0.0  # No overlap
            
            # Calculate proportion based on days
            sink_total_days = (sink_end - sink_start).days + 1
            overlap_days = (overlap_end - overlap_start).days + 1
            
            if sink_total_days <= 0:
                return 1.0
            
            return overlap_days / sink_total_days
        except (ValueError, TypeError):
            return 1.0  # If parsing fails, include full value
    
    total_sinks = 0.0
    for s in sinks_data:
        proportion = calculate_sink_proportion(s, sinks_start_date, sinks_end_date) if sinks_start_date and sinks_end_date else 1.0
        s['_proportion'] = proportion
        total_sinks += s.get("total_emissions_reduced", 0) * proportion
    
    # Filter emissions based on report_type
    # For scope_1_2 report: exclude scope3 emissions, include only biogenic scope1
    # For scope_1_2_3 report: include all emissions
    if request.report_type == "scope_1_2":
        filtered_emissions = []
        for e in emissions_data:
            scope = (e.get("scope") or "").lower()
            # Include scope1 and scope2
            if scope in ["scope1", "scope2"]:
                filtered_emissions.append(e)
            # Include biogenic only if it's scope1 (direct biogenic)
            elif scope == "biogenic":
                biogenic_selection = (e.get("biogenic_scope_selection") or "").lower()
                # Include only direct/scope1 biogenic emissions
                if biogenic_selection in ["scope1", "direct", ""]:
                    filtered_emissions.append(e)
            # Exclude scope3
        emissions_data = filtered_emissions
    # For scope_1_2_3: include everything (no filtering needed)
    
    # Fetch facility production data from production_quantities collection with proportional allocation
    import re
    
    async def get_production_for_period(facility_id_or_none, start_period, end_period, org_id):
        """
        Get production quantity for a report period with proportional allocation.
        Returns (quantity, unit) or (None, None) if not found.
        """
        query = {
            "organization_id": org_id,
            "is_deleted": {"$ne": True}
        }
        if facility_id_or_none:
            query["facility_id"] = facility_id_or_none
        else:
            query["facility_id"] = None  # Organization-level
        
        records = await db.production_quantities.find(query, {"_id": 0}).to_list(10000)
        if not records:
            return None, None
        
        # Parse report period range
        def parse_ym(period_str):
            match = re.match(r'(\d{4})-(\d{2})', period_str)
            if match:
                return int(match.group(1)), int(match.group(2))
            return None, None
        
        start_y, start_m = parse_ym(start_period)
        end_y, end_m = parse_ym(end_period)
        
        if not start_y or not end_y:
            return None, None
        
        # Generate list of months in report range
        report_months = []
        y, m = start_y, start_m
        while (y, m) <= (end_y, end_m):
            report_months.append((y, m))
            m += 1
            if m > 12:
                m = 1
                y += 1
        
        # Process each production record
        aggregated_qty = 0.0
        unit = None
        
        for record in records:
            period = record.get("reporting_period", "")
            qty = record.get("quantity", 0)
            rec_unit = record.get("unit", "")
            
            overlap_months = 0
            total_period_months = 1
            
            # Monthly format: YYYY-MM
            monthly_match = re.match(r'^(\d{4})-(\d{2})$', period)
            if monthly_match:
                rec_y, rec_m = int(monthly_match.group(1)), int(monthly_match.group(2))
                if (rec_y, rec_m) in report_months:
                    overlap_months = 1
                    total_period_months = 1
            
            # FY format
            fy_match = re.match(r'^FY\s*(\d{4})-(\d{2,4})$', period, re.IGNORECASE)
            if fy_match:
                fy_start_year = int(fy_match.group(1))
                fy_months = []
                for mo in range(4, 13):
                    fy_months.append((fy_start_year, mo))
                for mo in range(1, 4):
                    fy_months.append((fy_start_year + 1, mo))
                total_period_months = 12
                overlap_months = len(set(fy_months) & set(report_months))
            
            # CY format
            cy_match = re.match(r'^CY\s*(\d{4})$', period, re.IGNORECASE)
            if cy_match:
                cy_year = int(cy_match.group(1))
                cy_months = [(cy_year, mo) for mo in range(1, 13)]
                total_period_months = 12
                overlap_months = len(set(cy_months) & set(report_months))
            
            if overlap_months > 0:
                proportion = overlap_months / total_period_months
                aggregated_qty += qty * proportion
                # Capture unit from first overlapping record
                if unit is None:
                    unit = rec_unit
        
        return (round(aggregated_qty, 4) if aggregated_qty > 0 else None, unit)
    
    # Fetch production data from collection for each facility
    facility_production_data = {}
    for facility in facilities_data:
        fid = facility["id"]
        qty, unit = await get_production_for_period(
            fid, 
            request.reporting_period_start, 
            request.reporting_period_end,
            org_id
        )
        if qty and unit:
            facility_production_data[fid] = {'quantity': qty, 'unit': unit}
    
    # Fetch organization-level production quantity (facility_id = None)
    org_production_data = None
    org_qty, org_unit = await get_production_for_period(
        None,  # Organization level
        request.reporting_period_start, 
        request.reporting_period_end,
        org_id
    )
    if org_qty and org_unit:
        org_production_data = {'quantity': org_qty, 'unit': org_unit}
    
    # Generate report - pass backend URL for internal file access
    generator = GHGReportGenerator(backend_base_url='http://localhost:8001')
    report_buffer = generator.generate_report(
        organization=organization,
        facilities=facilities_data,
        emissions=emissions_data,
        reporting_period_start=request.reporting_period_start,
        reporting_period_end=request.reporting_period_end,
        include_previous_years=request.include_previous_years,
        sinks_total=total_sinks,
        sinks_data=sinks_data,
        facility_production=facility_production_data,
        org_production=org_production_data,
        report_type=request.report_type,
        is_complete_organization=request.is_complete_organization
    )
    
    # Generate filename based on format
    org_name = organization.get('name', 'Organization').replace(' ', '_')
    file_extension = "pdf" if request.output_format == "pdf" else "docx"
    filename = f"GHG_Inventory_Report_{org_name}_{request.reporting_period_start}_{request.reporting_period_end}.{file_extension}"
    
    # Convert to PDF if requested using Playwright
    if request.output_format == "pdf":
        try:
            import tempfile
            import os
            import mammoth
            
            # Save docx to temp file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                report_buffer.seek(0)
                temp_docx.write(report_buffer.read())
                temp_docx_path = temp_docx.name
            
            # Convert DOCX to HTML using mammoth
            with open(temp_docx_path, 'rb') as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html_content = result.value
            
            # Create styled HTML document
            styled_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    @page {{
                        size: A4;
                        margin: 20mm;
                    }}
                    body {{
                        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                        font-size: 11pt;
                        line-height: 1.5;
                        color: #333;
                        max-width: 100%;
                    }}
                    h1 {{ font-size: 18pt; color: #1a5f2a; margin-top: 20px; margin-bottom: 10px; }}
                    h2 {{ font-size: 14pt; color: #2d7d46; margin-top: 16px; margin-bottom: 8px; }}
                    h3 {{ font-size: 12pt; color: #3d9d56; margin-top: 12px; margin-bottom: 6px; }}
                    p {{ margin: 8px 0; text-align: justify; }}
                    table {{
                        width: 100%;
                        border-collapse: collapse;
                        margin: 10px 0;
                        font-size: 10pt;
                    }}
                    th, td {{
                        border: 1px solid #ddd;
                        padding: 6px 8px;
                        text-align: left;
                    }}
                    th {{
                        background-color: #f5f5f5;
                        font-weight: bold;
                    }}
                    img {{
                        max-width: 100%;
                        height: auto;
                        margin: 10px 0;
                    }}
                    .page-break {{
                        page-break-before: always;
                    }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            # Save HTML to temp file
            with tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w', encoding='utf-8') as temp_html:
                temp_html.write(styled_html)
                temp_html_path = temp_html.name
            
            # Use Playwright async API to generate PDF
            from playwright.async_api import async_playwright
            
            async with async_playwright() as p:
                browser = await p.chromium.launch(headless=True)
                page = await browser.new_page()
                await page.goto(f'file://{temp_html_path}')
                
                pdf_bytes = await page.pdf(
                    format='A4',
                    margin={'top': '20mm', 'bottom': '20mm', 'left': '15mm', 'right': '15mm'},
                    print_background=True
                )
                await browser.close()
            
            report_buffer = io.BytesIO(pdf_bytes)
            report_buffer.seek(0)
            
            # Cleanup temp files
            os.unlink(temp_docx_path)
            os.unlink(temp_html_path)
            
            logger.info("PDF generated successfully using Playwright")
            
        except Exception as e:
            # Fallback to docx if PDF conversion fails
            logger.error(f"PDF conversion error: {str(e)}")
            filename = filename.replace('.pdf', '.docx')
            report_buffer.seek(0)
    
    # Generate download token and store report
    download_token = str(uuid.uuid4())
    
    # Clean up old downloads (older than 5 minutes)
    current_time = datetime.now(timezone.utc)
    expired_tokens = [
        token for token, data in pending_downloads.items()
        if (current_time - data["created_at"]).total_seconds() > 300
    ]
    for token in expired_tokens:
        del pending_downloads[token]
    
    # Store new download
    report_buffer.seek(0)
    pending_downloads[download_token] = {
        "buffer": report_buffer.read(),
        "filename": filename,
        "created_at": current_time
    }
    
    return {"download_token": download_token, "filename": filename}


@router.get("/reports/download/{download_token}")
async def download_report(download_token: str):
    """Download a generated report using token"""
    if download_token not in pending_downloads:
        raise HTTPException(status_code=404, detail="Download link expired or invalid")
    
    download_data = pending_downloads[download_token]
    
    # Create a new BytesIO from the stored bytes
    buffer = io.BytesIO(download_data["buffer"])
    buffer.seek(0)
    
    # Note: Token is NOT deleted immediately - it will expire after 5 minutes
    # This allows retry if download fails in sandboxed environments
    
    # Determine content type from filename
    filename = download_data['filename']
    if filename.endswith('.pdf'):
        content_type = "application/pdf"
    else:
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    return StreamingResponse(
        buffer,
        media_type=content_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# ============== AI REPORT GENERATION ==============

class AIReportRequest(BaseModel):
    facility_ids: List[str]
    reporting_period_start: str
    reporting_period_end: str


async def aggregate_emissions_for_ai(organization_id: str, facility_ids: List[str], start_period: str, end_period: str) -> dict:
    """Aggregate emission data for AI report generation - applies equity share and temporal proportion if applicable"""
    # Get organization info
    org = await db.organizations.find_one({"id": organization_id}, {"_id": 0})
    if not org:
        return None
    
    # Check if equity share approach is used
    use_equity_share = org.get("org_boundaries_approach") == "equity_share"
    
    # Get facilities that belong to the organization AND are in the requested list
    facilities = await db.facilities.find({
        "id": {"$in": facility_ids},
        "organization_id": organization_id
    }, {"_id": 0}).to_list(100)
    
    if not facilities:
        return None
    
    # Build facility equity map
    facility_equity_map = {}
    for f in facilities:
        if use_equity_share:
            equity_pct = f.get("equity_share_percentage", 100.0) or 100.0
            facility_equity_map[f['id']] = equity_pct / 100.0
        else:
            facility_equity_map[f['id']] = 1.0
    
    # Use only the facility IDs that actually belong to this org
    valid_facility_ids = [f['id'] for f in facilities]
    
    # Query emission_records
    emissions = await db.emission_records.find({
        "facility_id": {"$in": valid_facility_ids}
    }, {"_id": 0}).to_list(10000)

    # --- DATE PARSING & PROPORTION LOGIC ---
    def parse_date_str(date_str, is_end=False):
        """Helper to safely parse dates and handle CY/FY/Month-end logic"""
        if not date_str:
            return None
        date_str = str(date_str).strip()
        
        upper_date = date_str.upper()

        # Handle "CY 2024" format
        if upper_date.startswith("CY"):
            try:
                year = int(upper_date.replace("CY", "").strip())
                return datetime(year, 12, 31) if is_end else datetime(year, 1, 1)
            except ValueError:
                pass

        # Handle "FY" formats (assuming standard April-March fiscal year)
        # e.g., "FY 2024", "FY24", "FY 2024-25", "FY 24-25"
        if upper_date.startswith("FY"):
            try:
                fy_str = upper_date.replace("FY", "").strip()
                
                # Format: "2024-25" or "24-25"
                if "-" in fy_str:
                    start_yr_str = fy_str.split("-")[0].strip()
                    # Normalize to 4-digit year
                    start_year = int(start_yr_str) if len(start_yr_str) == 4 else int("20" + start_yr_str[-2:])
                    return datetime(start_year + 1, 3, 31) if is_end else datetime(start_year, 4, 1)
                
                # Format single year: "2024" or "24"
                else:
                    # Normalize to 4-digit year
                    end_year = int(fy_str) if len(fy_str) == 4 else int("20" + fy_str[-2:])
                    # FY2024 typically means the year ending March 31, 2024 (April 2023 - March 2024)
                    return datetime(end_year, 3, 31) if is_end else datetime(end_year - 1, 4, 1)
            except ValueError:
                pass

        # Standard formats to attempt
        formats = [
            ("%Y-%m-%d", False),
            ("%Y-%m", True),
            ("%Y", True),
            ("%B %Y", True), # e.g. "April 2024"
            ("%b %Y", True)  # e.g. "Apr 2024"
        ]
        
        for fmt, requires_month_end in formats:
            try:
                dt = datetime.strptime(date_str, fmt)
                if is_end and requires_month_end:
                    if fmt in ["%Y-%m", "%B %Y", "%b %Y"]:
                        last_day = calendar.monthrange(dt.year, dt.month)[1]
                        return dt.replace(day=last_day)
                    elif fmt == "%Y":
                        return datetime(dt.year, 12, 31)
                return dt
            except ValueError:
                continue
        return None

    
    req_start_dt = parse_date_str(start_period, is_end=False)
    req_end_dt = parse_date_str(end_period, is_end=True)

    def calculate_proportion(period_str):
        """Calculates what fraction of the record's date range falls within the requested report range"""
        if not req_start_dt or not req_end_dt or not period_str:
            return 0.0
        # Normalize separators
        period_str = period_str.replace(' to ', ' - ')
        if ' - ' in period_str:
            parts = period_str.split(' - ')
            rec_start_dt = parse_date_str(parts[0], is_end=False)
            rec_end_dt = parse_date_str(parts[1], is_end=True)
        else:
            rec_start_dt = parse_date_str(period_str, is_end=False)
            rec_end_dt = parse_date_str(period_str, is_end=True)

        if not rec_start_dt or not rec_end_dt:
            return 0.0

        # Find the overlap window
        overlap_start = max(req_start_dt, rec_start_dt)
        overlap_end = min(req_end_dt, rec_end_dt)

        if overlap_start > overlap_end:
            return 0.0 # No overlap

        overlap_days = (overlap_end - overlap_start).days + 1
        total_record_days = (rec_end_dt - rec_start_dt).days + 1

        if total_record_days <= 0:
            return 0.0

        return overlap_days / total_record_days
    # ---------------------------------------
    
    # Filter and inject proportion into emissions
    filtered_emissions = []
    for e in emissions:
        prop = calculate_proportion(e.get('reporting_period', ''))
        if prop > 0:
            e['_proportion'] = prop
            filtered_emissions.append(e)
    
    if not filtered_emissions:
        return None
    
    # Helper to get CO2e value with BOTH equity share and time proportion applied
    def get_co2e(e):
        raw_value = e.get('calculated_co2e') or e.get('co2e_emissions') or e.get('total_emissions') or 0
        facility_id = e.get('facility_id')
        equity_factor = facility_equity_map.get(facility_id, 1.0)
        proportion_factor = e.get('_proportion', 1.0)
        return raw_value * equity_factor * proportion_factor
    
    scope1_total = 0
    scope2_total = 0
    scope3_total = 0
    biogenic_total = 0

    # Aggregate by scope 
    # scope1_total = sum(get_co2e(e) for e in filtered_emissions if e.get('scope') == 'scope1')
    # scope2_total = sum(get_co2e(e) for e in filtered_emissions if e.get('scope') == 'scope2')
    # scope3_total = sum(get_co2e(e) for e in filtered_emissions if e.get('scope') == 'scope3')
    # biogenic_total = sum(get_co2e(e) for e in filtered_emissions if e.get('scope') == 'biogenic')

    for e in filtered_emissions:
        raw_scope = (e.get('scope') or "").lower()
        safe_scope = raw_scope.replace(" ", "")
        
        co2e_val = get_co2e(e)
        
        if safe_scope == 'scope1':
            scope1_total += co2e_val
        elif safe_scope == 'scope2':
            scope2_total += co2e_val
        elif safe_scope == 'scope3':
                scope3_total += co2e_val
        elif safe_scope == 'biogenic':
            biogenic_total += co2e_val
    
    gross_emissions = scope1_total + scope2_total + scope3_total
    
    # Get sinks data 
    sinks = await db.sinks.find({
        "facility_id": {"$in": valid_facility_ids}
    }, {"_id": 0}).to_list(1000)
    
    # Filter and inject proportion into sinks
    filtered_sinks = []
    for s in sinks:
        sink_period = None
        if s.get('reporting_period'):
            sink_period = s['reporting_period']
        elif s.get('start_date') and s.get('end_date'):
            sink_period = f"{s['start_date']} to {s['end_date']}"
        elif s.get('start_date'):
            sink_period = s['start_date']
        elif s.get('reporting_year'):
            year = s['reporting_year']
            month = s.get('reporting_month', 0) + 1
            sink_period = f"{year}-{month:02d}"

        prop = calculate_proportion(sink_period)
        if prop > 0:
            s['_proportion'] = prop
            filtered_sinks.append(s)
    
    # Calculate total sinks
    total_sinks = 0
    sinks_breakdown = []
    facility_name_map = {f['id']: f['name'] for f in facilities}
    
    for s in filtered_sinks:
        sink_value = s.get('total_emissions_reduced', 0) or 0
        equity_factor = facility_equity_map.get(s.get('facility_id'), 1.0)
        proportion_factor = s.get('_proportion', 1.0)
        
        adjusted_value = sink_value * equity_factor * proportion_factor
        total_sinks += adjusted_value
        
        if adjusted_value > 0:
            sinks_breakdown.append({
                "sink_type": s.get('sink_type') or s.get('type') or 'Carbon Sink',
                "description": s.get('description') or '',
                "emissions_reduced_tco2e": round(adjusted_value, 4),
                "facility": facility_name_map.get(s.get('facility_id'), 'Unknown'),
                "period": s.get('reporting_period') or s.get('start_date', '')[:7] if s.get('start_date') else ''
            })
    
    # Aggregate by category 
    category_breakdown = {}
    for e in filtered_emissions:
        cat = e.get('category', 'Unknown')
        if cat not in category_breakdown:
            category_breakdown[cat] = {'co2e': 0, 'count': 0}
        category_breakdown[cat]['co2e'] += get_co2e(e)
        category_breakdown[cat]['count'] += 1
    
    sorted_categories = sorted(category_breakdown.items(), key=lambda x: x[1]['co2e'], reverse=True)
    
    # Aggregate by facility
    facility_breakdown = {}
    for e in filtered_emissions:
        fid = e.get('facility_id')
        if fid not in facility_breakdown:
            facility_breakdown[fid] = {'co2e': 0, 'count': 0, 'equity_pct': facility_equity_map.get(fid, 1.0) * 100}
        facility_breakdown[fid]['co2e'] += get_co2e(e)
        facility_breakdown[fid]['count'] += 1
    
    facility_data = [
        {
            'name': facility_name_map.get(fid, 'Unknown'), 
            'co2e': data['co2e'], 
            'count': data['count'],
            'equity_share_pct': data['equity_pct']
        }
        for fid, data in facility_breakdown.items()
    ]
    facility_data.sort(key=lambda x: x['co2e'], reverse=True)
    
    # Check for custom factors usage
    custom_factor_count = sum(1 for e in filtered_emissions if e.get('is_custom_factor'))
    override_count = sum(1 for e in filtered_emissions if e.get('override_calorific_value') or e.get('override_density'))
    
    aggregated_data = {
        "organization_name": org.get('name', 'Organization'),
        "reporting_period": f"{start_period} to {end_period}",
        "consolidation_approach": "Equity Share" if use_equity_share else "Control (Operational/Financial)",
        "equity_share_applied": use_equity_share,
        "facilities_count": len(facilities),
        "facility_names": [f['name'] for f in facilities],
        "total_emission_records": len(filtered_emissions),
        "emissions_summary": {
            "gross_emissions_tco2e": round(gross_emissions, 4),
            "scope1_tco2e": round(scope1_total, 4),
            "scope2_tco2e": round(scope2_total, 4),
            "scope3_tco2e": round(scope3_total, 4),
            "biogenic_tco2e": round(biogenic_total, 4),
            "carbon_sinks_tco2e": round(total_sinks, 4),
            "net_emissions_tco2e": round(gross_emissions - total_sinks, 4)
        },
        "scope1_percentage": round((scope1_total / gross_emissions * 100) if gross_emissions > 0 else 0, 1),
        "scope2_percentage": round((scope2_total / gross_emissions * 100) if gross_emissions > 0 else 0, 1),
        "breakdown_by_category": [
            {"category": cat, "co2e_tco2e": round(data['co2e'], 4), "record_count": data['count']}
            for cat, data in sorted_categories[:10]
        ],
        "breakdown_by_facility": facility_data[:10],
        "carbon_sinks_details": {
            "total_sinks_tco2e": round(total_sinks, 4),
            "sinks_count": len(filtered_sinks),
            "breakdown": sinks_breakdown[:10] if sinks_breakdown else []
        },
        "data_quality": {
            "custom_emission_factors_used": custom_factor_count,
            "parameter_overrides_used": override_count,
            "total_records": len(filtered_emissions)
        }
    }
    
    return aggregated_data

async def generate_ai_summary(aggregated_data: dict, mask_org_name: bool = True) -> str:
    """Generate executive summary using Claude AI
    
    Args:
        aggregated_data: The emissions data to analyze
        mask_org_name: If True, masks organization and facility names before sending to AI
    """
    
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="AI service not configured")
    
    # Store original names and create masking mappings
    original_org_name = aggregated_data.get("organization_name", "Organization")
    masked_org_name = "[THE ORGANIZATION]"
    
    # Create facility name mappings
    facility_name_mapping = {}  # {masked_name: original_name}
    
    # Create a copy of data with masked names for AI
    ai_data = aggregated_data.copy()
    if mask_org_name:
        ai_data["organization_name"] = masked_org_name
        
        # Mask facility names in breakdown_by_facility
        if "breakdown_by_facility" in ai_data:
            masked_facilities = []
            for i, facility in enumerate(ai_data["breakdown_by_facility"]):
                # original_name = facility.get("facility_name", f"Facility {i+1}")
                original_name = facility.get("name") or facility.get("facility_name") or f"Facility {i+1}"
                masked_name = f"[FACILITY_{i+1}]"
                facility_name_mapping[masked_name] = original_name
                
                masked_facility = facility.copy()
                masked_facility["facility_name"] = masked_name
                masked_facilities.append(masked_facility)
            ai_data["breakdown_by_facility"] = masked_facilities
        
        # Also mask in sinks_by_facility if present
        if "sinks_by_facility" in ai_data:
            masked_sinks = []
            for sink in ai_data["sinks_by_facility"]:
                original_name = sink.get("facility_name", "Unknown")
                # Find corresponding masked name or create new one
                masked_name = None
                for mname, oname in facility_name_mapping.items():
                    if oname == original_name:
                        masked_name = mname
                        break
                if not masked_name:
                    idx = len(facility_name_mapping) + 1
                    masked_name = f"[FACILITY_{idx}]"
                    facility_name_mapping[masked_name] = original_name
                
                masked_sink = sink.copy()
                masked_sink["facility_name"] = masked_name
                masked_sinks.append(masked_sink)
            ai_data["sinks_by_facility"] = masked_sinks
    
    equity_context = ""
    if aggregated_data.get("equity_share_applied"):
        equity_context = """
IMPORTANT CONTEXT: This organization uses the EQUITY SHARE consolidation approach. All emission figures have been adjusted 
based on each facility's equity share percentage. Mention this in your summary - that emissions are reported 
proportionally based on the organization's equity stake in each facility.
"""
    
    system_prompt = f"""You are an expert Chief Sustainability Officer (CSO) assistant writing an executive summary and strategic action plan for a corporate GHG emissions report.
You will be provided with pre-calculated, emissions data in JSON format.
{equity_context}
CORE REPORTING RULES:
1. STRICT DATA INTEGRITY: Do NOT calculate, invent, or estimate any metrics. Use ONLY the exact quantitative values provided in the JSON.
2. NO VERIFICATION CLAIMS: Do NOT use the word "verified" to describe the emissions or the data. Do not claim or imply that the data has undergone third-party verification.
3. Format the output using clear Markdown headings and bullet points for readability.
4. BULLET-LIST HIERARCHY ONLY: The PDF rendering engine does not support markdown tables or the "|" character. You must present all data breakdowns, scope summaries, category rankings, and priority timelines using a nested, bolded bullet-point hierarchy exactly as specified in the Output Structure Templates below.
5. Keep the tone objective, clinical for the data, and strategic for the recommendations.
6. The output of the emissions should always be shown in units tCO2e (tonnes of CO2 equivalent) with exactly 2 decimal places.
7. When referring to the organization, use "{masked_org_name}" exactly as provided - do not use any other name.
8. When referring to facilities, use the facility names exactly as provided in the data (e.g., [FACILITY_1], [FACILITY_2]).
9. All numerical values should be formatted to exactly 2 decimal places.


REQUIRED STRUCTURE:

### 1. Executive Emissions Overview
Provide a detailed summary of total gross emissions, net emissions, and the Scope 1, 2 & 3 breakdown. Mention the reporting period and number of facilities covered. Report Net GHG Emissions first, then mention Biogenic emissions separately if they exist. Note if custom emission factors or overrides were used (critical for audit transparency).

### 2. Primary Emission Drivers
Analyze the 'breakdown_by_category' data. Identify and explain the top sources driving the carbon footprint so stakeholders understand exactly where the emissions are coming from.

### 3. Outlier & Emission Anomaly Insights
Review the provided emissions data and identify any significant outliers, unusual emission concentrations, or anomalies across facilities, scopes, categories, or reporting periods.
- Highlight facilities or categories that contribute disproportionately to total emissions.
- Identify unusually high emission sources that may warrant further operational investigation or targeted reduction efforts.
- If no meaningful outliers are evident from the provided data, explicitly state that no significant emission anomalies were identified.
- Do NOT invent, estimate, or statistically derive values that are not directly supported by the provided data.
- Focus on management-relevant insights that may indicate operational inefficiencies, abnormal activity, data quality concerns, or concentrated emission hotspots.

### 4. Strategic Decarbonization & Reduction Pathways
Based strictly on the highest emitting categories identified above, provide 3 to 4 tailored, actionable recommendations to reduce emissions. 
- Tailor the advice: If mobile combustion is a primary driver, suggest fleet electrification or logistics optimization. If stationary combustion/electricity is high, suggest renewable energy procurement (PPAs) or HVAC efficiency upgrades.
- Where applicable for hard-to-abate emissions, include brief suggestions on carbon capture technology, transitioning to low-carbon alternative fuels, or investing in verified carbon sinks/offsets.

OUTPUT STRUCTURE TEMPLATES (MANDATORY):
Whenever displaying list metrics, breakdowns, or rankings, use the following exact text structures. Do not invent any other layout or use table characters.

Example for Scope / Category breakdowns:
• **[Scope/Category Name]**: [Value] tCO2e ([Percentage]%)

Example for Decarbonization Priorities / Recommendations:
• **Priority [Number]: [Category Name]**
  - **Current Emissions**: [Value] tCO2e
  - **Target Reduction**: [Value]%
  - **Timeline**: [Value]

"""
    
    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        
        message = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=10000,
            temperature=0.3,
            system=system_prompt,
            messages=[
                {
                    "role": "user",
                    "content": json.dumps(ai_data)
                }
            ]
        )
        
        ai_response = message.content[0].text
        
        # Unmask: Replace masked names with original names
        if mask_org_name:
            # Replace organization name
            ai_response = ai_response.replace(masked_org_name, original_org_name)
            ai_response = ai_response.replace("[THE ORGANIZATION]", original_org_name)
            ai_response = ai_response.replace("THE ORGANIZATION", original_org_name)
            ai_response = ai_response.replace("the organization", original_org_name)
            
            for masked_key, original_name in facility_name_mapping.items():
                # Exact match replacement
                ai_response = ai_response.replace(masked_key, original_name)
                
                # Extract the actual number from the key "[FACILITY_X]"
                match = re.search(r'\d+', masked_key)
                if match:
                    fac_num = match.group()
                    
                    # Catch AI variations using the exact number
                    variations = [
                        f"Facility {fac_num}",
                        f"facility {fac_num}",
                        f"FACILITY_{fac_num}",
                        f"Facility_{fac_num}",
                        f"facility_{fac_num}"
                    ]
                    
                    for var in variations:
                        ai_response = ai_response.replace(var, original_name)
                        ai_response = ai_response.replace(f"[{var}]", original_name)

        return ai_response
        
    except anthropic.APIError as e:
        logger.error(f"Anthropic API Error: {e}")
        error_msg = str(e)
        if "credit balance" in error_msg.lower() or "billing" in error_msg.lower():
            raise HTTPException(status_code=402, detail="AI service credits exhausted. Please add balance to your Anthropic account.")
        raise HTTPException(status_code=500, detail="Failed to generate AI summary. Please try again later.")


def generate_ai_report_pdf(aggregated_data: dict, ai_summary: str) -> io.BytesIO:
    """Generate a PDF report with AI executive summary"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Frame, PageTemplate, BaseDocTemplate
    from reportlab.pdfgen import canvas
    
    buffer = io.BytesIO()
    
    # Border color - darker blue (#1E3A5F)
    BORDER_COLOR = colors.HexColor('#1E3A5F')
    
    def add_page_border(canvas_obj, doc):
        """Draw border on each page"""
        canvas_obj.saveState()
        canvas_obj.setStrokeColor(BORDER_COLOR)
        canvas_obj.setLineWidth(2)
        # Draw rectangle with margin from edges
        margin = 20
        canvas_obj.rect(margin, margin, A4[0] - 2*margin, A4[1] - 2*margin)
        canvas_obj.restoreState()
    
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4, 
        topMargin=0.75*inch, 
        bottomMargin=0.75*inch,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        spaceAfter=20,
        textColor=colors.HexColor('#1a365d'),
        alignment=1  # Center
    )
    
    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Normal'],
        fontSize=12,
        textColor=colors.HexColor('#4a5568'),
        alignment=1,
        spaceAfter=30
    )
    
    section_style = ParagraphStyle(
        'Section',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#2d3748'),
        spaceBefore=20,
        spaceAfter=10
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#2d3748'),
        spaceAfter=12,
        leading=14
    )
    
    elements = []
    
    # Title
    elements.append(Paragraph("AI Executive Summary Report", title_style))
    elements.append(Paragraph(f"{aggregated_data['organization_name']}", subtitle_style))
    
    # Report metadata
    meta_data = [
        ["Reporting Period:", aggregated_data['reporting_period']],
        ["Consolidation Approach:", aggregated_data.get('consolidation_approach', 'Control')],
        ["Facilities Covered:", str(aggregated_data['facilities_count'])],
        ["Generated:", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")]
    ]
    
    meta_table = Table(meta_data, colWidths=[2*inch, 4*inch])
    meta_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#4a5568')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(meta_table)
    elements.append(Spacer(1, 20))
    
    # Emissions Summary Section
    elements.append(Paragraph("Emissions Summary", section_style))
    
    emissions = aggregated_data['emissions_summary']
    summary_data = [
        ["Metric", "Value (tCO2e)"],
        ["Gross Emissions (Scope 1 + 2 + 3)", f"{emissions['gross_emissions_tco2e']:,.2f}"],
        ["Scope 1 Emissions", f"{emissions['scope1_tco2e']:,.2f}"],
        ["Scope 2 Emissions", f"{emissions['scope2_tco2e']:,.2f}"],
        ["Scope 3 Emissions", f"{emissions['scope3_tco2e']:,.2f}"],
        ["Carbon Sinks", f"{emissions['carbon_sinks_tco2e']:,.2f}"],
        ["Net Emissions", f"{emissions['net_emissions_tco2e']:,.2f}"],
        ["Biogenic Emissions", f"{emissions['biogenic_tco2e']:,.2f}"],
    ]
    
    summary_table = Table(summary_data, colWidths=[3.5*inch, 2.5*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 20))
    
    # AI Executive Summary Section
    elements.append(Paragraph("AI Analysis & Recommendations", section_style))
    
    # Custom styles for markdown rendering
    heading_style = ParagraphStyle(
        'Heading',
        parent=styles['Heading3'],
        fontSize=11,
        textColor=colors.HexColor('#1a365d'),
        spaceBefore=12,
        spaceAfter=6,
        fontName='Helvetica-Bold'
    )
    
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#2d3748'),
        leftIndent=20,
        spaceAfter=4,
        leading=14,
        bulletIndent=10
    )
    
    # Process AI summary with markdown support
    # Clean special characters that don't render in PDF
    def clean_for_pdf(text):
        # CRITICAL XML ESCAPE: ReportLab Paragraphs will drop text or crash 
        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        # Replace subscript/superscript characters with ASCII equivalents
        replacements = {
            '₂': '2', '₃': '3', '₄': '4',
            '²': '2', '³': '3',
            'CO₂': 'CO2', 'tCO₂e': 'tCO2e',
            '–': '-', '—': '-',
            '"': '"',  # Left double quote
            
        }
        # Also replace right double quote (handled separately to avoid dict key collision)
        text = text.replace('"', '"')
        for old, new in replacements.items():
            text = text.replace(old, new)
        # Remove markdown formatting
        text = text.replace('**', '').replace('*', '')
        # Remove markdown heading markers
        text = text.lstrip('#').strip()
        return text


    lines = ai_summary.strip().split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            elements.append(Spacer(1, 6))
            continue
        
        # 1. Strip markdown bolding to test the core text
        test_line = line.replace('**', '').strip()
        
        # 2. Strip any accidental AI bullet markers to see the actual word
        unbulleted_test_line = test_line.lstrip('-•* ')
        
        # 3. Robust Heading Detection
        is_heading = (
            test_line.startswith('#') or 
            unbulleted_test_line.startswith('1. ') or 
            unbulleted_test_line.startswith('2. ') or 
            unbulleted_test_line.startswith('3. ') or 
            unbulleted_test_line.startswith('4. ') or 
            unbulleted_test_line.startswith('Recommendation') or 
            unbulleted_test_line.startswith('Conclusion') or
            unbulleted_test_line.startswith('Summary of Recommended')
        )
        
        if is_heading:
            # Process as a HEADING (forces stripping of accidental bullets/dashes)
            clean_text = clean_for_pdf(unbulleted_test_line)
            if clean_text:
                elements.append(Paragraph(clean_text, heading_style))
                
        elif line.startswith('-') or line.startswith('•') or line.startswith('*'):
            # Process as a TRUE BULLET POINT
            bullet_text = line.lstrip('-•* ').strip()
            clean_text = clean_for_pdf(bullet_text)
            if clean_text:
                elements.append(Paragraph(f"&bull; {clean_text}", bullet_style))
                
        else:
            # Process as a NORMAL PARAGRAPH
            clean_text = clean_for_pdf(line)
            if clean_text:
                elements.append(Paragraph(clean_text, body_style))
    
    elements.append(Spacer(1, 20))
    
    # Category Breakdown
    if aggregated_data.get('breakdown_by_category'):
        elements.append(Paragraph("Top 5 Emissions by Category", section_style))
        
        cat_data = [["Category", "Emissions (tCO2e)", "Records"]]
        for cat in aggregated_data['breakdown_by_category'][:5]:
            cat_data.append([
                cat['category'],
                f"{cat['co2e_tco2e']:,.2f}",
                str(cat['record_count'])
            ])
        
        cat_table = Table(cat_data, colWidths=[3*inch, 1.75*inch, 1.25*inch])
        cat_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (1, 0), (-1, -1), 'RIGHT'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
        ]))
        elements.append(cat_table)
    
    # Carbon Sinks Section
    sinks_details = aggregated_data.get('carbon_sinks_details', {})
    if sinks_details.get('total_sinks_tco2e', 0) > 0 or sinks_details.get('breakdown'):
        elements.append(Spacer(1, 15))
        elements.append(Paragraph("Top 5 Carbon Sinks & Offsets", section_style))
        
        if sinks_details.get('breakdown'):
            sinks_data = [["Sink Type", "Description", "CO2 Reduced (tCO2e)", "Facility"]]
            for sink in sinks_details['breakdown'][:5]:
                # Truncate facility name if too long to fit
                facility_name = sink.get('facility', 'Unknown')
                if len(facility_name) > 20:
                    facility_name = facility_name[:18] + '..'
                sinks_data.append([
                    sink.get('sink_type', 'Carbon Sink'),
                    (sink.get('description', '')[:25] + '..' if len(sink.get('description', '')) > 25 else sink.get('description', '')),
                    f"{sink.get('emissions_reduced_tco2e', 0):,.2f}",
                    facility_name
                ])
            
            # Adjusted column widths to fit facility names better
            sinks_table = Table(sinks_data, colWidths=[1.3*inch, 1.8*inch, 1.4*inch, 1.5*inch])
            sinks_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#047857')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),  # Slightly smaller font
                ('ALIGN', (2, 0), (2, -1), 'RIGHT'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#d1fae5')),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('WORDWRAP', (0, 0), (-1, -1), True),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            elements.append(sinks_table)
        else:
            elements.append(Paragraph(f"Total Carbon Sinks: {sinks_details.get('total_sinks_tco2e', 0):,.2f} tCO2e", body_style))
    
    # Build PDF with border on each page
    doc.build(elements, onFirstPage=add_page_border, onLaterPages=add_page_border)
    buffer.seek(0)
    return buffer


@router.post("/reports/ai-summary")
async def generate_ai_report_summary(
    request: AIReportRequest,
    current_user: dict = Depends(get_current_user)
):
    """Generate AI-powered executive summary PDF for emissions data"""
    
    if not request.facility_ids:
        raise HTTPException(status_code=400, detail="Please select at least one facility")
    
    if not request.reporting_period_start or not request.reporting_period_end:
        raise HTTPException(status_code=400, detail="Please specify reporting period")
    
    # Get organization from user
    organization_id = current_user.get('organization_id')
    if not organization_id:
        raise HTTPException(status_code=400, detail="User not associated with an organization")
    
    # Aggregate emissions data (with equity share applied if applicable)
    aggregated_data = await aggregate_emissions_for_ai(
        organization_id,
        request.facility_ids,
        request.reporting_period_start,
        request.reporting_period_end
    )
    if not aggregated_data:
        raise HTTPException(status_code=404, detail="No emission records found for the selected facilities and period")
    
    # Generate AI summary
    ai_summary = await generate_ai_summary(aggregated_data)
    
    # Generate PDF
    pdf_buffer = generate_ai_report_pdf(aggregated_data, ai_summary)
    
    # Create download token
    download_token = str(uuid.uuid4())
    org_name = aggregated_data['organization_name'].replace(' ', '_')
    filename = f"AI_Executive_Summary_{org_name}_{request.reporting_period_start}_to_{request.reporting_period_end}.pdf"
    
    current_time = datetime.now(timezone.utc)

    pending_downloads[download_token] = {
        "buffer": pdf_buffer.getvalue(),
        "filename": filename,
        "created_at": current_time
    }
    
    return {
        "success": True,
        "download_token": download_token,
        "filename": filename,
        "message": "AI Summary PDF generated successfully"
    }


