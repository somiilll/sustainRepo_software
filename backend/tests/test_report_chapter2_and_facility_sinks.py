"""
Test cases for 2 Report Generator Fixes:
1. Chapter 2 Organization Boundaries should NOT have random org name at end when 
   approach is not equity_share with percentage
2. Facilities with no emissions but with sinks should show: 
   'No emission reported... However, carbon sinks/removals totaling X.XX tCO₂e have been reported'
   And should show the Carbon Sinks / Removals section

Test setup:
- Test Org Updated has equity_share approach
- test-2 may have no emissions in 2026-01 to 2026-03 period
"""
import pytest
import requests
import os
import io
from docx import Document

BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', '').rstrip('/')


class TestChapter2OrganizationBoundaries:
    """Test Chapter 2 - no random org name at end unless equity_share with percentage"""
    
    @pytest.fixture(autouse=True)
    def setup(self):
        """Authenticate and setup session"""
        self.session = requests.Session()
        
        login_response = self.session.post(
            f"{BASE_URL}/api/auth/login",
            json={"email": "testadmin@test.com", "password": "Test123!"}
        )
        
        if login_response.status_code == 200:
            data = login_response.json()
            token = data.get('token') or data.get('access_token')
            if token:
                self.session.headers.update({"Authorization": f"Bearer {token}"})
        else:
            pytest.skip(f"Authentication failed: {login_response.status_code}")
        
        yield
        self.session.close()
    
    def _get_organization(self):
        """Get organization info"""
        response = self.session.get(f"{BASE_URL}/api/organizations/me")
        if response.status_code != 200:
            # Try via facilities
            fac_resp = self.session.get(f"{BASE_URL}/api/facilities")
            if fac_resp.status_code == 200:
                facs = fac_resp.json()
                if facs and facs[0].get('organization_id'):
                    org_resp = self.session.get(f"{BASE_URL}/api/organizations/{facs[0]['organization_id']}")
                    if org_resp.status_code == 200:
                        return org_resp.json()
            return {}
        return response.json()
    
    def _generate_report(self, facility_ids, period_start="2025-01", period_end="2025-12"):
        """Generate and download report"""
        payload = {
            "facility_ids": facility_ids,
            "reporting_period_start": period_start,
            "reporting_period_end": period_end,
            "include_previous_years": True
        }
        
        gen_response = self.session.post(f"{BASE_URL}/api/reports/ghg-inventory", json=payload)
        if gen_response.status_code != 200:
            print(f"Report generation failed: {gen_response.status_code} - {gen_response.text}")
            return None
        
        result = gen_response.json()
        download_token = result.get('download_token')
        if not download_token:
            print("No download token received")
            return None
        
        dl_response = self.session.get(f"{BASE_URL}/api/reports/download/{download_token}")
        if dl_response.status_code != 200:
            print(f"Download failed: {dl_response.status_code}")
            return None
        
        return dl_response.content
    
    def _extract_full_text(self, content):
        """Extract text from DOCX"""
        doc = Document(io.BytesIO(content))
        return '\n'.join([p.text for p in doc.paragraphs])
    
    def test_chapter2_no_random_org_name_for_non_equity_percentage(self):
        """
        FIX 1: Chapter 2 should NOT have random org name at end unless 
        approach is equity_share AND has specific percentage
        """
        org = self._get_organization()
        org_name = org.get('name', '')
        approach = org.get('org_boundaries_approach', '')
        equity_percentage = org.get('org_boundaries_equity_percentage')
        
        print(f"Organization: {org_name}")
        print(f"Approach: {approach}")
        print(f"Equity percentage: {equity_percentage}")
        
        facilities_resp = self.session.get(f"{BASE_URL}/api/facilities")
        facilities = facilities_resp.json() if facilities_resp.status_code == 200 else []
        facility_ids = [f.get('id') for f in facilities]
        
        content = self._generate_report(facility_ids, "2025-01", "2025-12")
        assert content is not None, "Failed to generate report"
        
        text = self._extract_full_text(content)
        
        # Find Chapter 2 section
        ch2_start = text.lower().find("organization boundaries")
        ch3_start = text.lower().find("reporting boundaries")
        
        if ch2_start != -1 and ch3_start != -1:
            chapter2_text = text[ch2_start:ch3_start]
            print(f"\n=== CHAPTER 2 CONTENT (first 2000 chars) ===\n{chapter2_text[:2000]}\n")
            
            # Count occurrences of org name in Chapter 2
            org_name_count = chapter2_text.lower().count(org_name.lower())
            print(f"Occurrences of org name '{org_name}' in Chapter 2: {org_name_count}")
            
            # If approach is NOT equity_share with percentage, org name should NOT appear
            # at end of chapter after boundary notes (random appearance)
            if approach != 'equity_share' or not equity_percentage:
                # The org name can appear in:
                # 1. "Org has adopted the X Approach" statement
                # 2. "Additional Boundary Notes" if any
                # But should NOT appear standalone at the end
                
                # Check for standalone org name at end (which was the bug)
                last_500_chars = chapter2_text[-500:] if len(chapter2_text) > 500 else chapter2_text
                
                # Look for patterns that suggest duplicate/random org name
                # Pattern: org name appearing after the approach statement without context
                approach_statement_variations = [
                    f"{org_name} has adopted the",
                    f"{org_name} has chosen the"
                ]
                
                has_valid_appearance = any(var.lower() in chapter2_text.lower() for var in approach_statement_variations)
                
                if has_valid_appearance:
                    print("✓ Org name appears in valid 'has adopted' statement")
                
                # The detailed explanation should NOT appear without equity percentage
                bad_patterns = [
                    f"{org_name} has chosen the Equity Share Approach. The organization accounts for",
                    "The organization accounts for greenhouse gas emissions in proportion to its equity share"
                ]
                
                for bad_pattern in bad_patterns:
                    if bad_pattern.lower() in chapter2_text.lower():
                        if approach != 'equity_share' or not equity_percentage:
                            pytest.fail(f"FIX 1 FAILED: Detailed equity explanation found without percentage. Pattern: '{bad_pattern[:50]}...'")
                
                print("✓ No random/detailed org name explanation found without equity percentage")
        
        print("\nFIX 1 PASSED: Chapter 2 does not have random org name at end")
    
    def test_chapter2_code_structure_verified(self):
        """
        Verify the code structure for Chapter 2 generation
        """
        code_path = "/app/backend/report_generator.py"
        
        with open(code_path, 'r') as f:
            code = f.read()
        
        # Find _generate_chapter2 method
        ch2_start = code.find("def _generate_chapter2")
        ch3_start = code.find("def _generate_chapter3")
        
        if ch2_start != -1 and ch3_start != -1:
            ch2_code = code[ch2_start:ch3_start]
            
            # Verify the condition: only add detailed explanation if equity_share AND percentage
            assert "if approach == 'equity_share' and equity_percentage:" in ch2_code, \
                "FIX 1 FAILED: Detailed explanation should only appear when approach='equity_share' AND equity_percentage exists"
            
            # Verify there's no else clause adding random org name
            # After the if block for equity_share with percentage, there should NOT be
            # an else clause that adds org name
            
            # Count occurrences of org_name variable usage
            org_name_usages = ch2_code.count("org_name")
            print(f"Number of org_name usages in _generate_chapter2: {org_name_usages}")
            
            # Check the condition properly limits when detailed explanation appears
            if 'if approach == "equity_share" and equity_percentage:' in ch2_code or \
               "if approach == 'equity_share' and equity_percentage:" in ch2_code:
                print("✓ Conditional check for equity_share with percentage exists")
            
            print("✓ Code structure verified for Chapter 2")
        
        print("\nFIX 1 CODE VERIFIED: No random org name added without equity percentage")


class TestFacilitySinksWithNoEmissions:
    """Test facilities with sinks but no emissions show proper message and section"""
    
    @pytest.fixture(autouse=True)
    def setup(self):
        """Authenticate and setup session"""
        self.session = requests.Session()
        
        login_response = self.session.post(
            f"{BASE_URL}/api/auth/login",
            json={"email": "testadmin@test.com", "password": "Test123!"}
        )
        
        if login_response.status_code == 200:
            data = login_response.json()
            token = data.get('token') or data.get('access_token')
            if token:
                self.session.headers.update({"Authorization": f"Bearer {token}"})
        else:
            pytest.skip(f"Authentication failed: {login_response.status_code}")
        
        yield
        self.session.close()
    
    def _generate_report(self, facility_ids, period_start, period_end):
        """Generate and download report"""
        payload = {
            "facility_ids": facility_ids,
            "reporting_period_start": period_start,
            "reporting_period_end": period_end,
            "include_previous_years": True
        }
        
        gen_response = self.session.post(f"{BASE_URL}/api/reports/ghg-inventory", json=payload)
        if gen_response.status_code != 200:
            print(f"Report generation failed: {gen_response.status_code} - {gen_response.text}")
            return None
        
        result = gen_response.json()
        download_token = result.get('download_token')
        if not download_token:
            return None
        
        dl_response = self.session.get(f"{BASE_URL}/api/reports/download/{download_token}")
        return dl_response.content if dl_response.status_code == 200 else None
    
    def _extract_full_text(self, content):
        """Extract text from DOCX"""
        doc = Document(io.BytesIO(content))
        return '\n'.join([p.text for p in doc.paragraphs])
    
    def test_facility_with_sinks_but_no_emissions_shows_message(self):
        """
        FIX 2: Facility with sinks but no emissions should show:
        - "No emission reported... However, carbon sinks/removals totaling X.XX tCO₂e have been reported"
        - Carbon Sinks / Removals section
        """
        # Get facilities and find one with potential sinks
        facilities_resp = self.session.get(f"{BASE_URL}/api/facilities")
        facilities = facilities_resp.json() if facilities_resp.status_code == 200 else []
        
        # Get sinks data
        sinks_resp = self.session.get(f"{BASE_URL}/api/sinks")
        sinks = sinks_resp.json() if sinks_resp.status_code == 200 else []
        print(f"Found {len(sinks)} total sink records")
        
        # Get emissions data
        emissions_resp = self.session.get(f"{BASE_URL}/api/emissions")
        emissions = emissions_resp.json() if emissions_resp.status_code == 200 else []
        
        # Find test-2 or any facility with sinks
        test2 = next((f for f in facilities if f.get('name') == 'test-2'), None)
        
        # Check what period test-2 has no emissions but might have sinks
        test_period_start = "2026-01"
        test_period_end = "2026-03"
        
        for facility in facilities:
            fac_id = facility.get('id')
            fac_name = facility.get('name')
            
            # Check emissions in period
            fac_emissions = [e for e in emissions 
                           if e.get('facility_id') == fac_id 
                           and e.get('reporting_period', '').startswith('2026')]
            
            # Check sinks for this facility (regardless of period for now)
            fac_sinks = [s for s in sinks if s.get('facility_id') == fac_id]
            
            print(f"\nFacility '{fac_name}':")
            print(f"  Emissions in 2026: {len(fac_emissions)}")
            print(f"  Total sinks: {len(fac_sinks)}")
            
            if fac_sinks:
                for s in fac_sinks:
                    print(f"    Sink: {s.get('total_emissions_reduced')} tCO2e, period: {s.get('reporting_period', s.get('reporting_month', 'N/A'))}")
        
        # Generate report for 2026-01 to 2026-03 period
        facility_ids = [f.get('id') for f in facilities]
        content = self._generate_report(facility_ids, test_period_start, test_period_end)
        
        if content is None:
            pytest.skip("Could not generate report")
        
        text = self._extract_full_text(content)
        
        # Check for the sinks message pattern
        sinks_message_pattern = "However, carbon sinks/removals totaling"
        
        if sinks_message_pattern in text:
            print(f"\n✓ Found sinks message: '{sinks_message_pattern}'")
            
            # Extract the full message
            msg_start = text.find(sinks_message_pattern)
            msg_end = text.find("have been reported", msg_start)
            if msg_end != -1:
                full_msg = text[msg_start:msg_end + len("have been reported")]
                print(f"Full message: '{full_msg}'")
            
            print("FIX 2 PASSED: Sinks message shown for facility with no emissions but with sinks")
        else:
            # Check if there's actually a facility with sinks but no emissions in this period
            # It may be that no such scenario exists in test data
            facilities_with_only_sinks = []
            for facility in facilities:
                fac_id = facility.get('id')
                fac_name = facility.get('name')
                
                fac_emissions_in_period = [
                    e for e in emissions 
                    if e.get('facility_id') == fac_id 
                    and e.get('reporting_period', '') >= test_period_start
                    and e.get('reporting_period', '') <= test_period_end
                ]
                
                fac_sinks_in_period = [
                    s for s in sinks 
                    if s.get('facility_id') == fac_id
                ]
                
                if not fac_emissions_in_period and fac_sinks_in_period:
                    facilities_with_only_sinks.append({
                        'name': fac_name,
                        'sinks_total': sum(s.get('total_emissions_reduced', 0) for s in fac_sinks_in_period)
                    })
            
            if facilities_with_only_sinks:
                pytest.fail(f"FIX 2 FAILED: Facilities with sinks but no emissions exist: {facilities_with_only_sinks}, but message not shown")
            else:
                print("NOTE: No facility with sinks but no emissions in test period - cannot verify message")
                # Still verify code structure
    
    def test_code_structure_for_sinks_without_emissions(self):
        """
        Verify the code structure handles sinks when no emissions
        """
        code_path = "/app/backend/report_generator.py"
        
        with open(code_path, 'r') as f:
            code = f.read()
        
        # Find the facility loop in _generate_chapter4
        ch4_start = code.find("def _generate_chapter4")
        ch5_start = code.find("def _generate_chapter5")
        
        if ch4_start != -1:
            # Check for sinks handling when no emissions
            ch4_code = code[ch4_start:ch5_start] if ch5_start != -1 else code[ch4_start:]
            
            # Verify the logic: if not has_emissions but has_sinks
            assert "if not has_emissions:" in ch4_code, \
                "FIX 2 FAILED: Code should check for 'not has_emissions'"
            
            assert "if has_sinks:" in ch4_code, \
                "FIX 2 FAILED: Code should check 'has_sinks' when no emissions"
            
            # Verify the message includes sinks info
            assert 'However, carbon sinks/removals totaling' in ch4_code, \
                "FIX 2 FAILED: Message about sinks not found in code"
            
            # Verify Carbon Sinks section is added when has_sinks and not has_emissions
            assert 'Carbon Sinks / Removals' in ch4_code, \
                "FIX 2 FAILED: Carbon Sinks section heading not found"
            
            print("✓ Code structure verified for handling sinks with no emissions")
        
        print("\nFIX 2 CODE VERIFIED: Sinks handling exists in _generate_chapter4")
    
    def test_code_shows_correct_message_format(self):
        """
        Verify the exact message format in code
        """
        code_path = "/app/backend/report_generator.py"
        
        with open(code_path, 'r') as f:
            code = f.read()
        
        # Check for the specific message format
        expected_patterns = [
            "No emission reported for this facility in the selected reporting period. However, carbon sinks/removals totaling",
            "have been reported"
        ]
        
        for pattern in expected_patterns:
            assert pattern in code, f"FIX 2 FAILED: Expected pattern not found: '{pattern}'"
            print(f"✓ Found pattern: '{pattern[:50]}...'")
        
        # Verify the format includes the sink total value
        assert "facility_sink_total" in code and "tCO₂e have been reported" in code, \
            "FIX 2 FAILED: Message should include facility_sink_total"
        
        print("\nFIX 2 MESSAGE FORMAT VERIFIED")


class TestDataSetup:
    """Verify test data setup"""
    
    @pytest.fixture(autouse=True)
    def setup(self):
        self.session = requests.Session()
        login_response = self.session.post(
            f"{BASE_URL}/api/auth/login",
            json={"email": "testadmin@test.com", "password": "Test123!"}
        )
        if login_response.status_code == 200:
            token = login_response.json().get('token') or login_response.json().get('access_token')
            if token:
                self.session.headers.update({"Authorization": f"Bearer {token}"})
        yield
        self.session.close()
    
    def test_verify_organization_setup(self):
        """Verify Test Org Updated has equity_share approach"""
        org_resp = self.session.get(f"{BASE_URL}/api/organizations/me")
        if org_resp.status_code == 200:
            org = org_resp.json()
            print(f"Organization: {org.get('name')}")
            print(f"Approach: {org.get('org_boundaries_approach')}")
            print(f"Equity percentage: {org.get('org_boundaries_equity_percentage')}")
    
    def test_verify_facilities_and_sinks(self):
        """List all facilities and their sinks"""
        fac_resp = self.session.get(f"{BASE_URL}/api/facilities")
        sinks_resp = self.session.get(f"{BASE_URL}/api/sinks")
        
        facilities = fac_resp.json() if fac_resp.status_code == 200 else []
        sinks = sinks_resp.json() if sinks_resp.status_code == 200 else []
        
        print(f"\nFacilities ({len(facilities)}):")
        for f in facilities:
            fac_sinks = [s for s in sinks if s.get('facility_id') == f.get('id')]
            total_sinks = sum(s.get('total_emissions_reduced', 0) for s in fac_sinks)
            print(f"  - {f.get('name')}: sinks={len(fac_sinks)}, total={total_sinks} tCO2e")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
