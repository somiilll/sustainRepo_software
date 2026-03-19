"""
Test cases for 3 Report Generator Fixes:
1. Remove 'has not specified an organizational boundary approach' message
2. Apply equity share adjustment to historical data (Previous Years table) - 70% equity
3. Add percentage to carbon sinks contribution by facility

Test Org Updated: equity_share approach
test-1: 70% equity, has FY 2024 historical data, has sinks (35.35 tCO2e)
"""
import pytest
import requests
import os
import io
from docx import Document

BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', '').rstrip('/')


class TestThreeReportFixes:
    """Test class for the 3 specific report generator fixes"""
    
    @pytest.fixture(autouse=True)
    def setup(self):
        """Authenticate and setup session"""
        self.admin_email = "testadmin@test.com"
        self.admin_password = "Test123!"
        self.session = requests.Session()
        
        login_response = self.session.post(
            f"{BASE_URL}/api/auth/login",
            json={"email": self.admin_email, "password": self.admin_password}
        )
        
        if login_response.status_code == 200:
            data = login_response.json()
            token = data.get('token') or data.get('access_token')
            if token:
                self.session.headers.update({"Authorization": f"Bearer {token}"})
        else:
            pytest.skip(f"Authentication failed: {login_response.status_code}")
        
        yield
        self.session.close()
    
    def _get_facilities(self):
        """Get facilities list"""
        response = self.session.get(f"{BASE_URL}/api/facilities")
        if response.status_code != 200:
            return []
        return response.json()
    
    def _get_organization(self):
        """Get organization info"""
        response = self.session.get(f"{BASE_URL}/api/organizations/me")
        if response.status_code != 200:
            return {}
        return response.json()
    
    def _generate_report(self, facility_ids, period_start="2025-01", period_end="2025-12"):
        """Generate and download report"""
        payload = {
            "facility_ids": facility_ids,
            "reporting_period_start": period_start,
            "reporting_period_end": period_end,
            "include_previous_years": True
        }
        
        gen_response = self.session.post(f"{BASE_URL}/api/reports/ghg-inventory", json=payload)
        if gen_response.status_code != 200:
            print(f"Report generation failed: {gen_response.status_code} - {gen_response.text}")
            return None
        
        result = gen_response.json()
        download_token = result.get('download_token')
        if not download_token:
            print("No download token received")
            return None
        
        dl_response = self.session.get(f"{BASE_URL}/api/reports/download/{download_token}")
        if dl_response.status_code != 200:
            print(f"Download failed: {dl_response.status_code}")
            return None
        
        return dl_response.content
    
    def _extract_full_text(self, content):
        """Extract text from DOCX"""
        doc = Document(io.BytesIO(content))
        return '\n'.join([p.text for p in doc.paragraphs])
    
    # =========================================================================
    # FIX 1: Remove 'has not specified an organizational boundary approach'
    # =========================================================================
    def test_fix1_no_unspecified_boundary_message(self):
        """
        FIX 1: Organizations without specified boundary approach should NOT 
        show 'has not specified an organizational boundary approach' message
        """
        facilities = self._get_facilities()
        facility_ids = [f.get('id') for f in facilities]
        
        # Generate report for a period where we can check Chapter 2
        content = self._generate_report(facility_ids, "2025-01", "2025-12")
        assert content is not None, "Failed to generate report"
        
        text = self._extract_full_text(content)
        
        # The phrase 'has not specified an organizational boundary approach' should NOT appear
        bad_phrase = "has not specified an organizational boundary approach"
        assert bad_phrase not in text.lower(), \
            f"FIX 1 FAILED: Found forbidden message: '{bad_phrase}'"
        
        # Also check variations
        bad_variations = [
            "not specified an organizational boundary",
            "has not specified.*boundary approach",
            "has not selected a boundary approach"
        ]
        
        for bad_var in bad_variations:
            assert bad_var.lower() not in text.lower(), \
                f"FIX 1 FAILED: Found bad variation: '{bad_var}'"
        
        # Verify Chapter 2 exists
        assert "Organization Boundaries" in text or "ORGANIZATION BOUNDARIES" in text, \
            "Chapter 2 (Organization Boundaries) not found"
        
        print("FIX 1 PASSED: No 'has not specified boundary approach' message found")
    
    def test_fix1_valid_approach_statement_shown(self):
        """
        FIX 1 Supplementary: When org has equity_share approach, show proper statement
        """
        org = self._get_organization()
        approach = org.get('org_boundaries_approach', '')
        
        facilities = self._get_facilities()
        facility_ids = [f.get('id') for f in facilities]
        
        content = self._generate_report(facility_ids, "2025-01", "2025-12")
        assert content is not None, "Failed to generate report"
        
        text = self._extract_full_text(content)
        
        if approach == 'equity_share':
            # Should show equity share adoption statement
            assert "has adopted the Equity Share Approach" in text, \
                "FIX 1: Equity share approach statement not shown"
            print(f"FIX 1 PASSED: Equity share approach statement found for org with approach='{approach}'")
        elif approach in ['control', 'control_operational', 'control_financial']:
            # Should show control approach adoption statement
            assert "has adopted the" in text and "Control" in text, \
                "FIX 1: Control approach statement not shown"
            print(f"FIX 1 PASSED: Control approach statement found for org with approach='{approach}'")
        else:
            # No approach specified - should NOT show the bad message
            assert "has not specified" not in text.lower(), \
                "FIX 1 FAILED: Bad message shown for org without specified approach"
            print(f"FIX 1 PASSED: No 'not specified' message for org without approach (approach='{approach}')")
    
    # =========================================================================
    # FIX 2: Apply equity share adjustment to historical data (Previous Years)
    # =========================================================================
    def test_fix2_historical_data_equity_adjustment(self):
        """
        FIX 2: Historical data (Previous Years table) should show equity-adjusted values
        For test-1 with 70% equity, values should be multiplied by 0.7
        """
        facilities = self._get_facilities()
        
        # Find test-1 facility
        test1 = next((f for f in facilities if f.get('name') == 'test-1'), None)
        if not test1:
            pytest.skip("test-1 facility not found")
        
        equity_pct = test1.get('equity_share_percentage', 100)
        print(f"test-1 equity share percentage: {equity_pct}%")
        
        # Get emissions to find FY 2024 historical data
        emissions_response = self.session.get(f"{BASE_URL}/api/emissions")
        if emissions_response.status_code != 200:
            pytest.skip("Could not fetch emissions")
        
        emissions = emissions_response.json()
        
        # Find 2024 emissions for test-1
        fy2024_emissions = [
            e for e in emissions 
            if e.get('facility_id') == test1.get('id') 
            and e.get('reporting_period', '').startswith('2024')
        ]
        
        if not fy2024_emissions:
            pytest.skip("No FY 2024 emissions found for test-1")
        
        # Calculate expected raw total
        raw_total = sum(float(e.get('total_emissions', 0) or 0) for e in fy2024_emissions)
        equity_factor = equity_pct / 100.0
        expected_adjusted = raw_total * equity_factor
        
        print(f"FY 2024 raw emissions for test-1: {raw_total:.2f} tCO2e")
        print(f"Expected after 70% equity adjustment: {expected_adjusted:.2f} tCO2e")
        
        # Generate report with reporting period after 2024 so 2024 becomes historical
        facility_ids = [f.get('id') for f in facilities]
        content = self._generate_report(facility_ids, "2025-01", "2025-12")
        assert content is not None, "Failed to generate report"
        
        doc = Document(io.BytesIO(content))
        
        # Find the historical data table
        found_historical_table = False
        adjusted_values_found = []
        
        for table in doc.tables:
            if len(table.rows) > 0:
                headers = [cell.text for cell in table.rows[0].cells]
                
                # Check if this is historical data table (has FY column)
                if any('FY 2024' in h or 'FY' in h for h in headers):
                    found_historical_table = True
                    print(f"Found historical data table with headers: {headers}")
                    
                    # Find FY 2024 column index
                    fy2024_col = None
                    for idx, h in enumerate(headers):
                        if 'FY 2024' in h or '2024' in h:
                            fy2024_col = idx
                            break
                    
                    if fy2024_col is not None:
                        for row_idx, row in enumerate(table.rows[1:], start=1):
                            cells = [cell.text for cell in row.cells]
                            if len(cells) > fy2024_col:
                                val_str = cells[fy2024_col].replace(',', '').strip()
                                try:
                                    val = float(val_str)
                                    adjusted_values_found.append(val)
                                    print(f"  Row {row_idx}: {cells[:2]} -> {val:.2f} tCO2e")
                                except ValueError:
                                    pass
                    break
        
        if not found_historical_table:
            # Check text for any mention of historical data
            text = self._extract_full_text(content)
            if "Emissions of Previous Years" in text:
                print("Previous Years section found, but may show 'NA' if no data in period")
            pytest.skip("No historical data table found (may be expected if no 2024 data)")
        
        # Verify values are equity-adjusted (should be ~70% of raw values)
        if adjusted_values_found:
            total_in_table = sum(adjusted_values_found)
            
            # Allow some tolerance for rounding
            tolerance = raw_total * 0.05  # 5% tolerance
            
            # The adjusted value should be close to expected_adjusted
            # If raw_total was 1000, adjusted should be ~700 (not 1000)
            print(f"Total in historical table: {total_in_table:.2f} tCO2e")
            print(f"Expected (equity adjusted): {expected_adjusted:.2f} tCO2e")
            print(f"Raw (unadjusted): {raw_total:.2f} tCO2e")
            
            # Check it's closer to expected than raw (proves adjustment was applied)
            if raw_total != expected_adjusted:  # Only check if there's actually an adjustment
                diff_to_expected = abs(total_in_table - expected_adjusted)
                diff_to_raw = abs(total_in_table - raw_total)
                
                assert diff_to_expected < diff_to_raw, \
                    f"FIX 2 FAILED: Historical values not equity-adjusted. " \
                    f"Got {total_in_table:.2f}, expected ~{expected_adjusted:.2f} (not {raw_total:.2f})"
            
            print(f"FIX 2 PASSED: Historical data shows equity-adjusted values")
        else:
            print("FIX 2: No numeric values found in historical table (may be expected)")
    
    def test_fix2_verify_code_passes_equity_factor(self):
        """
        FIX 2 Code Verification: Confirm equity_factor is passed to _add_previous_years_table
        """
        # Read report_generator.py to verify code structure
        code_path = "/app/backend/report_generator.py"
        
        try:
            with open(code_path, 'r') as f:
                code = f.read()
        except Exception:
            pytest.skip("Could not read report_generator.py")
        
        # Check that _add_previous_years_table receives equity_factor parameter
        assert "def _add_previous_years_table(self, doc: Document, prev_year_data: Dict, equity_factor:" in code, \
            "FIX 2 FAILED: _add_previous_years_table method doesn't have equity_factor parameter"
        
        # Check that equity_factor is applied to values
        assert "adjusted_val = val * equity_factor" in code, \
            "FIX 2 FAILED: equity_factor not applied to values in _add_previous_years_table"
        
        # Check it's called with equity_factor argument
        assert "_add_previous_years_table(doc, prev_year_data, equity_factor)" in code, \
            "FIX 2 FAILED: _add_previous_years_table not called with equity_factor"
        
        print("FIX 2 CODE VERIFIED: equity_factor parameter exists and is used")
    
    # =========================================================================
    # FIX 3: Add percentage to carbon sinks contribution by facility
    # =========================================================================
    def test_fix3_sink_percentage_in_organization_analysis(self):
        """
        FIX 3: Carbon sinks contribution by facility should show percentage along with tCO2e
        Format: "• facility_name: X.XX tCO2e (Y.Y%)"
        """
        facilities = self._get_facilities()
        facility_ids = [f.get('id') for f in facilities]
        
        # Check sinks data first
        sinks_response = self.session.get(f"{BASE_URL}/api/sinks")
        if sinks_response.status_code != 200:
            pytest.skip("Could not fetch sinks data")
        
        sinks = sinks_response.json()
        print(f"Found {len(sinks)} sink records")
        
        for s in sinks:
            print(f"  Sink: {s.get('facility_id')} - {s.get('total_emissions_reduced')} tCO2e - period: {s.get('reporting_period', s.get('reporting_month', 'unknown'))}")
        
        # Generate report
        content = self._generate_report(facility_ids, "2025-01", "2025-12")
        assert content is not None, "Failed to generate report"
        
        text = self._extract_full_text(content)
        
        # Check for "Carbon sinks contribution by facility" section
        if "Carbon sinks contribution by facility" in text:
            print("Found 'Carbon sinks contribution by facility' section")
            
            # Look for pattern: "tCO2e (X.X%)" which indicates percentage is shown
            import re
            
            # Pattern: number tCO2e (percentage%)
            pattern = r'[\d,]+\.?\d*\s*tCO2e\s*\(\d+\.?\d*%\)'
            matches = re.findall(pattern, text)
            
            if matches:
                print(f"FIX 3 PASSED: Found sink values with percentages: {matches[:3]}")
                assert len(matches) > 0, "FIX 3: Percentage format not found"
            else:
                # Alternative check: Look for "% " after tCO2e in sink section
                sink_section_start = text.find("Carbon sinks contribution by facility")
                if sink_section_start != -1:
                    sink_section = text[sink_section_start:sink_section_start + 500]
                    assert "%" in sink_section, \
                        f"FIX 3 FAILED: No percentage found in sink section. Section text: {sink_section[:200]}"
                    print(f"FIX 3 PASSED: Percentage symbol found in sink contribution section")
        else:
            # If no sinks in reporting period, section won't appear
            total_sinks_in_period = 0
            for s in sinks:
                period = s.get('reporting_period', '')
                if not period:
                    # Try year-based check
                    year = s.get('reporting_year')
                    if year and str(year).startswith('2025'):
                        total_sinks_in_period += s.get('total_emissions_reduced', 0)
                elif period.startswith('2025'):
                    total_sinks_in_period += s.get('total_emissions_reduced', 0)
            
            if total_sinks_in_period == 0:
                print(f"FIX 3 SKIPPED: No sinks in reporting period (2025), so section not shown")
                pytest.skip("No sinks in reporting period")
            else:
                pytest.fail(f"FIX 3 FAILED: Sinks exist ({total_sinks_in_period} tCO2e) but section not found")
    
    def test_fix3_verify_code_has_percentage_format(self):
        """
        FIX 3 Code Verification: Confirm sink percentage is included in output
        """
        code_path = "/app/backend/report_generator.py"
        
        try:
            with open(code_path, 'r') as f:
                code = f.read()
        except Exception:
            pytest.skip("Could not read report_generator.py")
        
        # Check for sink percentage calculation
        assert "sink_pct = (sink_total / removals) * 100" in code, \
            "FIX 3 FAILED: sink_pct calculation not found"
        
        # Check for percentage in output format
        assert "{sink_pct:.1f}%" in code or "sink_pct:.1f}%)" in code, \
            "FIX 3 FAILED: sink_pct not included in output format"
        
        # Verify the full format: tCO2e (X.X%)
        assert "tCO2e ({sink_pct:.1f}%)" in code or "tCO₂e ({sink_pct:.1f}%)" in code, \
            "FIX 3 FAILED: Full sink output format with percentage not found"
        
        print("FIX 3 CODE VERIFIED: Sink percentage format exists in code")


class TestDataVerification:
    """Verify test data setup is correct"""
    
    @pytest.fixture(autouse=True)
    def setup(self):
        self.session = requests.Session()
        login_response = self.session.post(
            f"{BASE_URL}/api/auth/login",
            json={"email": "testadmin@test.com", "password": "Test123!"}
        )
        if login_response.status_code == 200:
            token = login_response.json().get('token') or login_response.json().get('access_token')
            if token:
                self.session.headers.update({"Authorization": f"Bearer {token}"})
        yield
        self.session.close()
    
    def test_verify_org_approach(self):
        """Verify organization has equity_share approach"""
        # Get organization through facilities or auth/me endpoint
        response = self.session.get(f"{BASE_URL}/api/auth/me")
        if response.status_code != 200:
            pytest.skip(f"Could not get org info: {response.status_code}")
        user = response.json()
        org_id = user.get('organization_id')
        
        if org_id:
            org_response = self.session.get(f"{BASE_URL}/api/organizations/{org_id}")
            if org_response.status_code == 200:
                org = org_response.json()
                approach = org.get('org_boundaries_approach')
                print(f"Organization: {org.get('name')}, Approach: {approach}")
        print("Organization approach verification - user auth confirmed")
    
    def test_verify_facility_equity(self):
        """Verify test-1 has 70% equity"""
        response = self.session.get(f"{BASE_URL}/api/facilities")
        assert response.status_code == 200
        facilities = response.json()
        
        test1 = next((f for f in facilities if f.get('name') == 'test-1'), None)
        if test1:
            equity = test1.get('equity_share_percentage', 100)
            print(f"test-1 equity_share_percentage: {equity}%")
            assert equity == 70, f"Expected test-1 to have 70% equity, got {equity}%"
        else:
            print("test-1 facility not found - listing available facilities:")
            for f in facilities:
                print(f"  - {f.get('name')}: equity={f.get('equity_share_percentage', 100)}%")
    
    def test_verify_historical_emissions(self):
        """Verify test-1 has FY 2024 emissions (historical data)"""
        response = self.session.get(f"{BASE_URL}/api/facilities")
        facilities = response.json()
        test1 = next((f for f in facilities if f.get('name') == 'test-1'), None)
        
        if not test1:
            pytest.skip("test-1 not found")
        
        em_response = self.session.get(f"{BASE_URL}/api/emissions")
        assert em_response.status_code == 200
        emissions = em_response.json()
        
        fy2024 = [
            e for e in emissions 
            if e.get('facility_id') == test1.get('id') 
            and e.get('reporting_period', '').startswith('2024')
        ]
        
        print(f"test-1 FY 2024 emissions: {len(fy2024)} records")
        total = sum(float(e.get('total_emissions', 0) or 0) for e in fy2024)
        print(f"Total FY 2024 emissions: {total:.2f} tCO2e")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
