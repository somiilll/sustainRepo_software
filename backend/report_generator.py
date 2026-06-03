"""
GHG Inventory Report Generator
Generates DOCX reports based on the new 6-Chapter structure according to ISO 14064-1:2018
"""
import os
import io
import requests
from datetime import datetime, timezone, timedelta
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from collections import defaultdict
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np


class GHGReportGenerator:
    """Generates GHG Inventory Reports with 6-Chapter structure"""
    
    # Month order for sorting
    MONTH_ORDER = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
    
    def __init__(self, template_path: str = None, backend_base_url: str = None):
        """Initialize with template path"""
        self.template_path = template_path or os.path.join(
            os.path.dirname(__file__), 'templates', 'GHG_inventory_report.docx'
        )
        self.backend_base_url = backend_base_url or os.environ.get('BACKEND_URL', 'http://localhost:8001')
        self.report_date = datetime.now(timezone.utc).strftime("%B %d, %Y")
    
    # ==================== UTILITY METHODS ====================
    
    def _format_month(self, period_str: str) -> str:
        """Format month string from YYYY-MM to Mon-YYYY format"""
        try:
            if not period_str:
                return 'NA'
            if ' to ' in period_str:
                parts = period_str.split(' to ')
                return f"{self._format_month(parts[0])} to {self._format_month(parts[1])}"
            dt = datetime.strptime(period_str.strip(), "%Y-%m")
            return dt.strftime("%b-%Y")
        except (ValueError, TypeError):
            return period_str or 'NA'
    
    def _format_month_full(self, period_str: str) -> str:
        """Format month string from YYYY-MM to full format (e.g., January 2025)"""
        try:
            if not period_str:
                return 'NA'
            dt = datetime.strptime(period_str.strip(), "%Y-%m")
            return dt.strftime("%B %Y")
        except (ValueError, TypeError):
            return period_str or 'NA'
    
    def _format_reporting_period_with_dates(self, start_period: str, end_period: str) -> str:
        """Format reporting period with specific dates (e.g., '1st March 2023 – 30th April 2024')"""
        try:
            if not start_period or not end_period:
                return 'NA'
            
            # Parse start period
            start_dt = datetime.strptime(start_period.strip(), "%Y-%m")
            start_day = 1
            start_month = start_dt.strftime("%B")
            start_year = start_dt.year
            
            # Parse end period and get last day of month
            end_dt = datetime.strptime(end_period.strip(), "%Y-%m")
            # Get last day of the end month
            if end_dt.month == 12:
                next_month = datetime(end_dt.year + 1, 1, 1)
            else:
                next_month = datetime(end_dt.year, end_dt.month + 1, 1)
            end_day = (next_month - timedelta(days=1)).day
            end_month = end_dt.strftime("%B")
            end_year = end_dt.year
            
            # Helper for ordinal suffix
            def ordinal_suffix(day):
                if 11 <= day <= 13:
                    return 'th'
                return {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
            
            start_str = f"{start_day}{ordinal_suffix(start_day)} {start_month} {start_year}"
            end_str = f"{end_day}{ordinal_suffix(end_day)} {end_month} {end_year}"
            
            return f"{start_str} – {end_str}"
        except (ValueError, TypeError) as e:
            print(f"Error formatting reporting period: {e}")
            return f"{start_period or 'NA'} - {end_period or 'NA'}"
    
    def _format_number(self, value, decimals=2) -> str:
        """Format number to specified decimal places"""
        try:
            if value is None:
                return '0.00'
            num = float(value)
            return f"{num:,.{decimals}f}"
        except (ValueError, TypeError):
            return '0.00'
    
    def _get_value_or_na(self, obj: Dict, key: str, default='NA') -> str:
        """Get value from dict or return NA if empty/None"""
        if obj is None:
            return default
        val = obj.get(key)
        if val is None or val == '' or val == 'NA':
            return default
        return str(val)
    
    def _sort_months(self, months: List[str]) -> List[str]:
        """Sort months in chronological order (Jan -> Dec)"""
        def month_key(month_str):
            try:
                # Handle formats like "Jan-2026" or "2026-01"
                if '-' in month_str:
                    parts = month_str.split('-')
                    if len(parts[0]) == 4:  # YYYY-MM format
                        return datetime.strptime(month_str, "%Y-%m")
                    else:  # Mon-YYYY format
                        return datetime.strptime(month_str, "%b-%Y")
                return datetime.min
            except Exception:
                return datetime.min
        return sorted(months, key=month_key)
    
    def _deduplicate_list(self, items: List[str], case_insensitive: bool = True) -> List[str]:
        """Remove duplicates from list, optionally case-insensitive"""
        seen = set()
        result = []
        for item in items:
            if item is None:
                continue
            key = item.lower().strip() if case_insensitive else item.strip()
            if key not in seen:
                seen.add(key)
                result.append(item.strip())
        return result
    
    def _download_image(self, url: str) -> Optional[io.BytesIO]:
        """Download an image from URL and return as BytesIO"""
        if not url:
            return None
        
        try:
            import re
            
            # Handle internal API file URLs - fetch from R2 directly
            if '/api/files/' in url:
                match = re.search(r'/api/files/([a-f0-9\-]+)', url)
                if match:
                    file_id = match.group(1)
                    
                    try:
                        from r2_storage import get_r2_storage
                        from pymongo import MongoClient
                        
                        # Get file record from database - check BOTH collections
                        mongo_url = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
                        db_name = os.environ.get('DB_NAME', 'test_database')
                        client = MongoClient(mongo_url)
                        db = client[db_name]
                        
                        # Check 'uploaded_files' collection first (primary), then 'files' as fallback
                        file_record = db.uploaded_files.find_one({"id": file_id}, {"_id": 0})
                        if not file_record:
                            file_record = db.files.find_one({"id": file_id}, {"_id": 0})
                        
                        client.close()
                        
                        if file_record:
                            bucket_type = file_record.get('bucket_type')
                            r2_key = file_record.get('r2_key')
                            
                            if bucket_type and r2_key:
                                r2 = get_r2_storage()
                                # Generate presigned URL and download from it
                                presigned_url = r2.generate_presigned_url(
                                    bucket_type=bucket_type,
                                    key=r2_key,
                                    expiration=300  # 5 minutes
                                )
                                
                                response = requests.get(presigned_url, timeout=30)
                                
                                if response.status_code == 200:
                                    content_type = response.headers.get('content-type', '')
                                    if 'image' in content_type.lower() or self._is_image_content(response.content):
                                        return io.BytesIO(response.content)
                            
                    except Exception as e:
                        print(f"Error fetching from R2: {e}")
                    
                    return None
            
            # For external URLs (non-API URLs like direct image links), use HTTP request
            response = requests.get(url, timeout=15, allow_redirects=True)
            
            if response.status_code == 200:
                content_type = response.headers.get('content-type', '')
                if 'image' in content_type.lower() or self._is_image_content(response.content):
                    return io.BytesIO(response.content)
                        
        except Exception as e:
            print(f"Error downloading image from {url}: {e}")
        
        return None
    
    def _is_image_content(self, content: bytes) -> bool:
        """Check if content is an image based on magic bytes"""
        if len(content) < 8:
            return False
        header = content[:8]
        return (header[:4] == b'\x89PNG' or 
                header[:2] == b'\xff\xd8' or 
                header[:6] == b'GIF87a' or 
                header[:6] == b'GIF89a' or
                header[:2] == b'BM' or
                header[:4] == b'RIFF')
    
    def _get_file_path_from_db(self, file_id: str) -> Optional[str]:
        """Get file path from database record"""
        try:
            from pymongo import MongoClient
            mongo_url = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
            db_name = os.environ.get('DB_NAME', 'ghg_platform')
            
            client = MongoClient(mongo_url)
            db = client[db_name]
            
            file_record = db.uploaded_files.find_one({"id": file_id}, {"_id": 0})
            client.close()
            
            if file_record:
                return file_record.get('file_path')
            return None
        except Exception as e:
            print(f"Error getting file path from DB: {e}")
            return None
    
    # ==================== DOCUMENT FORMATTING ====================
    
    def _add_page_border(self, doc: Document):
        """Add blue border to all pages"""
        for section in doc.sections:
            sectPr = section._sectPr
            pgBorders = OxmlElement('w:pgBorders')
            pgBorders.set(qn('w:offsetFrom'), 'page')
            
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '24')  # Border width (3pt = 24 eighths of a point)
                border.set(qn('w:space'), '24')  # Space from page edge
                border.set(qn('w:color'), '1E3A5F')  # Darker blue color
                pgBorders.append(border)
            
            sectPr.append(pgBorders)
    
    def _set_document_font(self, doc: Document):
        """Set default document font to 12pt for all normal text"""
        style = doc.styles['Normal']
        style.font.size = Pt(12)
        style.font.name = 'Calibri'
    
    def _add_footer(self, doc: Document):
        """Add footer with 'Report generated by SustainRepo' to all sections"""
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            
            # Add footer paragraph with SustainRepo attribution
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("Report generated by SustainRepo")
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)  # Gray color
    
    def _add_styled_heading(self, doc: Document, text: str, level: int = 1):
        """Add a styled heading with proper font size hierarchy
        Level 1 (Chapter): 16pt, centered, uppercase
        Level 2 (x.y sections): 14pt, left aligned
        Level 3 (x.y.z subsections): 12pt, left aligned
        """
        if not text:
            text = "Untitled"
        # Check if this is a chapter heading
        is_chapter = text.lower().startswith('chapter')
        
        if is_chapter and level == 1:
            # Chapter headings: centered, uppercase, size 16
            heading = doc.add_heading(text.upper(), level=level)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                run.font.size = Pt(16)
                run.font.bold = True
        elif level == 2:
            # x.y section headings: 14pt, left aligned
            heading = doc.add_heading(text, level=level)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in heading.runs:
                run.font.size = Pt(14)
                run.font.bold = True
        elif level == 3:
            # x.y.z subsection headings: 12pt, left aligned
            heading = doc.add_heading(text, level=level)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in heading.runs:
                run.font.size = Pt(12)
                run.font.bold = True
        else:
            # Regular headings
            heading = doc.add_heading(text, level=level)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return heading
    
    def _add_paragraph_with_bold_label(self, doc: Document, label: str, value: str):
        """Add paragraph with bold label on one line and value on next line"""
        # Label line
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}:")
        run_label.bold = True
        
        # Value line
        value_text = value if value and value != 'NA' else 'NA'
        doc.add_paragraph(value_text)
        return p
    
    def _add_labeled_field(self, doc: Document, label: str, value: str):
        """Add a labeled field with label on one line and value on next line"""
        p = doc.add_paragraph()
        run = p.add_run(f"{label}:")
        run.bold = True
        
        value_text = value if value and value != 'NA' else 'NA'
        doc.add_paragraph(value_text)
        return p
    
    def _add_figure_caption(self, doc: Document, caption_text: str):
        """Add a centered, styled figure caption with light gray color"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption_text)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)  # Light gray color
        return p
    
    def _create_styled_table(self, doc: Document, headers: List[str], data: List[List[str]], 
                            col_widths: List[float] = None, bold_rows: List[int] = None) -> Any:
        """Create a styled table with headers and data
        
        Args:
            bold_rows: List of row indices (0-based, not counting header) that should be bold
        """
        num_cols = len(headers)
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        bold_rows = bold_rows or []
        
        # Set header row
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
            # Set background color for header
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'E8E8E8')
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        
        # Add data rows
        for row_idx, row_data in enumerate(data):
            row_cells = table.add_row().cells
            is_bold_row = row_idx in bold_rows
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data) if cell_data is not None else ''
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        if is_bold_row:
                            run.font.bold = True
        
        # Set column widths if provided
        if col_widths:
            for i, width in enumerate(col_widths):
                for cell in table.columns[i].cells:
                    cell.width = Inches(width)
        
        return table
    
    def _add_emissions_list_table(self, doc: Document, scope1_by_category: Dict[str, List[str]], scope2_processes: List[str], scope3_by_category: Dict[str, List[str]] = None):
        """Create a professionally formatted table for List of Emissions section with merged cells.
        
        Uses 2-column format (Category | Process/Fuel) with scope as section header for BOTH report types.
        """
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        if scope3_by_category is None:
            scope3_by_category = {}
        
        # Check if this is a Scope 3 report
        is_scope3_report = getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3'
        
        # Check if there are any Scope 1 emissions
        has_scope1 = any(scope1_by_category[cat] for cat in scope1_by_category)
        has_scope3 = bool(scope3_by_category) and is_scope3_report
        
        # Use 2-column format for BOTH report types
        self._add_emissions_list_table_2col_extended(doc, scope1_by_category, scope2_processes, scope3_by_category, has_scope1, has_scope3, is_scope3_report)

    def _add_emissions_list_table_4col(self, doc: Document, scope1_by_category: Dict[str, List[Dict]], 
                                        scope2_processes: List[Dict], scope3_by_category: Dict[str, List[Dict]],
                                        has_scope1: bool, has_scope3: bool, is_scope3_report: bool):
        """Create 4-column emissions list table with Person Responsible and Source of Information"""
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Calculate total rows needed
        total_rows = 1  # Header row
        
        # Scope 1 section
        total_rows += 1  # Scope 1 header
        if has_scope1:
            for cat_key in ['stationary_combustion', 'mobile_combustion', 'fugitive_emissions', 'other']:
                if scope1_by_category.get(cat_key):
                    total_rows += len(scope1_by_category[cat_key])
        else:
            total_rows += 1
        
        # Scope 2 section
        total_rows += 1  # Scope 2 header
        if scope2_processes and not (len(scope2_processes) == 1 and scope2_processes[0].get('process_fuel') == 'NA'):
            total_rows += len(scope2_processes)
        else:
            total_rows += 1
        
        # Scope 3 section
        if has_scope3:
            total_rows += 1  # Scope 3 header
            for cat_key in scope3_by_category:
                if scope3_by_category[cat_key]:
                    total_rows += len(scope3_by_category[cat_key])
        
        # Create table with 4 columns
        table = doc.add_table(rows=total_rows, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        col_widths = [Inches(2.0), Inches(3.0), Inches(1.5), Inches(1.5)]
        for col_idx, width in enumerate(col_widths):
            for cell in table.columns[col_idx].cells:
                cell.width = width
        
        current_row = 0
        
        # Header row
        headers = ['Category', 'Process Name – Fuel/Activity/Energy', 'Person Responsible', 'Source of Information']
        for col_idx, header in enumerate(headers):
            cell = table.rows[current_row].cells[col_idx]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '1E3A5F')
            cell._tc.get_or_add_tcPr().append(shading)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        current_row += 1
        
        # Helper to add scope header row (merged across all 4 columns)
        def add_scope_header(row_idx, text):
            row = table.rows[row_idx]
            row.cells[0].merge(row.cells[3])
            cell = row.cells[0]
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D4E6F1')
            cell._tc.get_or_add_tcPr().append(shading)
        
        # Helper to add data rows for a category
        def add_category_rows(start_row, cat_name, items):
            nonlocal current_row
            cat_start_row = start_row
            for idx, item in enumerate(items):
                row = table.rows[current_row]
                if idx == 0:
                    row.cells[0].text = cat_name
                    for paragraph in row.cells[0].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                row.cells[1].text = item.get('process_fuel', '')
                row.cells[2].text = item.get('responsible_person', '')
                row.cells[3].text = item.get('record_source', '')
                for col in range(1, 4):
                    for paragraph in row.cells[col].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                current_row += 1
            
            # Merge category cells if more than one row
            if len(items) > 1:
                start_cell = table.rows[cat_start_row].cells[0]
                end_cell = table.rows[cat_start_row + len(items) - 1].cells[0]
                start_cell.merge(end_cell)
                start_cell.text = cat_name
                start_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in start_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
        
        # Scope 1 Header
        add_scope_header(current_row, "Scope 1 Emissions")
        current_row += 1
        
        # Scope 1 Categories
        if has_scope1:
            for cat_key, cat_name in [('stationary_combustion', 'Stationary Combustion'), 
                                       ('mobile_combustion', 'Mobile Combustion'),
                                       ('fugitive_emissions', 'Fugitive Emissions'),
                                       ('other', 'Other')]:
                items = scope1_by_category.get(cat_key, [])
                if items:
                    add_category_rows(current_row, cat_name, items)
        else:
            row = table.rows[current_row]
            row.cells[0].merge(row.cells[3])
            row.cells[0].text = "No emission reported"
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.italic = True
                    run.font.size = Pt(10)
            current_row += 1
        
        # Scope 2 Header
        add_scope_header(current_row, "Scope 2 Emissions")
        current_row += 1
        
        # Scope 2 Processes
        has_scope2_data = scope2_processes and not (len(scope2_processes) == 1 and scope2_processes[0].get('process_fuel') == 'NA')
        if has_scope2_data:
            add_category_rows(current_row, "Purchased Energy", scope2_processes)
        else:
            row = table.rows[current_row]
            row.cells[0].merge(row.cells[3])
            row.cells[0].text = "No emission reported"
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.italic = True
                    run.font.size = Pt(10)
            current_row += 1
        
        # Scope 3 Section
        if has_scope3:
            add_scope_header(current_row, "Scope 3 Emissions")
            current_row += 1
            
            scope3_display_names = {
                'c1': 'C1 - Purchased Goods and Services', 'c2': 'C2 - Capital Goods',
                'c3': 'C3 - Fuel and Energy Related Activities', 'c4': 'C4 - Upstream Transportation and Distribution',
                'c5': 'C5 - Waste Generated in Operations', 'c6': 'C6 - Business Travel',
                'c7': 'C7 - Employee Commuting', 'c8': 'C8 - Upstream Leased Assets',
                'c9': 'C9 - Downstream Transportation and Distribution', 'c10': 'C10 - Processing of Sold Products',
                'c11': 'C11 - Use of Sold Products', 'c12': 'C12 - End-of-Life Treatment of Sold Products',
                'c13': 'C13 - Downstream Leased Assets', 'c14': 'C14 - Franchises', 'c15': 'C15 - Investments',
            }
            
            for cat_key in sorted(scope3_by_category.keys(), key=lambda x: int(x[1:]) if x.startswith('c') and x[1:].isdigit() else 99):
                items = scope3_by_category[cat_key]
                if items:
                    cat_display = scope3_display_names.get(cat_key, cat_key)
                    add_category_rows(current_row, cat_display, items)
        
        doc.add_paragraph()
    
    def _add_emissions_list_table_2col_extended(self, doc: Document, scope1_by_category: Dict[str, List[str]], 
                                                 scope2_processes: List[str], scope3_by_category: Dict[str, List[str]],
                                                 has_scope1: bool, has_scope3: bool, is_scope3_report: bool):
        """Create 2-column emissions list table (Category | Process/Fuel) for all report types"""
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Calculate total rows needed
        total_rows = 1  # Header row
        
        # Scope 1 section
        total_rows += 1  # Scope 1 header
        if has_scope1:
            for cat_key, cat_name in [('stationary_combustion', 'Stationary Combustion'), 
                                       ('mobile_combustion', 'Mobile Combustion'),
                                       ('fugitive_emissions', 'Fugitive Emissions'),
                                       ('other', 'Other')]:
                if scope1_by_category[cat_key]:
                    total_rows += len(scope1_by_category[cat_key])
        else:
            total_rows += 1  # "No emission reported" row
        
        # Scope 2 section
        total_rows += 1  # Scope 2 header
        if scope2_processes and scope2_processes != ["NA"]:
            total_rows += len(scope2_processes)
        else:
            total_rows += 1  # "No emission reported" row
        
        # Scope 3 section (only if has_scope3)
        if has_scope3:
            total_rows += 1  # Scope 3 header
            for cat_key in scope3_by_category:
                if scope3_by_category[cat_key]:
                    total_rows += len(scope3_by_category[cat_key])
        
        # Create table with 2 columns
        table = doc.add_table(rows=total_rows, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        for cell in table.columns[0].cells:
            cell.width = Inches(2.5)
        for cell in table.columns[1].cells:
            cell.width = Inches(4.5)
        
        current_row = 0
        
        # Header row - adapt based on report type
        if is_scope3_report:
            headers = ['Category', 'Process Name – Fuel/Activity/Energy']
        else:
            headers = ['Category', 'Process / Fuel']
        for col_idx, header in enumerate(headers):
            cell = table.rows[current_row].cells[col_idx]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
            # Header background - dark blue
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '1E3A5F')
            cell._tc.get_or_add_tcPr().append(shading)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        current_row += 1
        
        # Scope 1 Header Row - merge both cells
        scope1_header_row = table.rows[current_row]
        scope1_header_row.cells[0].merge(scope1_header_row.cells[1])
        cell = scope1_header_row.cells[0]
        cell.text = "Direct/Scope 1 Emissions" if not is_scope3_report else "Scope 1 Emissions"
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(12)
        # Light blue background for scope header
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D4E6F1')
        cell._tc.get_or_add_tcPr().append(shading)
        
        current_row += 1
        
        # Scope 1 Categories
        if has_scope1:
            for cat_key, cat_name in [('stationary_combustion', 'Stationary Combustion'), 
                                       ('mobile_combustion', 'Mobile Combustion'),
                                       ('fugitive_emissions', 'Fugitive Emissions'),
                                       ('other', 'Other')]:
                processes = scope1_by_category[cat_key]
                if not processes:
                    continue
                
                cat_start_row = current_row
                
                for idx, process in enumerate(processes):
                    row = table.rows[current_row]
                    
                    # Category column - only fill for first row, will merge later
                    if idx == 0:
                        row.cells[0].text = cat_name
                        for paragraph in row.cells[0].paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(12)
                    
                    # Process/Fuel column
                    row.cells[1].text = process
                    for paragraph in row.cells[1].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(12)
                    
                    current_row += 1
                
                # Merge category cells if more than one process
                if len(processes) > 1:
                    start_cell = table.rows[cat_start_row].cells[0]
                    end_cell = table.rows[cat_start_row + len(processes) - 1].cells[0]
                    start_cell.merge(end_cell)
                    # Re-apply formatting after merge
                    start_cell.text = cat_name
                    start_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for paragraph in start_cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(12)
        else:
            # No emissions row
            row = table.rows[current_row]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = "No emission reported"
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.italic = True
                    run.font.size = Pt(12)
            current_row += 1
        
        # Scope 2 Header Row - merge both cells
        scope2_header_row = table.rows[current_row]
        scope2_header_row.cells[0].merge(scope2_header_row.cells[1])
        cell = scope2_header_row.cells[0]
        cell.text = "Indirect/Scope 2 Emissions" if not is_scope3_report else "Scope 2 Emissions"
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(12)
        # Light blue background for scope header
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D4E6F1')
        cell._tc.get_or_add_tcPr().append(shading)
        
        current_row += 1
        
        # Scope 2 Processes
        if scope2_processes and scope2_processes != ["NA"]:
            scope2_start_row = current_row
            for idx, process in enumerate(scope2_processes):
                row = table.rows[current_row]
                if idx == 0:
                    row.cells[0].text = "Purchased Energy"
                    for paragraph in row.cells[0].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(12)
                row.cells[1].text = process
                for paragraph in row.cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                current_row += 1
            
            # Merge category cells if more than one process
            if len(scope2_processes) > 1:
                start_cell = table.rows[scope2_start_row].cells[0]
                end_cell = table.rows[scope2_start_row + len(scope2_processes) - 1].cells[0]
                start_cell.merge(end_cell)
                start_cell.text = "Purchased Energy"
                start_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in start_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(12)
        else:
            # No emissions row
            row = table.rows[current_row]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = "No emission reported"
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.italic = True
                    run.font.size = Pt(12)
            current_row += 1
        
        # Scope 3 Section - Only if has_scope3
        if has_scope3:
            # Scope 3 Header Row - merge both cells
            scope3_header_row = table.rows[current_row]
            scope3_header_row.cells[0].merge(scope3_header_row.cells[1])
            cell = scope3_header_row.cells[0]
            cell.text = "Scope 3 Emissions"
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
            # Light blue background for scope 3 header
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D4E6F1')
            cell._tc.get_or_add_tcPr().append(shading)
            
            current_row += 1
            
            # Scope 3 Category display names
            scope3_display_names = getattr(self, 'scope3_category_display', {
                'c1': 'C1 - Purchased Goods and Services',
                'c2': 'C2 - Capital Goods',
                'c3': 'C3 - Fuel and Energy Related Activities Not Included in Scope 1 or Scope 2',
                'c4': 'C4 - Upstream Transportation and Distribution',
                'c5': 'C5 - Waste Generated in Operations',
                'c6': 'C6 - Business Travel',
                'c7': 'C7 - Employee Commuting',
                'c8': 'C8 - Upstream Leased Assets',
                'c9': 'C9 - Downstream Transportation and Distribution',
                'c10': 'C10 - Processing of Sold Products',
                'c11': 'C11 - Use of Sold Products',
                'c12': 'C12 - End-of-Life Treatment of Sold Products',
                'c13': 'C13 - Downstream Leased Assets',
                'c14': 'C14 - Franchises',
                'c15': 'C15 - Investments',
                'other': 'Other Scope 3'
            })
            
            # Add Scope 3 categories
            for cat_key in sorted(scope3_by_category.keys(), key=lambda x: int(x[1:]) if x.startswith('c') and x[1:].isdigit() else 99):
                processes = scope3_by_category[cat_key]
                if not processes:
                    continue
                
                cat_display_name = scope3_display_names.get(cat_key, cat_key)
                cat_start_row = current_row
                
                for idx, process in enumerate(processes):
                    row = table.rows[current_row]
                    
                    # Category column - only fill for first row, will merge later
                    if idx == 0:
                        row.cells[0].text = cat_display_name
                        for paragraph in row.cells[0].paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(12)
                    
                    # Process/Activity column
                    row.cells[1].text = process
                    for paragraph in row.cells[1].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(12)
                    
                    current_row += 1
                
                # Merge category cells if more than one process
                if len(processes) > 1:
                    start_cell = table.rows[cat_start_row].cells[0]
                    end_cell = table.rows[cat_start_row + len(processes) - 1].cells[0]
                    start_cell.merge(end_cell)
                    # Re-apply formatting after merge
                    start_cell.text = cat_display_name
                    start_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for paragraph in start_cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(12)
        
        doc.add_paragraph()  # Add spacing after table
    
    def _add_emissions_list_table_2col(self, doc: Document, scope1_by_category: Dict[str, List[str]], 
                                        scope2_processes: List[str], has_scope1: bool):
        """Create 2-column emissions list table for Scope 1,2 reports (Category | Process/Fuel)"""
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Calculate total rows needed
        total_rows = 1  # Header row
        
        # Scope 1 section
        total_rows += 1  # Scope 1 header
        if has_scope1:
            for cat_key, cat_name in [('stationary_combustion', 'Stationary Combustion'), 
                                       ('mobile_combustion', 'Mobile Combustion'),
                                       ('fugitive_emissions', 'Fugitive Emissions'),
                                       ('other', 'Other')]:
                if scope1_by_category[cat_key]:
                    total_rows += len(scope1_by_category[cat_key])
        else:
            total_rows += 1  # "No emission reported" row
        
        # Scope 2 section
        total_rows += 1  # Scope 2 header
        if scope2_processes and scope2_processes != ["NA"]:
            total_rows += len(scope2_processes)
        else:
            total_rows += 1  # "No emission reported" row
        
        # Create table with 2 columns
        table = doc.add_table(rows=total_rows, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        for cell in table.columns[0].cells:
            cell.width = Inches(2.5)
        for cell in table.columns[1].cells:
            cell.width = Inches(4.5)
        
        current_row = 0
        
        # Header row
        headers = ['Category', 'Process / Fuel']
        for col_idx, header in enumerate(headers):
            cell = table.rows[current_row].cells[col_idx]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
            # Header background - dark blue
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '1E3A5F')
            cell._tc.get_or_add_tcPr().append(shading)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        current_row += 1
        
        # Scope 1 Header Row - merge both cells
        scope1_header_row = table.rows[current_row]
        scope1_header_row.cells[0].merge(scope1_header_row.cells[1])
        cell = scope1_header_row.cells[0]
        cell.text = "Direct/Scope 1 Emissions"
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(12)
        # Light blue background for scope header
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D4E6F1')
        cell._tc.get_or_add_tcPr().append(shading)
        
        current_row += 1
        
        # Scope 1 Categories
        if has_scope1:
            for cat_key, cat_name in [('stationary_combustion', 'Stationary Combustion'), 
                                       ('mobile_combustion', 'Mobile Combustion'),
                                       ('fugitive_emissions', 'Fugitive Emissions'),
                                       ('other', 'Other')]:
                processes = scope1_by_category[cat_key]
                if not processes:
                    continue
                
                cat_start_row = current_row
                
                for idx, process in enumerate(processes):
                    row = table.rows[current_row]
                    
                    # Category column - only fill for first row, will merge later
                    if idx == 0:
                        row.cells[0].text = cat_name
                        for paragraph in row.cells[0].paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(12)
                    
                    # Process/Fuel column
                    row.cells[1].text = process
                    for paragraph in row.cells[1].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(12)
                    
                    current_row += 1
                
                # Merge category cells if more than one process
                if len(processes) > 1:
                    start_cell = table.rows[cat_start_row].cells[0]
                    end_cell = table.rows[cat_start_row + len(processes) - 1].cells[0]
                    start_cell.merge(end_cell)
                    # Re-apply formatting after merge
                    start_cell.text = cat_name
                    start_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for paragraph in start_cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(12)
        else:
            # No emissions row
            row = table.rows[current_row]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = "No emission reported"
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.italic = True
                    run.font.size = Pt(12)
            current_row += 1
        
        # Scope 2 Header Row - merge both cells
        scope2_header_row = table.rows[current_row]
        scope2_header_row.cells[0].merge(scope2_header_row.cells[1])
        cell = scope2_header_row.cells[0]
        cell.text = "Indirect/Scope 2 Emissions"
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(12)
        # Light blue background for scope header
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D4E6F1')
        cell._tc.get_or_add_tcPr().append(shading)
        
        current_row += 1
        
        # Scope 2 Processes
        if scope2_processes and scope2_processes != ["NA"]:
            scope2_start_row = current_row
            for idx, process in enumerate(scope2_processes):
                row = table.rows[current_row]
                if idx == 0:
                    row.cells[0].text = "Purchased Energy"
                    for paragraph in row.cells[0].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(12)
                row.cells[1].text = process
                for paragraph in row.cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                current_row += 1
            
            # Merge category cells if more than one process
            if len(scope2_processes) > 1:
                start_cell = table.rows[scope2_start_row].cells[0]
                end_cell = table.rows[scope2_start_row + len(scope2_processes) - 1].cells[0]
                start_cell.merge(end_cell)
                start_cell.text = "Purchased Energy"
                start_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in start_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(12)
        else:
            # No emissions row
            row = table.rows[current_row]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = "No emission reported"
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.italic = True
                    run.font.size = Pt(12)
            current_row += 1
        
        doc.add_paragraph()  # Add spacing after table
    
    def _add_emissions_list_table_3col(self, doc: Document, scope1_by_category: Dict[str, List[str]], 
                                        scope2_processes: List[str], scope3_by_category: Dict[str, List[str]],
                                        has_scope1: bool, has_scope3: bool):
        """Create 3-column emissions list table for Scope 1,2,3 reports (Scope | Category | Process)"""
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Calculate total rows needed
        total_rows = 1  # Header row
        
        # Scope 1 section
        total_rows += 1  # Scope 1 header
        if has_scope1:
            for cat_key, cat_name in [('stationary_combustion', 'Stationary Combustion'), 
                                       ('mobile_combustion', 'Mobile Combustion'),
                                       ('fugitive_emissions', 'Fugitive Emissions'),
                                       ('other', 'Other')]:
                if scope1_by_category[cat_key]:
                    total_rows += len(scope1_by_category[cat_key])
        else:
            total_rows += 1  # "No emission reported" row
        
        # Scope 2 section
        total_rows += 1  # Scope 2 header
        if scope2_processes and scope2_processes != ["NA"]:
            total_rows += len(scope2_processes)
        else:
            total_rows += 1  # "No emission reported" row
        
        # Scope 3 section
        if has_scope3:
            total_rows += 1  # Scope 3 header
            for cat_key in scope3_by_category:
                total_rows += len(scope3_by_category[cat_key])
        
        # Create table with 3 columns
        table = doc.add_table(rows=total_rows, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        for cell in table.columns[0].cells:
            cell.width = Inches(2.0)
        for cell in table.columns[1].cells:
            cell.width = Inches(2.0)
        for cell in table.columns[2].cells:
            cell.width = Inches(3.0)
        
        current_row = 0
        
        # Header row
        headers = ['Scope', 'Category', 'Process Name – Fuel/Activity/Energy']
        for col_idx, header in enumerate(headers):
            cell = table.rows[current_row].cells[col_idx]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
            # Header background
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '1E3A5F')  # Dark blue
            cell._tc.get_or_add_tcPr().append(shading)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        
        current_row += 1
        
        # Scope 1 Header Row - merge all 3 cells
        scope1_header_row = table.rows[current_row]
        scope1_header_row.cells[0].merge(scope1_header_row.cells[2])
        cell = scope1_header_row.cells[0]
        cell.text = "Scope 1 Emissions"
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(12)
        # Light blue background for scope header
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D4E6F1')
        cell._tc.get_or_add_tcPr().append(shading)
        
        current_row += 1
        
        # Scope 1 Categories
        if has_scope1:
            for cat_key, cat_name in [('stationary_combustion', 'Stationary Combustion'), 
                                       ('mobile_combustion', 'Mobile Combustion'),
                                       ('fugitive_emissions', 'Fugitive Emissions'),
                                       ('other', 'Other')]:
                processes = scope1_by_category[cat_key]
                if not processes:
                    continue
                
                cat_start_row = current_row
                
                for idx, process in enumerate(processes):
                    row = table.rows[current_row]
                    
                    # First column is empty (Scope already shown in header)
                    row.cells[0].text = ""
                    
                    # Category column - only fill for first row, will merge later
                    if idx == 0:
                        row.cells[1].text = cat_name
                        for paragraph in row.cells[1].paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(12)
                    
                    # Process/Fuel column
                    row.cells[2].text = process
                    for paragraph in row.cells[2].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(12)
                    
                    current_row += 1
                
                # Merge category cells if more than one process
                if len(processes) > 1:
                    start_cell = table.rows[cat_start_row].cells[1]
                    end_cell = table.rows[cat_start_row + len(processes) - 1].cells[1]
                    start_cell.merge(end_cell)
                    # Re-apply formatting after merge
                    start_cell.text = cat_name
                    start_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for paragraph in start_cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(12)
        else:
            # No emissions row
            row = table.rows[current_row]
            row.cells[0].text = ""
            row.cells[1].merge(row.cells[2])
            row.cells[1].text = "No emission reported"
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.italic = True
                    run.font.size = Pt(12)
            current_row += 1
        
        # Scope 2 Header Row - merge all 3 cells
        scope2_header_row = table.rows[current_row]
        scope2_header_row.cells[0].merge(scope2_header_row.cells[2])
        cell = scope2_header_row.cells[0]
        cell.text = "Scope 2 Emissions"
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(12)
        # Light blue background for scope header
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D4E6F1')
        cell._tc.get_or_add_tcPr().append(shading)
        
        current_row += 1
        
        # Scope 2 Processes
        if scope2_processes and scope2_processes != ["NA"]:
            scope2_start_row = current_row
            for idx, process in enumerate(scope2_processes):
                row = table.rows[current_row]
                row.cells[0].text = ""
                if idx == 0:
                    row.cells[1].text = "Purchased Energy"
                    for paragraph in row.cells[1].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(12)
                row.cells[2].text = process
                for paragraph in row.cells[2].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                current_row += 1
            
            # Merge category cells if more than one process
            if len(scope2_processes) > 1:
                start_cell = table.rows[scope2_start_row].cells[1]
                end_cell = table.rows[scope2_start_row + len(scope2_processes) - 1].cells[1]
                start_cell.merge(end_cell)
                start_cell.text = "Purchased Energy"
                start_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in start_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(12)
        else:
            # No emissions row
            row = table.rows[current_row]
            row.cells[0].text = ""
            row.cells[1].merge(row.cells[2])
            row.cells[1].text = "No emission reported"
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.italic = True
                    run.font.size = Pt(12)
            current_row += 1
        
        # Scope 3 Section - Only if has_scope3
        if has_scope3:
            # Scope 3 Header Row - merge all 3 cells
            scope3_header_row = table.rows[current_row]
            scope3_header_row.cells[0].merge(scope3_header_row.cells[2])
            cell = scope3_header_row.cells[0]
            cell.text = "Scope 3 Emissions"
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
            # Light blue background for scope 3 header
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D4E6F1')
            cell._tc.get_or_add_tcPr().append(shading)
            
            current_row += 1
            
            # Scope 3 Category display names
            scope3_display_names = getattr(self, 'scope3_category_display', {
                'c1': 'C1 - Purchased Goods and Services',
                'c2': 'C2 - Capital Goods',
                'c3': 'C3 - Fuel and Energy Related Activities Not Included in Scope 1 or Scope 2',
                'c4': 'C4 - Upstream Transportation and Distribution',
                'c5': 'C5 - Waste Generated in Operations',
                'c6': 'C6 - Business Travel',
                'c7': 'C7 - Employee Commuting',
                'c8': 'C8 - Upstream Leased Assets',
                'c9': 'C9 - Downstream Transportation and Distribution',
                'c10': 'C10 - Processing of Sold Products',
                'c11': 'C11 - Use of Sold Products',
                'c12': 'C12 - End-of-Life Treatment of Sold Products',
                'c13': 'C13 - Downstream Leased Assets',
                'c14': 'C14 - Franchises',
                'c15': 'C15 - Investments',
                'other': 'Other Scope 3'
            })
            
            # Add Scope 3 categories
            for cat_key in sorted(scope3_by_category.keys(), key=lambda x: int(x[1:]) if x.startswith('c') and x[1:].isdigit() else 99):
                processes = scope3_by_category[cat_key]
                if not processes:
                    continue
                
                cat_display_name = scope3_display_names.get(cat_key, cat_key)
                cat_start_row = current_row
                
                for idx, process in enumerate(processes):
                    row = table.rows[current_row]
                    
                    # First column is empty (Scope already shown in header)
                    row.cells[0].text = ""
                    
                    # Category column - only fill for first row, will merge later
                    if idx == 0:
                        row.cells[1].text = cat_display_name
                        for paragraph in row.cells[1].paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(12)
                    
                    # Process/Activity column
                    row.cells[2].text = process
                    for paragraph in row.cells[2].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(12)
                    
                    current_row += 1
                
                # Merge category cells if more than one process
                if len(processes) > 1:
                    start_cell = table.rows[cat_start_row].cells[1]
                    end_cell = table.rows[cat_start_row + len(processes) - 1].cells[1]
                    start_cell.merge(end_cell)
                    # Re-apply formatting after merge
                    start_cell.text = cat_display_name
                    start_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for paragraph in start_cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(12)
        
        doc.add_paragraph()  # Add spacing after table
    
    def _add_process_overview_table(self, doc: Document, facility_emissions: List[Dict]):
        """Create a table showing unique processes with their descriptions"""
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Extract unique processes with descriptions
        # Note: Only includes Scope 1 and Scope 2 emissions (excludes biogenic)
        unique_processes = {}
        
        for em in facility_emissions:
            # Skip biogenic emissions to match List of Emissions table
            scope = (em.get('scope') or '').lower()
            if 'biogenic' in scope or 'bio' in scope:
                continue
                
            # Get process descriptions from the emission record
            process_descriptions = em.get('process_descriptions', [])
            process_names = em.get('process_names', [])
            
            # If we have process_descriptions (new format)
            if process_descriptions:
                for pd in process_descriptions:
                    name = pd.get('name', '').strip()
                    desc = pd.get('description', '').strip()
                    # if name and name not in unique_processes:
                    #     unique_processes[name] = desc
                    if name:
                        key = (name, desc)  # 👈 composite uniqueness
                        unique_processes[key] = None
            # Fallback to process_names (old format - no description)
            elif process_names:
                for name in process_names:
                    if isinstance(name, str):
                        name = name.strip()
                        # if name and name not in unique_processes:
                        #     unique_processes[name] = ''
                        key = (name, '')  # no description fallback
                        unique_processes[key] = None
        
        if not unique_processes:
            p = doc.add_paragraph()
            run = p.add_run("No process information available.")
            run.italic = True
            return
        
        # Create table
        table = doc.add_table(rows=len(unique_processes) + 1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        for cell in table.columns[0].cells:
            cell.width = Inches(2.5)
        for cell in table.columns[1].cells:
            cell.width = Inches(4.5)
        
        # Header row
        headers = ['Process Name', 'Description']
        for col_idx, header in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
            # Header background
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '1E3A5F')  # Dark blue
            cell._tc.get_or_add_tcPr().append(shading)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        
        # Data rows
        # for row_idx, (process_name, description) in enumerate(unique_processes.items(), 1):
        #     row = table.rows[row_idx]
            
        #     # Process name cell
        #     row.cells[0].text = process_name
        #     for paragraph in row.cells[0].paragraphs:
        #         for run in paragraph.runs:
        #             run.font.bold = True
        #             run.font.size = Pt(12)
            
        #     # Description cell
        #     row.cells[1].text = description if description else '-'
        #     for paragraph in row.cells[1].paragraphs:
        #         for run in paragraph.runs:
        #             run.font.size = Pt(12)
        #             if not description:
        #                 run.font.italic = True
        for row_idx, ((process_name, description), _) in enumerate(unique_processes.items(), 1):
            row = table.rows[row_idx]

            row.cells[0].text = process_name
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)

            row.cells[1].text = description if description else '-'
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    if not description:
                        run.font.italic = True
                
        doc.add_paragraph()  # Add spacing after table
    
    # ==================== DATA PROCESSING ====================
    
    def _filter_emissions_by_period(self, emissions: List[Dict], start_period: str, end_period: str) -> List[Dict]:
        """Filter emissions to only include records within the reporting period.
        
        For yearly records (CY/FY), only include if there's actual overlap with the target period.
        """
        if not emissions:
            return []
        
        def normalize_period_for_comparison(period: str) -> str:
            """Normalize period to YYYY-MM format for comparison"""
            if not period:
                return None
            period = period.strip()
            # YYYY-MM already
            if len(period) == 7 and period[4] == '-':
                return period
            # YYYY format - use January
            if len(period) == 4:
                return f"{period}-01"
            return period
        
        filtered = []
        for em in emissions:
            period = em.get('reporting_period') or ''
            if not period:
                continue
            
            frequency_type = em.get('frequency_type', 'monthly')
            
            # Handle yearly records (CY2025, FY 2025-2026, FY 2025-26 formats)
            if frequency_type == 'yearly':
                # Use proration factor to check if there's any overlap
                factor, _ = self._calculate_proration_factor(period, start_period, end_period)
                if factor > 0:
                    filtered.append(em)
                continue
            
            # Handle single month or range (original logic for monthly records)
            if ' to ' in period:
                em_start, em_end = period.split(' to ')
            else:
                em_start = em_end = period
            
            em_start = em_start.strip()
            em_end = em_end.strip()
            
            # Normalize periods for comparison
            em_start_normalized = normalize_period_for_comparison(em_start)
            em_end_normalized = normalize_period_for_comparison(em_end)
            
            if em_start_normalized and em_end_normalized:
                # Check if emission period overlaps with reporting period
                if em_start_normalized <= end_period and em_end_normalized >= start_period:
                    filtered.append(em)
        
        return filtered
    
    def _calculate_proration_factor(self, emission_period: str, target_start: str, target_end: str) -> Tuple[float, bool]:
        """
        Calculate the proration factor for a yearly emission record based on the overlap
        with the target reporting period.
        
        Args:
            emission_period: The emission's reporting period (e.g., "CY 2025", "FY 2024-2025")
            target_start: Target reporting period start (e.g., "2024-04" for April 2024)
            target_end: Target reporting period end (e.g., "2025-03" for March 2025)
        
        Returns:
            Tuple of (proration_factor, is_prorated)
            - proration_factor: float between 0 and 1 representing the fraction to apply
            - is_prorated: True if proration was applied (not full overlap)
        
        Example for FY 2024-2025 (April 2024 - March 2025):
        - CY 2025 data: 3 months overlap (Jan-Mar) → factor = 3/12 = 0.25
        - CY 2024 data: 9 months overlap (Apr-Dec) → factor = 9/12 = 0.75  
        - FY 2024-2025 data: Full 12 months → factor = 1.0
        """
        if not emission_period or not target_start or not target_end:
            return 1.0, False
        
        emission_period = emission_period.strip()
        
        # Parse target period start and end as months
        try:
            target_start_year = int(target_start[:4])
            target_start_month = int(target_start[5:7]) if len(target_start) >= 7 else 1
            target_end_year = int(target_end[:4])
            target_end_month = int(target_end[5:7]) if len(target_end) >= 7 else 12
        except (ValueError, IndexError):
            return 1.0, False
        
        # Calculate target period months (list of (year, month) tuples)
        target_months = set()
        current_year, current_month = target_start_year, target_start_month
        while (current_year, current_month) <= (target_end_year, target_end_month):
            target_months.add((current_year, current_month))
            current_month += 1
            if current_month > 12:
                current_month = 1
                current_year += 1
        
        # Parse emission period to get its months
        emission_months = set()
        
        if emission_period.startswith("CY") or emission_period.startswith("CY "):
            # Calendar Year: CY 2025 or CY2025 → Jan-Dec of that year
            try:
                cy_year = int(emission_period.replace("CY", "").strip()[:4])
                for month in range(1, 13):
                    emission_months.add((cy_year, month))
            except (ValueError, IndexError):
                return 1.0, False
                
        elif emission_period.startswith("FY") or emission_period.startswith("FY "):
            # Financial Year: FY 2024-2025 or FY 2024-25 → Apr of first year to Mar of second year
            try:
                fy_part = emission_period.replace("FY", "").strip()
                if "-" in fy_part:
                    years = fy_part.split("-")
                    fy_start_year = int(years[0].strip())
                    end_year_str = years[1].strip()
                    if len(end_year_str) == 2:
                        fy_end_year = int(f"{str(fy_start_year)[:2]}{end_year_str}")
                    else:
                        fy_end_year = int(end_year_str)
                else:
                    fy_start_year = int(fy_part[:4])
                    fy_end_year = fy_start_year + 1
                
                # FY runs April to March
                # Apr-Dec of start year
                for month in range(4, 13):
                    emission_months.add((fy_start_year, month))
                # Jan-Mar of end year
                for month in range(1, 4):
                    emission_months.add((fy_end_year, month))
            except (ValueError, IndexError):
                return 1.0, False
        else:
            # Not a yearly period format - no proration needed
            return 1.0, False
        
        # Calculate overlap
        if not emission_months:
            return 1.0, False
        
        overlap_months = emission_months.intersection(target_months)
        total_emission_months = len(emission_months)
        overlap_count = len(overlap_months)
        
        if overlap_count == 0:
            return 0.0, True
        
        proration_factor = overlap_count / total_emission_months
        is_prorated = proration_factor < 1.0
        
        return proration_factor, is_prorated
    
    def _apply_proration_to_emissions(self, emissions: List[Dict], target_start: str, target_end: str) -> Tuple[List[Dict], bool]:
        """
        Apply proration to yearly emissions based on the target reporting period.
        Monthly emissions within the period are not prorated.
        
        Returns:
            Tuple of (prorated_emissions_list, has_any_prorated_data)
        """
        if not emissions:
            return [], False
        
        prorated_emissions = []
        has_prorated = False
        
        for em in emissions:
            em_copy = dict(em)
            frequency_type = em.get('frequency_type', 'monthly')
            
            if frequency_type == 'yearly':
                period = em.get('reporting_period', '')
                factor, is_prorated = self._calculate_proration_factor(period, target_start, target_end)
                
                if is_prorated:
                    has_prorated = True
                    # Mark as prorated
                    em_copy['_is_prorated'] = True
                    em_copy['_proration_factor'] = factor
                    
                    # Apply proration to emission values
                    original_total = float(em.get('total_emissions', 0) or em.get('co2e_emissions', 0) or 0)
                    em_copy['total_emissions'] = original_total * factor
                    em_copy['_original_total_emissions'] = original_total
                    
                    # Also prorate individual gas components if present
                    for gas_key in ['co2_emissions', 'ch4_emissions', 'n2o_emissions', 'co2e_emissions']:
                        if gas_key in em and em[gas_key]:
                            original_val = float(em[gas_key] or 0)
                            em_copy[gas_key] = original_val * factor
                            em_copy[f'_original_{gas_key}'] = original_val
                else:
                    em_copy['_is_prorated'] = False
            else:
                # Monthly records - no proration
                em_copy['_is_prorated'] = False
            
            prorated_emissions.append(em_copy)
        
        return prorated_emissions, has_prorated
    
    def _deduplicate_emissions(self, emissions: List[Dict]) -> List[Dict]:
        """
        Deduplicate emissions based on exact record ID only.
        
        Previously this method removed monthly records when yearly records existed for the
        same category/year, but this was too aggressive - monthly and yearly records for
        the same category can represent different activities/employees and should both be included.
        
        Now we only deduplicate exact duplicates (same record ID).
        """
        if not emissions:
            return []
        
        # Only deduplicate exact duplicates by record ID
        seen_ids = set()
        deduped = []
        for e in emissions:
            record_id = e.get('id')
            if record_id and record_id in seen_ids:
                continue
            if record_id:
                seen_ids.add(record_id)
            deduped.append(e)
        
        return deduped
    
    def _get_emissions_by_facility(self, emissions: List[Dict], facility_id: str) -> List[Dict]:
        """Get emissions for a specific facility"""
        return [em for em in emissions if em.get('facility_id') == facility_id]
    
    def _get_fuel_from_emission(self, em: Dict) -> str:
        """Get fuel/activity name from emission record, checking multiple possible fields.
        For Scope 3, prioritize scope3_activity over sub_category."""
        scope = (em.get('scope') or '').lower()
        
        # For Scope 3 emissions, prioritize scope3_activity field
        if 'scope3' in scope or 'scope 3' in scope or scope == '3':
            # Check scope3_activity first (updated when user changes activity)
            scope3_activity = em.get('scope3_activity')
            if scope3_activity:
                return scope3_activity
            # Also check dynamic_field_values for scope3_activity
            dynamic_fields = em.get('dynamic_field_values', {})
            if dynamic_fields:
                scope3_act = dynamic_fields.get('scope3_activity', {})
                if isinstance(scope3_act, dict) and scope3_act.get('value'):
                    return scope3_act.get('value')
        
        return (em.get('fuel_type') or em.get('fuel') or 
                em.get('sub_category') or 'Unknown')
    
    def _get_category_from_emission(self, em: Dict) -> str:
        """Get category from emission record, checking multiple possible fields"""
        return (em.get('category') or em.get('emission_category') or 'Unknown')
    
    def _get_process_names_from_emission(self, em: Dict) -> List[str]:
        """Get process names from emission record"""
        process_names = em.get('process_names', [])
        if process_names:
            return process_names if isinstance(process_names, list) else [process_names]
        
        # Fallback to single process_name field
        process = em.get('process_name') or em.get('name_of_process')
        return [process] if process else []
    
    def _calculate_facility_totals(self, facility_emissions: List[Dict], facility_id: str = None) -> Dict:
        """Calculate totals for a facility, including sinks deduction"""
        totals = {
            'scope1': 0.0,
            'scope2': 0.0,
            'scope3': 0.0,
            'biogenic': 0.0,
            'removals': 0.0,
            'by_month': defaultdict(lambda: {'scope1': 0.0, 'scope2': 0.0, 'scope3': 0.0}),
            'by_category': defaultdict(float),
            'by_fuel': defaultdict(float),
            'by_category_fuel': defaultdict(lambda: defaultdict(float)),  # {category: {fuel: emissions}}
            'by_scope_category_fuel': defaultdict(lambda: defaultdict(lambda: defaultdict(float))),  # {scope: {category: {fuel: emissions}}}
            'scope1_co2': 0.0,
            'scope1_ch4': 0.0,
            'scope1_n2o': 0.0,
            'scope2_co2': 0.0,
            'scope2_ch4': 0.0,
            'scope2_n2o': 0.0,
            'scope1_by_category': defaultdict(float),
            'scope1_by_fuel': defaultdict(float),
            'scope3_by_category': defaultdict(float),
        }
        
        for em in facility_emissions:
            scope = (em.get('scope') or '').lower()
            tco2e = float(em.get('total_emissions', 0) or 0)
            category = self._get_category_from_emission(em)
            fuel = self._get_fuel_from_emission(em)
            period = em.get('reporting_period') or ''
            
            # Track by_category and by_fuel for ALL scopes (used for charts)
            totals['by_category'][category] += tco2e
            totals['by_fuel'][fuel] += tco2e
            totals['by_category_fuel'][category][fuel] += tco2e
            
            # Determine scope label for by_scope_category_fuel
            if 'scope 1' in scope or 'scope1' in scope or scope == '1':
                totals['by_scope_category_fuel']['scope1'][category][fuel] += tco2e
                totals['scope1'] += tco2e
                # Track scope1-specific category and fuel breakdowns
                totals['scope1_by_category'][category] += tco2e
                totals['scope1_by_fuel'][fuel] += tco2e
                # Individual gas components from actual data
                totals['scope1_co2'] += float(em.get('co2_emissions', 0) or 0)
                totals['scope1_ch4'] += float(em.get('ch4_emissions', 0) or 0)
                totals['scope1_n2o'] += float(em.get('n2o_emissions', 0) or 0)
            elif 'scope 2' in scope or 'scope2' in scope or scope == '2':
                totals['scope2'] += tco2e
                totals['by_scope_category_fuel']['scope2'][category][fuel] += tco2e
                # Individual gas components from actual data (not hardcoded to CO2)
                totals['scope2_co2'] += float(em.get('co2_emissions', 0) or 0)
                totals['scope2_ch4'] += float(em.get('ch4_emissions', 0) or 0)
                totals['scope2_n2o'] += float(em.get('n2o_emissions', 0) or 0)
            elif 'biogenic' in scope:
                totals['biogenic'] += tco2e
                totals['by_scope_category_fuel']['biogenic'][category][fuel] += tco2e
            elif 'scope 3' in scope or 'scope3' in scope or scope == '3':
                totals['scope3'] += tco2e
                totals['by_scope_category_fuel']['scope3'][category][fuel] += tco2e
                totals['scope3_by_category'][category] += tco2e
            
            # Track by month (skip yearly records as they don't have monthly breakdown)
            frequency_type = em.get('frequency_type', 'monthly')
            if period and frequency_type != 'yearly':
                month_key = self._format_month(period.split(' to ')[0] if ' to ' in period else period)
                if 'scope 1' in scope or 'scope1' in scope or scope == '1':
                    totals['by_month'][month_key]['scope1'] += tco2e
                elif 'scope 2' in scope or 'scope2' in scope or scope == '2':
                    totals['by_month'][month_key]['scope2'] += tco2e
                elif 'scope 3' in scope or 'scope3' in scope or scope == '3':
                    totals['by_month'][month_key]['scope3'] += tco2e
        
        # Calculate sinks/removals for this facility
        if hasattr(self, 'sinks_data') and self.sinks_data and facility_id:
            facility_sinks = [s for s in self.sinks_data if s.get('facility_id') == facility_id]
            totals['removals'] = sum(s.get('total_emissions_reduced', 0) for s in facility_sinks)
        elif hasattr(self, 'sinks_total'):
            # If no facility_id provided, use total sinks distributed
            totals['removals'] = getattr(self, 'sinks_total', 0)
        
        totals['total'] = totals['scope1'] + totals['scope2']
        totals['total_ghg'] = totals['total'] - totals['removals']
        
        return totals
    
    def _get_emission_processes(self, facility_emissions: List[Dict]) -> Tuple[List[str], List[str]]:
        """Get unique emission processes for Scope 1 and Scope 2"""
        scope1_processes = []
        scope2_processes = []
        
        for em in facility_emissions:
            scope = (em.get('scope') or '').lower()
            process_names = self._get_process_names_from_emission(em)
            fuel = self._get_fuel_from_emission(em)
            
            for process in process_names:
                if process and fuel:
                    process_fuel = f"{process} - {fuel}"
                    if 'scope1' in scope or 'scope 1' in scope or scope == '1':
                        scope1_processes.append(process_fuel)
                    elif 'scope2' in scope or 'scope 2' in scope or scope == '2':
                        scope2_processes.append(process_fuel)
            
            # If no process names but has fuel, use category
            if not process_names and fuel:
                category = self._get_category_from_emission(em)
                process_fuel = f"{category} - {fuel}"
                if 'scope1' in scope or 'scope 1' in scope or scope == '1':
                    scope1_processes.append(process_fuel)
                elif 'scope2' in scope or 'scope 2' in scope or scope == '2':
                    scope2_processes.append(process_fuel)
        
        # Deduplicate (case insensitive)
        scope1_processes = self._deduplicate_list(scope1_processes, case_insensitive=True)
        scope2_processes = self._deduplicate_list(scope2_processes, case_insensitive=True)
        
        # Show "NA" if no processes found for a scope
        if not scope2_processes:
            scope2_processes = ["NA"]
        
        return scope1_processes, scope2_processes
    
    def _get_emission_processes_by_category(self, facility_emissions: List[Dict]) -> Dict[str, List[str]]:
        """Get emission processes for Scope 1 segregated by category (Stationary Combustion, Mobile Combustion, Fugitive emissions)"""
        categories = {
            'stationary_combustion': [],
            'mobile_combustion': [],
            'fugitive_emissions': [],
            'other': []
        }
        
        # Category mapping for normalization
        category_mapping = {
            'stationary combustion': 'stationary_combustion',
            'stationary_combustion': 'stationary_combustion',
            'stationarycombustion': 'stationary_combustion',
            'mobile combustion': 'mobile_combustion',
            'mobile_combustion': 'mobile_combustion',
            'mobilecombustion': 'mobile_combustion',
            'fugitive emissions': 'fugitive_emissions',
            'fugitive_emissions': 'fugitive_emissions',
            'fugitiveemissions': 'fugitive_emissions',
            'fugitive': 'fugitive_emissions',
        }
        
        for em in facility_emissions:
            scope = (em.get('scope') or '').lower()
            if not ('scope1' in scope or 'scope 1' in scope or scope == '1'):
                continue
                
            process_names = self._get_process_names_from_emission(em)
            fuel = self._get_fuel_from_emission(em)
            category = self._get_category_from_emission(em)
            
            # Normalize category
            cat_lower = (category or '').lower().strip()
            cat_key = category_mapping.get(cat_lower, 'other')
            
            for process in process_names:
                if process and fuel:
                    process_fuel = f"{process} - {fuel}"
                    categories[cat_key].append(process_fuel)
                elif fuel:
                    # No process name - just use fuel
                    categories[cat_key].append(fuel)
                elif process:
                    # Has process but no fuel
                    categories[cat_key].append(process)
            
            # If no process names but has fuel
            if not process_names and fuel:
                categories[cat_key].append(fuel)
            
            # If no process names and no fuel, use "NA"
            if not process_names and not fuel:
                categories[cat_key].append("NA")
        
        # Deduplicate each category
        for key in categories:
            categories[key] = self._deduplicate_list(categories[key], case_insensitive=True)
        
        return categories

    def _get_emission_processes_by_category_extended(self, facility_emissions: List[Dict]) -> Dict[str, List[Dict]]:
        """Get emission processes for Scope 1 with Person Responsible and Source of Information.
        Returns Dict[str, List[Dict]] where each dict has keys: process_fuel, responsible_person, record_source
        """
        # Track unique process_fuel -> {responsible_persons: set, record_sources: set}
        categories_data = {
            'stationary_combustion': {},
            'mobile_combustion': {},
            'fugitive_emissions': {},
            'other': {}
        }
        
        category_mapping = {
            'stationary combustion': 'stationary_combustion',
            'stationary_combustion': 'stationary_combustion',
            'stationarycombustion': 'stationary_combustion',
            'mobile combustion': 'mobile_combustion',
            'mobile_combustion': 'mobile_combustion',
            'mobilecombustion': 'mobile_combustion',
            'fugitive emissions': 'fugitive_emissions',
            'fugitive_emissions': 'fugitive_emissions',
            'fugitiveemissions': 'fugitive_emissions',
            'fugitive': 'fugitive_emissions',
        }
        
        for em in facility_emissions:
            scope = (em.get('scope') or '').lower()
            if not ('scope1' in scope or 'scope 1' in scope or scope == '1'):
                continue
                
            process_names = self._get_process_names_from_emission(em)
            fuel = self._get_fuel_from_emission(em)
            category = self._get_category_from_emission(em)
            responsible = em.get('responsible_person') or ''
            record_source = em.get('record_source') or ''
            
            cat_lower = (category or '').lower().strip()
            cat_key = category_mapping.get(cat_lower, 'other')
            
            process_fuels = []
            for process in process_names:
                if process and fuel:
                    process_fuels.append(f"{process} - {fuel}")
                elif fuel:
                    process_fuels.append(fuel)
                elif process:
                    process_fuels.append(process)
            
            if not process_names and fuel:
                process_fuels.append(fuel)
            if not process_names and not fuel:
                process_fuels.append("NA")
            
            for pf in process_fuels:
                pf_lower = pf.lower()
                if pf_lower not in categories_data[cat_key]:
                    categories_data[cat_key][pf_lower] = {
                        'process_fuel': pf,
                        'responsible_persons': set(),
                        'record_sources': set()
                    }
                if responsible:
                    categories_data[cat_key][pf_lower]['responsible_persons'].add(responsible)
                if record_source:
                    categories_data[cat_key][pf_lower]['record_sources'].add(record_source)
        
        # Convert to final format
        result = {}
        for cat_key, data in categories_data.items():
            result[cat_key] = []
            for pf_lower, info in data.items():
                result[cat_key].append({
                    'process_fuel': info['process_fuel'],
                    'responsible_person': ', '.join(sorted(info['responsible_persons'])) or 'NA',
                    'record_source': ', '.join(sorted(info['record_sources'])) or 'NA'
                })
        
        return result

    def _get_scope3_processes_by_category(self, facility_emissions: List[Dict]) -> Dict[str, List[str]]:
        """Get emission processes for Scope 3 segregated by GHG Protocol category (C1-C15)"""
        # Initialize all 15 Scope 3 categories
        categories = {f'c{i}': [] for i in range(1, 16)}
        categories['other'] = []
        
        # Category display names
        self.scope3_category_display = {
            'c1': 'C1 - Purchased Goods and Services',
            'c2': 'C2 - Capital Goods',
            'c3': 'C3 - Fuel and Energy Related Activities Not Included in Scope 1 or Scope 2',
            'c4': 'C4 - Upstream Transportation and Distribution',
            'c5': 'C5 - Waste Generated in Operations',
            'c6': 'C6 - Business Travel',
            'c7': 'C7 - Employee Commuting',
            'c8': 'C8 - Upstream Leased Assets',
            'c9': 'C9 - Downstream Transportation and Distribution',
            'c10': 'C10 - Processing of Sold Products',
            'c11': 'C11 - Use of Sold Products',
            'c12': 'C12 - End-of-Life Treatment of Sold Products',
            'c13': 'C13 - Downstream Leased Assets',
            'c14': 'C14 - Franchises',
            'c15': 'C15 - Investments',
            'other': 'Other Scope 3'
        }
        
        def get_category_key(category_str: str) -> str:
            """Extract category key (c1-c15) from category string"""
            if not category_str:
                return 'other'
            
            cat_lower = category_str.lower().strip()
            
            # Try exact prefix match first (e.g., "c15 - investments" starts with "c15")
            # Check longer prefixes first to avoid c1 matching c10, c11, etc.
            for i in range(15, 0, -1):  # Check c15 down to c1
                prefix = f'c{i}'
                # Check if starts with "c15 " or "c15-" or equals "c15"
                if cat_lower.startswith(f'{prefix} ') or cat_lower.startswith(f'{prefix}-') or cat_lower == prefix:
                    return prefix
            
            # Try matching by full category name
            category_name_mapping = {
                'purchased goods and services': 'c1',
                'capital goods': 'c2',
                'fuel and energy related activities': 'c3',
                'upstream transportation and distribution': 'c4',
                'upstream transportation': 'c4',
                'waste generated in operations': 'c5',
                'business travel': 'c6',
                'employee commuting': 'c7',
                'upstream leased assets': 'c8',
                'downstream transportation and distribution': 'c9',
                'downstream transportation': 'c9',
                'processing of sold products': 'c10',
                'use of sold products': 'c11',
                'end-of-life treatment of sold products': 'c12',
                'end-of-life treatment': 'c12',
                'downstream leased assets': 'c13',
                'franchises': 'c14',
                'investments': 'c15',
            }
            
            for name, key in category_name_mapping.items():
                if name in cat_lower:
                    return key
            
            return 'other'
        
        for em in facility_emissions:
            scope = (em.get('scope') or '').lower()
            if not ('scope3' in scope or 'scope 3' in scope or scope == '3'):
                continue
            
            category = self._get_category_from_emission(em)
            fuel = self._get_fuel_from_emission(em)
            process_names = self._get_process_names_from_emission(em)
            
            # Get category key using improved matching
            cat_key = get_category_key(category)
            
            # Build process name in format: {Process Name} - {Activity Name} or just {Activity Name} if no process
            for process in process_names:
                if process and fuel and fuel != 'Unknown':
                    process_info = f"{process} - {fuel}"
                elif fuel and fuel != 'Unknown':
                    # No process name - just use fuel/activity
                    process_info = fuel
                elif process:
                    process_info = process
                else:
                    process_info = "NA"
                
                categories[cat_key].append(process_info)
            
            # If no process names but has fuel, just use fuel
            if not process_names:
                if fuel and fuel != 'Unknown':
                    process_info = fuel
                else:
                    process_info = "NA"
                categories[cat_key].append(process_info)
        
        # Deduplicate each category and remove empty ones
        result = {}
        for key in categories:
            deduped = self._deduplicate_list(categories[key], case_insensitive=True)
            if deduped:
                result[key] = deduped
        
        return result

    def _get_scope3_processes_by_category_extended(self, facility_emissions: List[Dict]) -> Dict[str, List[Dict]]:
        """Get Scope 3 processes with Person Responsible and Source of Information.
        Returns Dict[str, List[Dict]] where each dict has keys: process_fuel, responsible_person, record_source
        """
        # Initialize all 15 Scope 3 categories
        categories_data = {f'c{i}': {} for i in range(1, 16)}
        categories_data['other'] = {}
        
        def get_category_key(category_str: str) -> str:
            if not category_str:
                return 'other'
            cat_lower = category_str.lower().strip()
            for i in range(15, 0, -1):
                prefix = f'c{i}'
                if cat_lower.startswith(f'{prefix} ') or cat_lower.startswith(f'{prefix}-') or cat_lower == prefix:
                    return prefix
            category_name_mapping = {
                'purchased goods and services': 'c1', 'capital goods': 'c2',
                'fuel and energy related activities': 'c3', 'upstream transportation and distribution': 'c4',
                'upstream transportation': 'c4', 'waste generated in operations': 'c5',
                'business travel': 'c6', 'employee commuting': 'c7', 'upstream leased assets': 'c8',
                'downstream transportation and distribution': 'c9', 'downstream transportation': 'c9',
                'processing of sold products': 'c10', 'use of sold products': 'c11',
                'end-of-life treatment of sold products': 'c12', 'end-of-life treatment': 'c12',
                'downstream leased assets': 'c13', 'franchises': 'c14', 'investments': 'c15',
            }
            for name, key in category_name_mapping.items():
                if name in cat_lower:
                    return key
            return 'other'
        
        for em in facility_emissions:
            scope = (em.get('scope') or '').lower()
            if not ('scope3' in scope or 'scope 3' in scope or scope == '3'):
                continue
            
            category = self._get_category_from_emission(em)
            fuel = self._get_fuel_from_emission(em)
            process_names = self._get_process_names_from_emission(em)
            responsible = em.get('responsible_person') or ''
            record_source = em.get('record_source') or ''
            
            cat_key = get_category_key(category)
            
            process_infos = []
            for process in process_names:
                if process and fuel and fuel != 'Unknown':
                    process_infos.append(f"{process} - {fuel}")
                elif fuel and fuel != 'Unknown':
                    process_infos.append(fuel)
                elif process:
                    process_infos.append(process)
                else:
                    process_infos.append("NA")
            
            if not process_names:
                process_infos.append(fuel if fuel and fuel != 'Unknown' else "NA")
            
            for pf in process_infos:
                pf_lower = pf.lower()
                if pf_lower not in categories_data[cat_key]:
                    categories_data[cat_key][pf_lower] = {
                        'process_fuel': pf,
                        'responsible_persons': set(),
                        'record_sources': set()
                    }
                if responsible:
                    categories_data[cat_key][pf_lower]['responsible_persons'].add(responsible)
                if record_source:
                    categories_data[cat_key][pf_lower]['record_sources'].add(record_source)
        
        # Convert to final format, removing empty categories
        result = {}
        for cat_key, data in categories_data.items():
            if data:
                result[cat_key] = []
                for pf_lower, info in data.items():
                    result[cat_key].append({
                        'process_fuel': info['process_fuel'],
                        'responsible_person': ', '.join(sorted(info['responsible_persons'])) or 'NA',
                        'record_source': ', '.join(sorted(info['record_sources'])) or 'NA'
                    })
        
        return result

    def _get_scope2_processes_extended(self, facility_emissions: List[Dict]) -> List[Dict]:
        """Get Scope 2 processes with Person Responsible and Source of Information."""
        processes_data = {}
        
        for em in facility_emissions:
            scope = (em.get('scope') or '').lower()
            if not ('scope2' in scope or 'scope 2' in scope or scope == '2'):
                continue
            
            process_names = self._get_process_names_from_emission(em)
            fuel = self._get_fuel_from_emission(em)
            responsible = em.get('responsible_person') or ''
            record_source = em.get('record_source') or ''
            
            process_fuels = []
            for process in process_names:
                if process and fuel:
                    process_fuels.append(f"{process} - {fuel}")
                elif fuel:
                    process_fuels.append(fuel)
                elif process:
                    process_fuels.append(process)
            
            if not process_names and fuel:
                process_fuels.append(fuel)
            if not process_names and not fuel:
                process_fuels.append("NA")
            
            for pf in process_fuels:
                pf_lower = pf.lower()
                if pf_lower not in processes_data:
                    processes_data[pf_lower] = {
                        'process_fuel': pf,
                        'responsible_persons': set(),
                        'record_sources': set()
                    }
                if responsible:
                    processes_data[pf_lower]['responsible_persons'].add(responsible)
                if record_source:
                    processes_data[pf_lower]['record_sources'].add(record_source)
        
        result = []
        for pf_lower, info in processes_data.items():
            result.append({
                'process_fuel': info['process_fuel'],
                'responsible_person': ', '.join(sorted(info['responsible_persons'])) or 'NA',
                'record_source': ', '.join(sorted(info['record_sources'])) or 'NA'
            })
        
        return result if result else [{'process_fuel': 'NA', 'responsible_person': 'NA', 'record_source': 'NA'}]

    def _get_unique_fuels(self, facility_emissions: List[Dict]) -> Tuple[List[str], List[str]]:
        """Get unique fuel names for Scope 1 and Scope 2"""
        scope1_fuels = []
        scope2_fuels = []
        
        for em in facility_emissions:
            scope = (em.get('scope') or '').lower()
            fuel = self._get_fuel_from_emission(em)
            
            if fuel and fuel != 'Unknown':
                if 'scope1' in scope or 'scope 1' in scope or scope == '1':
                    scope1_fuels.append(fuel)
                elif 'scope2' in scope or 'scope 2' in scope or scope == '2':
                    scope2_fuels.append(fuel)
        
        scope1_fuels = self._deduplicate_list(scope1_fuels, case_insensitive=True)
        scope2_fuels = self._deduplicate_list(scope2_fuels, case_insensitive=True)
        
        if not scope2_fuels:
            scope2_fuels = ["NA"]
        
        return scope1_fuels, scope2_fuels
    
    def _get_previous_year_data(self, emissions: List[Dict], current_start: str, current_end: str = None) -> List[Dict]:
        """
        Get previous period emissions data (emissions before the current reporting period).
        
        For FY 2024-2025 (April 2024 - March 2025):
        - CY 2024 (Jan-Dec 2024): 3 months (Jan-Mar 2024) fall before → prorate as 3/12
        - CY 2023: Fully before → include fully
        - FY 2023-2024: Fully before → include fully
        - Monthly records before 2024-04: Include fully
        
        Returns a list of dicts with: scope, category, fuel, reporting_period, tco2e, is_prorated
        """
        try:
            prev_year_records = []
            
            # Parse current reporting period start
            try:
                current_start_year = int(current_start[:4])
                current_start_month = int(current_start[5:7]) if len(current_start) >= 7 else 1
            except (ValueError, IndexError):
                return []
            
            # Build set of (year, month) tuples for current reporting period
            current_period_months = set()
            if current_end:
                try:
                    current_end_year = int(current_end[:4])
                    current_end_month = int(current_end[5:7]) if len(current_end) >= 7 else 12
                    
                    cy, cm = current_start_year, current_start_month
                    while (cy, cm) <= (current_end_year, current_end_month):
                        current_period_months.add((cy, cm))
                        cm += 1
                        if cm > 12:
                            cm = 1
                            cy += 1
                except (ValueError, IndexError):
                    pass
            
            for em in emissions:
                period = em.get('reporting_period') or ''
                if not period:
                    continue
                
                period = period.strip()
                frequency_type = em.get('frequency_type', 'monthly')
                
                # Get emission months based on period format
                emission_months = set()
                
                if frequency_type == 'yearly' or period.startswith('CY') or period.startswith('FY'):
                    # Yearly record - calculate months
                    if period.startswith('CY') or period.startswith('CY '):
                        # Calendar Year
                        try:
                            cy_year = int(period.replace('CY', '').strip()[:4])
                            for month in range(1, 13):
                                emission_months.add((cy_year, month))
                        except (ValueError, IndexError):
                            continue
                    elif period.startswith('FY') or period.startswith('FY '):
                        # Financial Year
                        try:
                            fy_part = period.replace('FY', '').strip()
                            if '-' in fy_part:
                                years = fy_part.split('-')
                                fy_start_year = int(years[0].strip())
                                end_year_str = years[1].strip()
                                if len(end_year_str) == 2:
                                    fy_end_year = int(f"{str(fy_start_year)[:2]}{end_year_str}")
                                else:
                                    fy_end_year = int(end_year_str)
                            else:
                                fy_start_year = int(fy_part[:4])
                                fy_end_year = fy_start_year + 1
                            
                            # FY runs April to March
                            for month in range(4, 13):
                                emission_months.add((fy_start_year, month))
                            for month in range(1, 4):
                                emission_months.add((fy_end_year, month))
                        except (ValueError, IndexError):
                            continue
                    else:
                        continue
                else:
                    # Monthly record
                    em_period_start = period.split(' to ')[0].strip() if ' to ' in period else period.strip()
                    try:
                        em_year = int(em_period_start[:4])
                        em_month = int(em_period_start[5:7]) if len(em_period_start) >= 7 else 1
                        emission_months.add((em_year, em_month))
                    except (ValueError, IndexError):
                        continue
                
                if not emission_months:
                    continue
                
                # Calculate months that fall BEFORE the current reporting period
                months_before_current = set()
                for (y, m) in emission_months:
                    if (y, m) < (current_start_year, current_start_month):
                        months_before_current.add((y, m))
                
                if not months_before_current:
                    # No months fall before current period - skip
                    continue
                
                # Calculate proration
                total_emission_months = len(emission_months)
                months_before_count = len(months_before_current)
                
                proration_factor = months_before_count / total_emission_months
                is_prorated = proration_factor < 1.0
                
                # Get emission value and apply proration
                tco2e = float(em.get('total_emissions', 0) or em.get('co2e_emissions', 0) or 0)
                prorated_tco2e = tco2e * proration_factor
                
                # Build record
                scope = em.get('scope', '')
                scope_lower = scope.lower() if scope else ''
                if 'scope1' in scope_lower or 'scope 1' in scope_lower or scope == '1':
                    scope_display = 'Scope 1'
                elif 'scope2' in scope_lower or 'scope 2' in scope_lower or scope == '2':
                    scope_display = 'Scope 2'
                elif 'scope3' in scope_lower or 'scope 3' in scope_lower or scope == '3':
                    scope_display = 'Scope 3'
                elif 'biogenic' in scope_lower:
                    scope_display = 'Biogenic'
                else:
                    scope_display = scope or 'Unknown'
                
                category = self._get_category_from_emission(em)
                fuel = self._get_fuel_from_emission(em)
                
                prev_year_records.append({
                    'scope': scope_display,
                    'category': category,
                    'fuel': fuel,
                    'reporting_period': period,
                    'tco2e': prorated_tco2e,
                    'is_prorated': is_prorated,
                    'original_tco2e': tco2e,
                    'proration_factor': proration_factor
                })
            
            # Sort by scope, then category, then fuel
            def sort_key(r):
                scope_order = {'Scope 1': 1, 'Scope 2': 2, 'Scope 3': 3, 'Biogenic': 4}.get(r['scope'], 9)
                return (scope_order, r['category'], r['fuel'], r['reporting_period'])
            
            prev_year_records.sort(key=sort_key)
            
            return prev_year_records
        except Exception as e:
            print(f"Error getting previous period data: {e}")
            import traceback
            traceback.print_exc()
            return []
    
    def _add_scope3_methodology_section(self, doc: Document, emissions: List[Dict]):
        """Add Scope 3 methodology analysis section with donut chart similar to dashboard"""
        
        # Filter scope 3 emissions
        scope3_emissions = [e for e in emissions if (e.get('scope') or '').lower() == 'scope3']
        
        if not scope3_emissions:
            return  # No Scope 3 emissions to analyze
        
        # Analyze emissions by methodology
        methodology_totals = {
            'activity_basis': {'count': 0, 'total': 0.0},
            'spend_basis': {'count': 0, 'total': 0.0},
            'supplier_basis': {'count': 0, 'total': 0.0},
            'other': {'count': 0, 'total': 0.0}
        }
        
        for em in scope3_emissions:
            method = (em.get('calculation_method_scope3') or 'other').lower()
            co2e = float(em.get('total_emissions') or em.get('co2e_emissions') or 0)
            
            # Normalize method name
            if 'activity' in method:
                method_key = 'activity_basis'
            elif 'spend' in method:
                method_key = 'spend_basis'
            elif 'supplier' in method:
                method_key = 'supplier_basis'
            else:
                method_key = 'other'
            
            methodology_totals[method_key]['count'] += 1
            methodology_totals[method_key]['total'] += co2e
        
        # Calculate grand total
        grand_total = sum(m['total'] for m in methodology_totals.values())
        
        if grand_total > 0:
            p = doc.add_paragraph()
            run = p.add_run("Scope 3 Methodology Analysis")
            run.bold = True
            
            doc.add_paragraph()
            
            # Overall methodology breakdown table
            headers = ['Methodology', 'Records', 'Emissions (tCO₂e)', 'Percentage']
            data = []
            
            method_labels = {
                'activity_basis': 'Average Data Based',
                'spend_basis': 'Spend Based',
                'supplier_basis': 'Supplier Based',
                'other': 'Other/Unspecified'
            }
            
            method_colors = {
                'activity_basis': '#10B981',  # Emerald
                'spend_basis': '#F59E0B',     # Amber
                'supplier_basis': '#3B82F6',  # Blue
                'other': '#6B7280'            # Gray
            }
            
            chart_data = []
            for method_key, label in method_labels.items():
                method_data = methodology_totals[method_key]
                if method_data['total'] > 0:
                    pct = (method_data['total'] / grand_total) * 100
                    data.append([
                        label,
                        str(method_data['count']),
                        self._format_number(method_data['total']),
                        f"{pct:.1f}%"
                    ])
                    chart_data.append({
                        'label': label,
                        'value': method_data['total'],
                        'percentage': pct,
                        'color': method_colors[method_key]
                    })
            
            # Add total row
            total_count = sum(m['count'] for m in methodology_totals.values())
            data.append(['Total', str(total_count), self._format_number(grand_total), '100.0%'])
            
            self._create_styled_table(doc, headers, data, bold_rows=[len(data)-1])
            
            # Generate donut chart for methodology analysis
            if len(chart_data) > 0:
                self._add_methodology_donut_chart(doc, chart_data, grand_total)
        
        doc.add_paragraph()
    
    def _add_methodology_donut_chart(self, doc: Document, chart_data: List[Dict], total: float):
        """Add a donut chart for Scope 3 methodology analysis similar to dashboard style"""
        try:
            fig, ax = plt.subplots(figsize=(8, 5))
            
            labels = [d['label'] for d in chart_data]
            sizes = [d['value'] for d in chart_data]
            colors = [d['color'] for d in chart_data]
            
            # Create donut chart
            wedges, texts, autotexts = ax.pie(
                sizes, 
                labels=None,  # We'll add labels manually
                colors=colors,
                autopct=lambda pct: f'{pct:.1f}%' if pct > 5 else '',
                startangle=90,
                pctdistance=0.75,
                wedgeprops=dict(width=0.5, edgecolor='white', linewidth=2)
            )
            
            # Style the percentage labels
            for autotext in autotexts:
                autotext.set_fontsize(9)
                autotext.set_fontweight('bold')
                autotext.set_color('white')
            
            # Add center text
            ax.text(0, 0, f'{self._format_number(total)}\ntCO₂e', 
                    ha='center', va='center', fontsize=12, fontweight='bold')
            
            # Add legend on the right side
            legend_labels = [f'{d["label"]}: {d["percentage"]:.1f}%' for d in chart_data]
            ax.legend(wedges, legend_labels, loc='center left', bbox_to_anchor=(1.1, 0.5), fontsize=9)
            
            ax.set_title('Scope 3 Emissions by Methodology', fontsize=11, fontweight='bold', pad=10)
            
            plt.tight_layout()
            
            # Save to buffer
            buf = io.BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
            buf.seek(0)
            plt.close(fig)
            
            # Add to document
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(buf, width=Inches(5.5))
            
            self._add_figure_caption(doc, "Figure: Scope 3 Emissions by Methodology")
            
        except Exception as e:
            print(f"Error creating methodology donut chart: {e}")
            import traceback
            traceback.print_exc()
    
    def _get_base_year_emissions_for_entity(self, entity_type: str, entity_id: str) -> Optional[Dict]:
        """Get base year emissions data for a facility or organization from the database.
        For Scope 1,2,3 reports, fetches both scope12 and scope3 base year records separately."""
        try:
            from pymongo import MongoClient
            mongo_url = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
            db_name = os.environ.get('DB_NAME', 'ghg_platform')
            
            client = MongoClient(mongo_url)
            db = client[db_name]
            
            query = {}
            if entity_type == 'facility':
                query['facility_id'] = entity_id
            else:
                query['organization_id'] = entity_id
                query['facility_id'] = None  # Organization-level record
            
            report_type = getattr(self, 'report_type', 'scope_1_2')
            
            if report_type == 'scope_1_2_3':
                # For Scope 1,2,3 reports, fetch both scope12 and scope3 records separately
                # Note: Old records may have scope_group=None instead of 'scope12', so we check for both
                scope12_query = {**query, '$or': [{'scope_group': 'scope12'}, {'scope_group': None}, {'scope_group': {'$exists': False}}]}
                scope3_query = {**query, 'scope_group': 'scope3'}
                
                scope12_record = db.base_year_emissions.find_one(scope12_query, {"_id": 0})
                scope3_record = db.base_year_emissions.find_one(scope3_query, {"_id": 0})
                
                client.close()
                
                if not scope12_record and not scope3_record:
                    return None
                
                # Build a structured result with clear separation
                result = {
                    'has_scope12_base_year': scope12_record is not None,
                    'has_scope3_base_year': scope3_record is not None,
                    'scope12_base_year': scope12_record.get('base_year') if scope12_record else None,
                    'scope3_base_year': scope3_record.get('base_year') if scope3_record else None,
                    'scope12_emissions_data': scope12_record.get('emissions_data', []) if scope12_record else [],
                    'scope3_emissions_data': scope3_record.get('emissions_data', []) if scope3_record else [],
                    # Combined emissions_data for the table display
                    'emissions_data': (scope12_record.get('emissions_data', []) if scope12_record else []) + 
                                      (scope3_record.get('emissions_data', []) if scope3_record else []),
                    # Use scope12 base year as primary, fallback to scope3 if scope12 doesn't exist
                    'base_year': scope12_record.get('base_year') if scope12_record else (scope3_record.get('base_year') if scope3_record else 'N/A'),
                }
                
                return result
            else:
                # For Scope 1,2 reports, just fetch scope12 record (or legacy None scope_group)
                query['$or'] = [{'scope_group': 'scope12'}, {'scope_group': None}, {'scope_group': {'$exists': False}}]
                base_year_record = db.base_year_emissions.find_one(query, {"_id": 0})
                client.close()
                if base_year_record:
                    base_year_record['has_scope12_base_year'] = True
                    base_year_record['scope12_base_year'] = base_year_record.get('base_year')
                    # Populate scope12_emissions_data for consistency with Scope 1,2,3 reports
                    base_year_record['scope12_emissions_data'] = base_year_record.get('emissions_data', [])
                return base_year_record
        except Exception as e:
            print(f"Error getting base year emissions: {e}")
            return None
    
    def _add_base_year_emissions_section(self, doc: Document, base_year_data: Dict, current_totals: Dict, 
                                         entity_name: str, equity_factor: float, use_equity_share: bool,
                                         reporting_period_start: str = None, reporting_period_end: str = None):
        """Add base year emissions table and comparison analysis to the document.
        
        Handles Scope 1&2 and Scope 3 comparisons independently with their own base years.
        Shows separate tables for Scope 1,2 and Scope 3 base year data.
        """
        report_type = getattr(self, 'report_type', 'scope_1_2')
        
        # Get flags and base years
        has_scope12_base_year = base_year_data.get('has_scope12_base_year', False)
        has_scope3_base_year = base_year_data.get('has_scope3_base_year', False)
        scope12_base_year = base_year_data.get('scope12_base_year')
        scope3_base_year_str = base_year_data.get('scope3_base_year')
        base_year = base_year_data.get('base_year', 'N/A')
        
        # Get emissions data
        scope12_emissions_data = base_year_data.get('scope12_emissions_data', [])
        scope3_emissions_data = base_year_data.get('scope3_emissions_data', [])
        
        # If no base year data at all, show message and return
        if not has_scope12_base_year and not has_scope3_base_year:
            p = doc.add_paragraph()
            run = p.add_run("Base year not defined.")
            run.italic = True
            return
        
        # Helper to check if reporting period matches base year
        def is_reporting_period_base_year(by_str):
            """Check if the reporting period matches the base year."""
            if not by_str or not reporting_period_start or not reporting_period_end:
                return False
            # Handle FY format like "FY 2021-2022"
            if by_str.startswith('FY ') or by_str.startswith('CY '):
                try:
                    years_part = by_str.split(' ')[1]  # "2021-2022"
                    start_year = int(years_part.split('-')[0])
                    end_year_short = years_part.split('-')[1]
                    end_year = int(f"{str(start_year)[:2]}{end_year_short}") if len(end_year_short) == 2 else int(end_year_short)
                    
                    rp_start_year = int(reporting_period_start.split('-')[0])
                    rp_end_year = int(reporting_period_end.split('-')[0])
                    
                    if by_str.startswith('FY '):
                        return rp_start_year >= start_year and rp_end_year <= end_year + 1
                    else:  # CY
                        return rp_start_year == start_year and rp_end_year == start_year
                except:
                    pass
            return False
        
        # Calculate base year emissions totals from the actual data
        scope1_2_base_year_total = 0.0
        scope3_base_year_total = 0.0
        biogenic_base_year = 0.0
        
        # Process scope12 emissions
        for em in scope12_emissions_data:
            scope = em.get('scope', '').lower()
            tco2e = float(em.get('tco2e', 0) or 0)
            if 'biogenic' in scope:
                biogenic_base_year += tco2e
            else:
                scope1_2_base_year_total += tco2e
        
        # Process scope3 emissions
        for em in scope3_emissions_data:
            tco2e = float(em.get('tco2e', 0) or 0)
            scope3_base_year_total += tco2e
        
        # Get current period totals
        current_scope1 = current_totals.get('scope1', 0)
        current_scope2 = current_totals.get('scope2', 0)
        current_scope3 = current_totals.get('scope3', 0)
        current_removals = current_totals.get('removals', 0)
        current_scope1_2 = current_scope1 + current_scope2
        
        # Apply equity factor to base year totals
        scope1_2_base_year_display = scope1_2_base_year_total * equity_factor if use_equity_share else scope1_2_base_year_total
        scope3_base_year_display = scope3_base_year_total * equity_factor if use_equity_share else scope3_base_year_total
        
        if report_type == 'scope_1_2_3':
            # ==================== SCOPE 1 & 2 BASE YEAR SECTION ====================
            p = doc.add_paragraph()
            run = p.add_run("Scope 1 & 2 Base Year Emissions")
            run.bold = True
            run.font.size = Pt(12)
            doc.add_paragraph()
            
            if has_scope12_base_year and scope12_base_year:
                p = doc.add_paragraph()
                run = p.add_run(f"Base Year: ")
                run.bold = True
                p.add_run(str(scope12_base_year))
                doc.add_paragraph()
                
                # Create Scope 1&2 base year table
                if scope12_emissions_data:
                    headers = ['Scope', 'Category', 'Subcategory', 'Emissions (tCO₂e)']
                    data = []
                    for em in scope12_emissions_data:
                        scope = em.get('scope', '')
                        category = em.get('category', '')
                        subcategory = em.get('subcategory', '')
                        tco2e = float(em.get('tco2e', 0) or 0)
                        display_tco2e = tco2e * equity_factor if use_equity_share else tco2e
                        data.append([scope, category, subcategory, self._format_number(display_tco2e)])
                    
                    if data:
                        data.append(['', '', 'Total Scope 1 & 2 Base Year', self._format_number(scope1_2_base_year_display)])
                        self._create_styled_table(doc, headers, data, bold_rows=[len(data)-1])
                        doc.add_paragraph()
                
                # Scope 1&2 Comparison
                if not is_reporting_period_base_year(scope12_base_year):
                    change_1_2 = current_scope1_2 - scope1_2_base_year_display
                    # Fix: Calculate percentage correctly, handle 0 base year case
                    if scope1_2_base_year_display > 0:
                        change_pct_1_2 = (change_1_2 / scope1_2_base_year_display) * 100
                    elif current_scope1_2 > 0:
                        change_pct_1_2 = 100.0  # 100% increase from 0
                    else:
                        change_pct_1_2 = 0.0
                    
                    comparison_headers = ['Period', 'Scope 1 & 2 Emissions (tCO₂e)']
                    comparison_data = [
                        [f"Base Year ({scope12_base_year})", self._format_number(scope1_2_base_year_display)],
                        ["Current Reporting Period", self._format_number(current_scope1_2)],
                        ["Change", f"{self._format_number(change_1_2)} ({'+' if change_1_2 >= 0 else ''}{change_pct_1_2:.1f}%)"]
                    ]
                    self._create_styled_table(doc, comparison_headers, comparison_data, bold_rows=[2])
                    
                    # Add comparison chart for Scope 1&2
                    if scope1_2_base_year_display > 0 or current_scope1_2 > 0:
                        try:
                            chart_buffer = self._create_base_year_comparison_chart(
                                scope1_2_base_year_display, current_scope1_2, 
                                f"Base Year ({scope12_base_year})", "Current Period",
                                "Scope 1 & 2 Base Year vs Current Period"
                            )
                            doc.add_paragraph()
                            doc.add_picture(chart_buffer, width=Inches(5.5))
                            doc.add_paragraph()
                        except Exception as e:
                            print(f"Error adding Scope 1&2 comparison chart: {e}")
                    
                    # Analysis text
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    if change_1_2 > 0:
                        p.add_run(f"Scope 1 & 2 emissions for {entity_name} have increased by {self._format_number(abs(change_1_2))} tCO₂e ({abs(change_pct_1_2):.1f}%) compared to the base year ({scope12_base_year}).")
                    elif change_1_2 < 0:
                        p.add_run(f"Scope 1 & 2 emissions for {entity_name} have decreased by {self._format_number(abs(change_1_2))} tCO₂e ({abs(change_pct_1_2):.1f}%) compared to the base year ({scope12_base_year}).")
                    else:
                        p.add_run(f"Scope 1 & 2 emissions for {entity_name} have remained stable compared to the base year ({scope12_base_year}).")
                else:
                    p = doc.add_paragraph()
                    p.add_run("The reporting period is the base year.")
                    p.runs[0].italic = True
            else:
                p = doc.add_paragraph()
                run = p.add_run("Scope 1 & 2 Base year not defined.")
                run.italic = True
            
            doc.add_paragraph()
            
            # ==================== SCOPE 3 BASE YEAR SECTION ====================
            p = doc.add_paragraph()
            run = p.add_run("Scope 3 Base Year Emissions")
            run.bold = True
            run.font.size = Pt(12)
            doc.add_paragraph()
            
            if has_scope3_base_year and scope3_base_year_str:
                p = doc.add_paragraph()
                run = p.add_run(f"Base Year: ")
                run.bold = True
                p.add_run(str(scope3_base_year_str))
                doc.add_paragraph()
                
                # Create Scope 3 base year table
                if scope3_emissions_data:
                    headers = ['Scope', 'Category', 'Subcategory', 'Emissions (tCO₂e)']
                    data = []
                    for em in scope3_emissions_data:
                        scope = em.get('scope', 'Scope 3')
                        category = em.get('category', '')
                        subcategory = em.get('subcategory', '')
                        tco2e = float(em.get('tco2e', 0) or 0)
                        display_tco2e = tco2e * equity_factor if use_equity_share else tco2e
                        data.append([scope, category, subcategory, self._format_number(display_tco2e)])
                    
                    if data:
                        data.append(['', '', 'Total Scope 3 Base Year', self._format_number(scope3_base_year_display)])
                        self._create_styled_table(doc, headers, data, bold_rows=[len(data)-1])
                        doc.add_paragraph()
                
                # Scope 3 Comparison
                if not is_reporting_period_base_year(scope3_base_year_str):
                    change_3 = current_scope3 - scope3_base_year_display
                    # Fix: Calculate percentage correctly, handle 0 base year case
                    if scope3_base_year_display > 0:
                        change_pct_3 = (change_3 / scope3_base_year_display) * 100
                    elif current_scope3 > 0:
                        change_pct_3 = 100.0  # 100% increase from 0
                    else:
                        change_pct_3 = 0.0
                    
                    comparison_headers_3 = ['Period', 'Scope 3 Emissions (tCO₂e)']
                    comparison_data_3 = [
                        [f"Base Year ({scope3_base_year_str})", self._format_number(scope3_base_year_display)],
                        ["Current Reporting Period", self._format_number(current_scope3)],
                        ["Change", f"{self._format_number(change_3)} ({'+' if change_3 >= 0 else ''}{change_pct_3:.1f}%)"]
                    ]
                    self._create_styled_table(doc, comparison_headers_3, comparison_data_3, bold_rows=[2])
                    
                    # Add comparison chart for Scope 3
                    if scope3_base_year_display > 0 or current_scope3 > 0:
                        try:
                            chart_buffer = self._create_base_year_comparison_chart(
                                scope3_base_year_display, current_scope3, 
                                f"Base Year ({scope3_base_year_str})", "Current Period",
                                "Scope 3 Base Year vs Current Period"
                            )
                            doc.add_paragraph()
                            doc.add_picture(chart_buffer, width=Inches(5.5))
                            doc.add_paragraph()
                        except Exception as e:
                            print(f"Error adding Scope 3 comparison chart: {e}")
                    
                    # Analysis text
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    if change_3 > 0:
                        p.add_run(f"Scope 3 emissions have increased by {self._format_number(abs(change_3))} tCO₂e ({abs(change_pct_3):.1f}%) compared to the base year ({scope3_base_year_str}).")
                    elif change_3 < 0:
                        p.add_run(f"Scope 3 emissions have decreased by {self._format_number(abs(change_3))} tCO₂e ({abs(change_pct_3):.1f}%) compared to the base year ({scope3_base_year_str}).")
                    else:
                        p.add_run(f"Scope 3 emissions have remained stable compared to the base year ({scope3_base_year_str}).")
                else:
                    p = doc.add_paragraph()
                    p.add_run("The reporting period is the base year.")
                    p.runs[0].italic = True
            else:
                p = doc.add_paragraph()
                run = p.add_run("Scope 3 Base year not defined.")
                run.italic = True
            
        else:
            # ==================== SCOPE 1 & 2 ONLY REPORT ====================
            if has_scope12_base_year and scope12_base_year:
                p = doc.add_paragraph()
                run = p.add_run(f"Base Year: ")
                run.bold = True
                p.add_run(str(scope12_base_year))
                doc.add_paragraph()
                
                # Create base year table
                if scope12_emissions_data:
                    headers = ['Scope', 'Category', 'Subcategory', 'Emissions (tCO₂e)']
                    data = []
                    for em in scope12_emissions_data:
                        scope = em.get('scope', '')
                        category = em.get('category', '')
                        subcategory = em.get('subcategory', '')
                        tco2e = float(em.get('tco2e', 0) or 0)
                        display_tco2e = tco2e * equity_factor if use_equity_share else tco2e
                        data.append([scope, category, subcategory, self._format_number(display_tco2e)])
                    
                    if data:
                        total_display = (scope1_2_base_year_total + biogenic_base_year) * (equity_factor if use_equity_share else 1)
                        data.append(['', '', 'Total Base Year Emissions', self._format_number(total_display)])
                        self._create_styled_table(doc, headers, data, bold_rows=[len(data)-1])
                        doc.add_paragraph()
                
                # Comparison
                if not is_reporting_period_base_year(scope12_base_year):
                    current_net_ghg = current_scope1_2 - current_removals
                    total_base_year_display = scope1_2_base_year_display + (biogenic_base_year * equity_factor if use_equity_share else biogenic_base_year)
                    
                    change = current_net_ghg - total_base_year_display
                    # Fix: Calculate percentage correctly
                    if total_base_year_display > 0:
                        change_pct = (change / total_base_year_display) * 100
                    elif current_net_ghg > 0:
                        change_pct = 100.0
                    else:
                        change_pct = 0.0
                    
                    comparison_headers = ['Period', 'Net GHG Emissions (tCO₂e)']
                    comparison_data = [
                        [f"Base Year ({scope12_base_year})", self._format_number(total_base_year_display)],
                        ["Current Reporting Period", self._format_number(current_net_ghg)],
                        ["Change", f"{self._format_number(change)} ({'+' if change >= 0 else ''}{change_pct:.1f}%)"]
                    ]
                    self._create_styled_table(doc, comparison_headers, comparison_data, bold_rows=[2])
                    
                    # Add comparison chart
                    if total_base_year_display > 0 or current_net_ghg > 0:
                        try:
                            chart_buffer = self._create_base_year_comparison_chart(
                                total_base_year_display, current_net_ghg, 
                                f"Base Year ({scope12_base_year})", "Current Period",
                                "Base Year vs Current Period Emissions"
                            )
                            doc.add_paragraph()
                            doc.add_picture(chart_buffer, width=Inches(5.5))
                            doc.add_paragraph()
                        except Exception as e:
                            print(f"Error adding base year comparison chart: {e}")
                    
                    # Analysis text
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    if change > 0:
                        p.add_run(f"The emissions for {entity_name} have increased by {self._format_number(abs(change))} tCO₂e ({abs(change_pct):.1f}%) compared to the base year ({scope12_base_year}). ")
                        p.add_run("This increase may be attributed to factors such as increased production, expansion of operations, changes in fuel mix, or other operational changes.")
                    elif change < 0:
                        p.add_run(f"The emissions for {entity_name} have decreased by {self._format_number(abs(change))} tCO₂e ({abs(change_pct):.1f}%) compared to the base year ({scope12_base_year}). ")
                        p.add_run("This reduction demonstrates progress in emission management.")
                    else:
                        p.add_run(f"The emissions for {entity_name} have remained stable compared to the base year ({scope12_base_year}).")
                else:
                    p = doc.add_paragraph()
                    p.add_run("The reporting period is the base year.")
                    p.runs[0].italic = True
            else:
                p = doc.add_paragraph()
                run = p.add_run("Base year not defined.")
                run.italic = True
    
    def _create_base_year_comparison_chart(self, base_year_value: float, current_value: float, 
                                           base_year_label: str, current_label: str, title: str) -> io.BytesIO:
        """Create a bar chart comparing base year vs current period emissions"""
        fig, ax = plt.subplots(figsize=(6, 4))
        
        labels = [base_year_label, current_label]
        values = [base_year_value, current_value]
        colors = ['#3498db', '#27ae60']  # Blue for base year, green for current
        
        bars = ax.bar(labels, values, color=colors, width=0.5, edgecolor='black', linewidth=0.5)
        
        # Add value labels on bars
        for bar, value in zip(bars, values):
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height,
                    f'{value:,.2f}',
                    ha='center', va='bottom', fontsize=10, fontweight='bold')
        
        ax.set_ylabel('tCO₂e', fontsize=11)
        ax.set_title(title, fontsize=12, fontweight='bold')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        
        # Set y-axis to start from 0
        ax.set_ylim(bottom=0)
        
        plt.tight_layout()
        
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        buffer.seek(0)
        plt.close(fig)
        
        return buffer
    
    # ==================== CHART GENERATION ====================
    
    def _create_scope_comparison_chart(self, scope1: float, scope2: float, scope3: float = None) -> io.BytesIO:
        """Create Scope 1 vs Scope 2 (and optionally Scope 3) comparison chart"""
        fig, ax = plt.subplots(figsize=(7, 4.5))
        
        if scope3 is not None and scope3 > 0:
            # Scope 1 vs 2 vs 3 chart
            labels = ['Scope 1\n(Direct)', 'Scope 2\n(Indirect)', 'Scope 3\n(Value Chain)']
            values = [scope1, scope2, scope3]
            colors = ['#3498db', '#e74c3c', '#27ae60']
            title = 'Scope 1 vs Scope 2 vs Scope 3 Emissions Comparison'
        else:
            # Scope 1 vs 2 chart only
            labels = ['Scope 1\n(Direct)', 'Scope 2\n(Indirect)']
            values = [scope1, scope2]
            colors = ['#3498db', '#e74c3c']
            title = 'Scope 1 vs Scope 2 Emissions Comparison'
        
        bars = ax.bar(labels, values, color=colors, edgecolor='black', linewidth=1.2)
        
        # Calculate proper offset for text labels to avoid overlap
        max_val = max(values) if max(values) > 0 else 1
        text_offset = max_val * 0.05
        
        for bar, val in zip(bars, values):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + text_offset,
                    f'{val:,.2f}', ha='center', va='bottom', fontsize=9, fontweight='bold')
        
        ax.set_ylabel('tCO2e', fontsize=10)
        ax.set_title(title, fontsize=11, fontweight='bold')
        ax.grid(axis='y', alpha=0.3)
        
        # Add extra space at the top to prevent text overlap with chart border
        y_max = max_val + text_offset + (max_val * 0.15)
        ax.set_ylim(0, y_max)
        
        plt.tight_layout(pad=1.5)
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=120, bbox_inches='tight')
        buf.seek(0)
        plt.close(fig)
        return buf
    
    def _create_scope123_comparison_chart(self, scope12_base: float, scope12_current: float,
                                           scope3_base: float, scope3_current: float,
                                           base_year: str) -> io.BytesIO:
        """Create Scope 1&2 vs Scope 3 base year comparison chart (grouped bar chart)"""
        fig, ax = plt.subplots(figsize=(8, 5))
        
        # Define labels and values
        categories = ['Scope 1 & 2', 'Scope 3']
        base_values = [scope12_base, scope3_base]
        current_values = [scope12_current, scope3_current]
        
        x = np.arange(len(categories))
        width = 0.35
        
        # Create grouped bars
        bars1 = ax.bar(x - width/2, base_values, width, label=f'Base Year ({base_year})', 
                       color='#3498db', edgecolor='black', linewidth=1.2)
        bars2 = ax.bar(x + width/2, current_values, width, label='Current Reporting Period', 
                       color='#e74c3c', edgecolor='black', linewidth=1.2)
        
        # Calculate proper offset for text labels
        all_values = base_values + current_values
        max_val = max(all_values) if max(all_values) > 0 else 1
        text_offset = max_val * 0.03
        
        # Add value labels on top of bars
        for bars in [bars1, bars2]:
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2, height + text_offset,
                        f'{height:,.2f}', ha='center', va='bottom', fontsize=8, fontweight='bold')
        
        ax.set_ylabel('tCO₂e', fontsize=10)
        ax.set_title('Base Year vs Current Period - Emissions Comparison', fontsize=11, fontweight='bold')
        ax.set_xticks(x)
        ax.set_xticklabels(categories)
        ax.legend(loc='upper right')
        ax.grid(axis='y', alpha=0.3)
        
        # Add extra space at the top to prevent text overlap
        y_max = max_val + text_offset + (max_val * 0.2)
        ax.set_ylim(0, y_max)
        
        plt.tight_layout(pad=1.5)
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=120, bbox_inches='tight')
        buf.seek(0)
        plt.close(fig)
        return buf
    
    def _create_category_chart(self, categories: Dict[str, float]) -> io.BytesIO:
        """Create category-wise emission distribution chart with clean legend below.
        
        Features:
        - No labels around the pie chart (cleaner look)
        - Intelligent grouping of small categories into "Others" 
        - Clean legend box below chart with full category names
        - Color swatches and percentages in legend
        """
        if not categories:
            categories = {'No Data': 0}
        
        labels = list(categories.keys())
        values = list(categories.values())
        total = sum(values) if values else 1
        
        # Configuration for "Others" grouping
        MIN_CATEGORY_PERCENT = 2.0  # Categories below this % get grouped into "Others"
        
        # Separate significant categories from small ones
        significant_data = []
        others_data = []
        
        for label, value in zip(labels, values):
            pct = (value / total) * 100 if total > 0 else 0
            if pct >= MIN_CATEGORY_PERCENT:
                significant_data.append((label, value, pct))
            else:
                others_data.append((label, value, pct))
        
        # Calculate combined "Others" percentage
        others_total = sum(v for _, v, _ in others_data)
        others_pct = (others_total / total) * 100 if total > 0 else 0
        
        # Always group small categories into "Others"
        if others_data and others_pct > 0:
            significant_data.append((f'Others ({len(others_data)} categories)', others_total, others_pct))
        
        # Sort by value descending
        significant_data.sort(key=lambda x: x[1], reverse=True)
        
        if significant_data:
            labels = [item[0] for item in significant_data]
            values = [item[1] for item in significant_data]
            percentages = [item[2] for item in significant_data]
        else:
            labels, values, percentages = ['No Data'], [1], [100.0]
        
        # Create figure with space for legend below
        fig = plt.figure(figsize=(8, 6.5))
        
        # Create GridSpec for pie chart (top) and legend (bottom)
        gs = fig.add_gridspec(2, 1, height_ratios=[3, 1.2], hspace=0.05)
        ax_pie = fig.add_subplot(gs[0])
        ax_legend = fig.add_subplot(gs[1])
        
        # Generate colors
        base_colors = plt.cm.Set3(np.linspace(0, 1, max(12, len(labels))))
        colors = base_colors[:len(labels)]
        
        # Special color for "Others" - make it gray
        for i, label in enumerate(labels):
            if 'Others' in label:
                colors[i] = (0.7, 0.7, 0.7, 1.0)  # Gray for "Others"
        
        # Create pie chart WITHOUT labels (clean look)
        wedges, _ = ax_pie.pie(
            values, 
            colors=colors, 
            startangle=90,
            wedgeprops=dict(width=1, edgecolor='white', linewidth=1.5)
        )
        
        # Add percentage labels INSIDE the pie slices (only for significant portions)
        for i, (wedge, pct) in enumerate(zip(wedges, percentages)):
            if pct >= 3:  # Only show percentage for slices >= 3%
                # Calculate position for text inside the slice
                ang = (wedge.theta2 - wedge.theta1) / 2. + wedge.theta1
                x = 0.6 * np.cos(np.deg2rad(ang))
                y = 0.6 * np.sin(np.deg2rad(ang))
                ax_pie.text(x, y, f'{pct:.1f}%', ha='center', va='center', 
                           fontsize=8, fontweight='bold', color='black')
        
        ax_pie.set_title('Category-wise Emission Distribution', fontsize=11, fontweight='bold', pad=10)
        
        # Create legend box below the chart
        ax_legend.axis('off')
        
        # Calculate layout for legend items
        num_items = len(labels)
        cols = 2 if num_items <= 8 else 3
        rows = (num_items + cols - 1) // cols
        
        # Create legend items
        legend_y_start = 0.95
        legend_y_step = 0.95 / max(rows, 1)
        col_width = 1.0 / cols
        
        for i, (label, value, pct) in enumerate(zip(labels, values, percentages)):
            row = i // cols
            col = i % cols
            
            x_pos = col * col_width + 0.02
            y_pos = legend_y_start - (row * legend_y_step)
            
            # Color swatch
            ax_legend.add_patch(plt.Rectangle((x_pos, y_pos - 0.04), 0.03, 0.06, 
                                               facecolor=colors[i], edgecolor='black', linewidth=0.5))
            
            # Truncate very long labels for legend display
            display_label = label[:35] + '...' if len(label) > 35 else label
            
            # Label text with percentage
            ax_legend.text(x_pos + 0.045, y_pos - 0.01, 
                          f'{display_label} ({pct:.1f}%)', 
                          fontsize=7, va='center', ha='left',
                          wrap=True)
        
        ax_legend.set_xlim(0, 1)
        ax_legend.set_ylim(0, 1)
        
        plt.tight_layout(pad=1.5)
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=120, bbox_inches='tight', pad_inches=0.3, facecolor='white')
        buf.seek(0)
        plt.close(fig)
        return buf
    
    def _create_fuel_chart(self, fuels: Dict[str, float]) -> io.BytesIO:
        """Create fuel-wise emission distribution chart as bar chart with totals on top"""
        # Adjust figure width based on number of fuels
        width = max(6, min(10, len(fuels) * 0.8))
        fig, ax = plt.subplots(figsize=(width, 4.5))
        
        if not fuels:
            fuels = {'No Data': 0}
        
        # Sort by value descending for better readability
        sorted_fuels = dict(sorted(fuels.items(), key=lambda x: x[1], reverse=True))
        
        # Limit to top 10 fuels if too many, group rest as "Other"
        if len(sorted_fuels) > 10:
            top_fuels = dict(list(sorted_fuels.items())[:9])
            other_value = sum(list(sorted_fuels.values())[9:])
            if other_value > 0:
                top_fuels['Other'] = other_value
            sorted_fuels = top_fuels
        
        labels = list(sorted_fuels.keys())
        values = list(sorted_fuels.values())
        
        # Use shorter labels if they're too long
        short_labels = [l[:12] + '...' if len(l) > 12 else l for l in labels]
        
        # Use distinct colors for different fuels
        colors = plt.cm.Pastel1(np.linspace(0, 1, len(labels)))
        
        bars = ax.bar(short_labels, values, color=colors, edgecolor='black', linewidth=1.0)
        
        # Calculate proper offset for text labels
        max_val = max(values) if values and max(values) > 0 else 1
        text_offset = max_val * 0.03
        
        # Add value labels on top of each bar (skip very small values to avoid clutter)
        total = sum(values)
        for bar, val in zip(bars, values):
            if val / total >= 0.02 or len(values) <= 5:  # Show label if >= 2% or few bars
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + text_offset,
                        f'{val:,.2f}', ha='center', va='bottom', fontsize=7, fontweight='bold')
        
        ax.set_ylabel('tCO₂e', fontsize=10)
        ax.set_title('Fuel-wise Emission Distribution', fontsize=11, fontweight='bold')
        ax.grid(axis='y', alpha=0.3)
        
        # Always rotate labels if more than 3 items for readability
        if len(labels) > 3:
            plt.xticks(rotation=45, ha='right', fontsize=7)
        else:
            plt.xticks(fontsize=9)
        
        # Add extra space at the top to prevent text overlap
        y_max = max_val + text_offset + (max_val * 0.18)
        ax.set_ylim(0, y_max)
        
        plt.tight_layout(pad=1.5)
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=120, bbox_inches='tight')
        buf.seek(0)
        plt.close(fig)
        return buf
    
    def _create_monthly_trend_chart(self, monthly_data: Dict) -> io.BytesIO:
        """Create monthly emission trend chart"""
        fig, ax = plt.subplots(figsize=(10, 5))
        
        if not monthly_data:
            monthly_data = {'No Data': {'scope1': 0, 'scope2': 0}}
        
        months = self._sort_months(list(monthly_data.keys()))
        scope1_vals = [monthly_data[m]['scope1'] for m in months]
        scope2_vals = [monthly_data[m]['scope2'] for m in months]
        
        x = np.arange(len(months))
        width = 0.35
        
        bars1 = ax.bar(x - width/2, scope1_vals, width, label='Scope 1', color='#3498db')
        bars2 = ax.bar(x + width/2, scope2_vals, width, label='Scope 2', color='#e74c3c')
        
        ax.set_xlabel('Month', fontsize=10)
        ax.set_ylabel('tCO2e', fontsize=10)
        ax.set_title('Monthly Emission Trend', fontsize=11, fontweight='bold', pad=10)
        ax.set_xticks(x)
        ax.set_xticklabels(months, rotation=45, ha='right', fontsize=8)
        ax.legend(fontsize=8, loc='upper right')
        ax.grid(axis='y', alpha=0.3)
        
        # Add margin at top for text if needed
        max_val = max(max(scope1_vals) if scope1_vals else 0, max(scope2_vals) if scope2_vals else 0)
        if max_val > 0:
            ax.set_ylim(0, max_val * 1.15)
        
        plt.tight_layout(pad=1.5)
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=120, bbox_inches='tight')
        buf.seek(0)
        plt.close(fig)
        return buf
    
    def _create_facility_comparison_chart(self, facility_totals: Dict[str, float]) -> io.BytesIO:
        """Create facility comparison chart"""
        fig, ax = plt.subplots(figsize=(9, 5))
        
        if not facility_totals:
            facility_totals = {'No Data': 0}
        
        labels = list(facility_totals.keys())
        values = list(facility_totals.values())
        colors = plt.cm.tab10(np.linspace(0, 1, len(labels)))
        
        # Truncate long facility names
        short_labels = [l[:15] + '...' if len(l) > 15 else l for l in labels]
        
        bars = ax.bar(short_labels, values, color=colors, edgecolor='black')
        
        max_val = max(values) if values and max(values) > 0 else 1
        for bar, val in zip(bars, values):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max_val*0.02,
                    f'{val:,.2f}', ha='center', va='bottom', fontsize=8, fontweight='bold')
        
        ax.set_ylabel('tCO2e', fontsize=10)
        ax.set_title('Facility-wise Emission Comparison', fontsize=11, fontweight='bold', pad=10)
        ax.set_xticklabels(short_labels, rotation=45, ha='right', fontsize=8)
        ax.grid(axis='y', alpha=0.3)
        
        # Add margin at top for text labels
        ax.set_ylim(0, max_val * 1.15)
        
        plt.tight_layout(pad=1.5)
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=120, bbox_inches='tight')
        buf.seek(0)
        plt.close(fig)
        return buf
    
    # ==================== CHAPTER GENERATORS ====================
    
    def _generate_cover_page(self, doc: Document, organization: Dict, reporting_period_start: str, 
                            reporting_period_end: str):
        """Generate cover page with logo and basic info"""
        # Company Name FIRST
        doc.add_paragraph()
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = company_para.add_run(self._get_value_or_na(organization, 'name'))
        run.font.size = Pt(24)
        run.font.bold = True
        
        # Company Logo BELOW the company name
        logo_url = organization.get('logo')
        if logo_url:
            try:
                logo_buffer = self._download_image(logo_url)
                if logo_buffer:
                    doc.add_paragraph()  # Spacing
                    logo_para = doc.add_paragraph()
                    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = logo_para.add_run()
                    run.add_picture(logo_buffer, width=Inches(2.5))
            except Exception as e:
                print(f"Error adding logo: {e}")
        
        # Add extra spacing between logo and report title
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Report Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("Greenhouse Gas (GHG) Inventory Report")
        run.font.size = Pt(18)
        run.font.bold = True
        
        # Reporting Period - Use new format with specific dates (e.g., "1st March 2023 – 30th April 2024")
        period_para = doc.add_paragraph()
        period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        formatted_period = self._format_reporting_period_with_dates(reporting_period_start, reporting_period_end)
        run = period_para.add_run(f"Reporting Period: {formatted_period}")
        run.font.size = Pt(14)
        
        # Subtitle - Only for Scope 1,2 reports
        if getattr(self, 'report_type', 'scope_1_2') != 'scope_1_2_3':
            doc.add_paragraph()
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = subtitle_para.add_run("Prepared as per ISO 14064-1:2018")
            run.font.size = Pt(12)
            run.font.italic = True
        
        # Date of Report Generation (only on cover page)
        doc.add_paragraph()
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"Date of Report Generation: {self.report_date}")
        run.font.size = Pt(11)
        run.font.italic = True
        
        doc.add_page_break()
        
        # Disclaimer - Above Abbreviations
        self._add_styled_heading(doc, "DISCLAIMER", level=1)
        
        org_name = organization.get('name') or 'the Company'
        disclaimer_text = f"This report is generated through SustainRepo platform. The data presented in this report has been provided by {org_name} through the SustainRepo platform. While the information has been compiled as submitted, SustainRepo does not independently verify the accuracy or completeness of the data provided. Accordingly, SustainRepo shall not be held responsible for any inaccuracies, misstatements, or omissions in the information, nor for any resulting consequences, including reputational or financial loss arising from reliance on this report."
        
        p = doc.add_paragraph()
        p.add_run(disclaimer_text)
        p.paragraph_format.space_after = Pt(12)
        
        doc.add_paragraph()
        
        # Abbreviations - in TABLE format
        self._add_styled_heading(doc, "ABBREVIATIONS", level=1)
        
        abbreviations = [
            ("GHG", "Greenhouse Gas"),
            ("CO₂", "Carbon Dioxide"),
            ("CH₄", "Methane"),
            ("N₂O", "Nitrous Oxide"),
            ("tCO₂e", "Tonnes of Carbon Dioxide Equivalent"),
            ("ISO", "International Organization for Standardization"),
            ("Scope 1", "Direct GHG Emissions"),
            ("Scope 2", "Indirect GHG Emissions from Purchased Energy"),
            ("Scope 3", "Other Indirect GHG Emissions (Value Chain)"),
            ("GWP", "Global Warming Potential"),
            ("NCV", "Net Calorific Value"),
            ("EF", "Emission Factor"),
            ("IPCC", "Intergovernmental Panel on Climate Change"),
        ]
        
        # Create abbreviations table
        abbr_headers = ['Abbreviation', 'Full Form']
        abbr_data = [[abbr, meaning] for abbr, meaning in abbreviations]
        self._create_styled_table(doc, abbr_headers, abbr_data, col_widths=[1.5, 5.0])
        
        # Page break - Chapter 1 starts on new page
        doc.add_page_break()
    
    def _generate_chapter1(self, doc: Document, organization: Dict, facilities: List[Dict]):
        """Chapter 1: GENERAL DESCRIPTION OF THE ORGANIZATION AND INVENTORY OBJECTIVES"""
        self._add_styled_heading(doc, "Chapter 1: GENERAL DESCRIPTION OF THE ORGANIZATION AND INVENTORY OBJECTIVES", level=1)
        
        # 1.1 Organization
        self._add_styled_heading(doc, "1.1 Organization", level=2)
        
        # Address in structured format
        p = doc.add_paragraph()
        run = p.add_run("1. Address:")
        run.bold = True
        
        # Street Address
        p = doc.add_paragraph()
        p.add_run("   Street Address: ")
        p.add_run(self._get_value_or_na(organization, 'corporate_address'))
        
        # City
        p = doc.add_paragraph()
        p.add_run("   City: ")
        p.add_run(self._get_value_or_na(organization, 'city'))
        
        # Pincode
        p = doc.add_paragraph()
        p.add_run("   Pincode: ")
        p.add_run(self._get_value_or_na(organization, 'pincode'))
        
        # State
        p = doc.add_paragraph()
        p.add_run("   State: ")
        p.add_run(self._get_value_or_na(organization, 'state'))
        
        # Country
        p = doc.add_paragraph()
        p.add_run("   Country: ")
        p.add_run(self._get_value_or_na(organization, 'country'))
        
        self._add_paragraph_with_bold_label(doc, "2. Organization Description", 
                                           self._get_value_or_na(organization, 'general_description'))
        self._add_paragraph_with_bold_label(doc, "3. Mission of the organization", 
                                           self._get_value_or_na(organization, 'mission'))
        self._add_paragraph_with_bold_label(doc, "4. Vision of the organization", 
                                           self._get_value_or_na(organization, 'vision'))
        self._add_paragraph_with_bold_label(doc, "5. Process Description", 
                                           self._get_value_or_na(organization, 'process_description'))
        
        # NEW SECTION 6: Importance of GHG Reporting
        p = doc.add_paragraph()
        run = p.add_run("6. Importance of GHG Reporting:")
        run.bold = True
        
        p = doc.add_paragraph()
        p.add_run("The need for accounting and reporting greenhouse gas (GHG) emissions has become increasingly important from both a corporate sustainability and operational efficiency perspective. Furthermore, it demonstrates industry commitment to consistent and transparent GHG accounting and reporting practices across organizations and programs. It also encourages stakeholder engagement, feedback, and dialogue, supporting collaborative efforts toward GHG mitigation.")
        
        p = doc.add_paragraph()
        p.add_run("The information generated through GHG reporting can be used to improve business processes, strengthen strategies, and guide actionable initiatives for emission reduction, while enhancing overall environmental performance.")
        
        # NEW SECTION 7: Introduction to the GHG Protocol
        p = doc.add_paragraph()
        run = p.add_run("7. Introduction to the GHG Protocol:")
        run.bold = True
        
        p = doc.add_paragraph()
        p.add_run("The Greenhouse Gas (GHG) Protocol is the world's most widely used framework for measuring, managing, and reporting greenhouse gas emissions. Developed through a partnership between the World Resources Institute (WRI) and the World Business Council for Sustainable Development (WBCSD), it provides standardized methodologies that enable organizations to prepare comprehensive and consistent GHG inventories.")
        
        p = doc.add_paragraph()
        p.add_run("The GHG Protocol helps organizations identify emission sources across their operations and value chains, categorize emissions into Scope 1, Scope 2, and Scope 3, and report emissions in a transparent and comparable manner. By following the GHG Protocol, organizations can better understand their carbon footprint, establish emission reduction targets, track progress over time, support sustainability initiatives, and meet stakeholder, customer, investor, and regulatory reporting requirements. The framework also serves as the foundation for many corporate climate disclosure programs, carbon management strategies, and net-zero commitments worldwide.")
        
        # NEW SECTION 8: Importance of GHG Management Systems
        p = doc.add_paragraph()
        run = p.add_run("8. Importance of GHG Management Systems:")
        run.bold = True
        
        p = doc.add_paragraph()
        p.add_run("In the current regulatory and environmental landscape, assessing and mitigating GHG emissions is essential for statutory compliance. A structured GHG accounting, management, and reporting system allows organizations to monitor energy usage at specific sources, identify emission reduction opportunities, and implement internal efficiency projects.")
        
        p = doc.add_paragraph()
        p.add_run("Such systems not only improve environmental performance but also help reduce operational costs by eliminating inefficiencies and avoiding duplication in data and processes.")
        
        # Continue with renumbered fields
        # Person Responsible with designation and contact details
        person_responsible = self._get_value_or_na(organization, 'person_responsible')
        person_designation = organization.get('person_responsible_designation', '')
        person_contact = organization.get('person_responsible_contact', '')
        
        # Build person responsible text with optional designation and contact
        person_text = person_responsible
        if person_designation:
            person_text += f"\nDesignation: {person_designation}"
        if person_contact:
            person_text += f"\nContact: {person_contact}"
        
        self._add_paragraph_with_bold_label(doc, "9. Person Responsible", person_text)
        self._add_paragraph_with_bold_label(doc, "10. Purpose of Reporting", 
                                           self._get_value_or_na(organization, 'report_purpose'))
        # self._add_paragraph_with_bold_label(doc, "11. Reporting Frequency", 
        #                                    self._get_value_or_na(organization, 'reporting_frequency', 'Yearly').capitalize())
        reporting_frequency = self._get_value_or_na(
            organization,
            'reporting_frequency',
            'Yearly'
        ).capitalize()

        self._add_paragraph_with_bold_label(
            doc,
            "11. Reporting Frequency",
            f"The organization reports greenhouse gas emissions and removals on a "
            f"{reporting_frequency.lower()} basis to support consistent monitoring, "
            f"performance evaluation, and disclosure of climate-related information."
        )
        self._add_paragraph_with_bold_label(doc, "12. Number of Facilities", str(len(facilities)))
        self._add_paragraph_with_bold_label(doc, "13. Other Information", 
                                           self._get_value_or_na(organization, 'other_information'))
        
        # NEW: GHG Accounting Principles (after Other Information)
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("There are five principles on which GHG accounting and reporting is based:")
        
        ghg_principles = [
            ("Relevance:", "Boundaries should be defined to ensure the report meets user needs."),
            ("Completeness:", "All emissions and removals within boundaries must be included; any exclusions must be disclosed and justified."),
            ("Consistency:", "Use consistent methodologies to allow meaningful comparisons; any changes must be disclosed and justified."),
            ("Transparency:", "All relevant information, assumptions, and methodologies should be clearly documented for reliability."),
            ("Accuracy:", "Ensure sufficient accuracy for decision-making; uncertainties should be minimized as far as practical.")
        ]
        
        for principle_name, principle_desc in ghg_principles:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f"{principle_name} ")
            run.bold = True
            p.add_run(principle_desc)
        
        doc.add_paragraph()
        
        # 1.2 Facilities
        self._add_styled_heading(doc, "1.2 Facilities", level=2)
        
        for i, facility in enumerate(facilities, 1):
            facility_name = self._get_value_or_na(facility, 'name')
            self._add_styled_heading(doc, f"1.2.{i} {facility_name}", level=3)
            
            # self._add_labeled_field(doc, "a) Sector/Industry", 
            #                        self._get_value_or_na(facility, 'sector'))

            sector = self._get_value_or_na(facility, 'sector')
            p = doc.add_paragraph()
            p.add_run("a) Sector/Industry:").bold = True

            p = doc.add_paragraph()
            p.add_run(
                f"This facility is classified under the {sector} sector. The greenhouse gas "
                f"inventory covers emissions and removals associated with activities undertaken "
                f"within this sector during the reporting period."
            )
            
            # Facility address in structured format
            p = doc.add_paragraph()
            p.add_run("b) Address:").bold = True
            
            p = doc.add_paragraph()
            p.add_run("   Street Address: ")
            p.add_run(self._get_value_or_na(facility, 'address'))
            
            p = doc.add_paragraph()
            p.add_run("   City: ")
            p.add_run(self._get_value_or_na(facility, 'city'))
            
            p = doc.add_paragraph()
            p.add_run("   Pincode: ")
            p.add_run(self._get_value_or_na(facility, 'pincode'))
            
            p = doc.add_paragraph()
            p.add_run("   State: ")
            p.add_run(self._get_value_or_na(facility, 'state'))
            
            p = doc.add_paragraph()
            p.add_run("   Country: ")
            p.add_run(self._get_value_or_na(facility, 'country'))
            
            self._add_labeled_field(doc, "c) Products/Services", 
                                   self._get_value_or_na(facility, 'products_services') or 
                                   self._get_value_or_na(facility, 'products_manufactured'))
            
            self._add_labeled_field(doc, "d) Machinery and Equipment", 
                                   self._get_value_or_na(facility, 'machinery_equipment') or 
                                   self._get_value_or_na(facility, 'machinery_used'))
            
            self._add_labeled_field(doc, "e) Process Description", 
                                   self._get_value_or_na(facility, 'process_description'))
            
            # Person Responsible with designation and contact details
            facility_person = self._get_value_or_na(facility, 'responsible_person')
            facility_designation = facility.get('responsible_person_designation', '')
            facility_contact = facility.get('responsible_person_contact', '')
            
            facility_person_text = facility_person
            if facility_designation:
                facility_person_text += f"\n   Designation: {facility_designation}"
            if facility_contact:
                facility_person_text += f"\n   Contact: {facility_contact}"
            
            self._add_labeled_field(doc, "f) Person Responsible", facility_person_text)
            
            # self._add_labeled_field(doc, "g) Monitoring Frequency", 
            #                        self._get_value_or_na(facility, 'monitoring_frequency', 'Monthly').capitalize())
            

            monitoring_frequency = self._get_value_or_na(
                facility,
                'monitoring_frequency',
                'Monthly'
            ).capitalize()

            p = doc.add_paragraph()
            p.add_run("g) Monitoring Frequency:").bold = True

            p = doc.add_paragraph()
            p.add_run(
                f"The facility monitors greenhouse gas emission sources and associated activity "
                f"data on a {monitoring_frequency.lower()} basis as part of its GHG inventory management process."
            )

            # self._add_labeled_field(doc, "h) Reporting Frequency", 
            #                        self._get_value_or_na(facility, 'reporting_frequency', 'Monthly').capitalize())
            
            reporting_frequency = self._get_value_or_na(
                facility,
                'reporting_frequency',
                'Monthly'
            ).capitalize()

            p = doc.add_paragraph()
            p.add_run("h) Reporting Frequency:").bold = True

            p = doc.add_paragraph()
            p.add_run(
                f"The facility monitors and reports greenhouse gas emissions on a {reporting_frequency.lower()} basis "
                f"in accordance with the organization's GHG accounting and reporting requirements."
            )

            self._add_labeled_field(doc, "i) Other Information", 
                                   self._get_value_or_na(facility, 'other_information') or 
                                   self._get_value_or_na(facility, 'remarks'))
        
        doc.add_page_break()
    
    def _generate_chapter2(self, doc: Document, organization: Dict):
        """Chapter 2: Organization Boundaries"""
        self._add_styled_heading(doc, "Chapter 2: Organization Boundaries", level=1)
        
        # Introduction text
        p = doc.add_paragraph()
        p.add_run("In greenhouse gas (GHG) accounting, organizations must first establish their organizational boundary to determine which operations, facilities, or subsidiaries are included in the GHG inventory. The organizational boundary defines the extent of the operations for which the organization is responsible for reporting emissions.")
        
        p = doc.add_paragraph()
        p.add_run("According to internationally recognized frameworks such as the GHG Protocol Corporate Standard and ISO 14064-1:2018, organizations can determine their organizational boundaries using two primary approaches:")
        
        # Equity Share Approach
        p = doc.add_paragraph()
        run = p.add_run("Equity Share Approach")
        run.bold = True
        
        p = doc.add_paragraph()
        p.add_run("Under the Equity Share Approach, an organization accounts for greenhouse gas emissions from operations based on its proportionate share of equity ownership in those operations.")
        
        p = doc.add_paragraph()
        p.add_run("This means that the organization reports emissions in proportion to the percentage of ownership or economic interest it holds in a facility, joint venture, or subsidiary. The emissions attributed to the organization are therefore aligned with its financial stake in the operation.")
        
        p = doc.add_paragraph()
        p.add_run("This approach is particularly useful for organizations involved in joint ventures, partnerships, or shared ownership arrangements, where multiple entities have a financial interest in the same operation. By allocating emissions proportionally, the equity share approach ensures that emissions reporting reflects the economic reality of ownership and investment.")
        
        # Control Approach
        p = doc.add_paragraph()
        run = p.add_run("Control Approach")
        run.bold = True
        
        p = doc.add_paragraph()
        p.add_run("Under the Control Approach, an organization accounts for 100% of the greenhouse gas emissions from operations over which it exercises control, regardless of its equity ownership in those operations.")
        
        p = doc.add_paragraph()
        p.add_run("Control can be defined in two ways:")
        
        p = doc.add_paragraph()
        run = p.add_run("Operational Control: ")
        run.bold = True
        p.add_run("The organization has the authority to introduce and implement operating policies at the facility or operation. In this case, the organization reports 100% of the emissions from that operation, even if its ownership stake is less than 100%.")
        
        p = doc.add_paragraph()
        run = p.add_run("Financial Control: ")
        run.bold = True
        p.add_run("The organization has the ability to direct the financial and operating policies of the operation with the intention of gaining economic benefits from its activities.")
        
        p = doc.add_paragraph()
        p.add_run("Under the control approach, emissions from operations where the organization does not have operational or financial control are not included in its GHG inventory, even if the organization holds a partial ownership stake.")
        
        # Importance section
        p = doc.add_paragraph()
        run = p.add_run("Importance of Selecting a Consistent Approach")
        run.bold = True
        
        p = doc.add_paragraph()
        p.add_run("Organizations must select one of these approaches and apply it consistently across their GHG inventory to ensure transparency, comparability, and consistency in emissions reporting. The chosen approach should also be clearly documented in the GHG report along with any assumptions or criteria used to determine ownership or control.")
        
        p = doc.add_paragraph()
        p.add_run("Establishing a well-defined organizational boundary is a critical first step in the GHG accounting process, as it determines which emission sources are included in the inventory and ensures that emissions are reported accurately in accordance with recognized international standards.")
        
        # Add a simple statement about which approach was chosen
        approach = (organization.get('org_boundaries_approach') or '').lower()
        org_name = self._get_value_or_na(organization, 'name')
        
        if approach == 'equity_share':
            p = doc.add_paragraph()
            run = p.add_run(f"{org_name} has adopted the Equity Share Approach for this GHG inventory.")
            run.bold = True

            p = doc.add_paragraph()
            p.add_run(
                "The organization accounts for greenhouse gas emissions from operations "
                "according to its equity share in each operation. Emissions are attributed "
                "based on the organization's percentage ownership or economic interest in "
                "the operation, regardless of operational or financial control."
            )

        elif approach in ['control', 'control_operational', 'control_financial', 'control_both']:
            if approach == 'control_operational':
                approach_name = "Operational Control"
                approach_desc = "The organization accounts for 100% of greenhouse gas emissions from operations over which it exercises operational control, i.e., full authority to introduce and implement operating policies."
            elif approach == 'control_financial':
                approach_name = "Financial Control"
                approach_desc = "The organization accounts for 100% of greenhouse gas emissions from operations over which it exercises financial control, i.e., the ability to direct the financial and operating policies of an operation."
            elif approach == 'control_both':
                approach_name = "Operational & Financial Control"
                approach_desc = "The organization accounts for 100% of greenhouse gas emissions from operations over which it exercises both operational and financial control."
            else:
                approach_name = "Control"
                approach_desc = "The organization accounts for 100% of greenhouse gas emissions from operations over which it has control."
            
            p = doc.add_paragraph()
            run = p.add_run(f"{org_name} has adopted the {approach_name} Approach for this GHG inventory.")
            run.bold = True
            
            p = doc.add_paragraph()
            p.add_run(approach_desc)
        
        # Organization's detailed boundary approach explanation (only if approach is specified with equity percentage)
        equity_percentage = organization.get('org_boundaries_equity_percentage')
        additional_notes = organization.get('org_boundaries') or organization.get('org_boundaries_notes', '')
        
        # Only add detailed explanation if equity share with specific percentage
        if approach == 'equity_share' and equity_percentage:
            p = doc.add_paragraph()
            run = p.add_run(f"{org_name}")
            run.bold = True
            p.add_run(" has chosen the ")
            run2 = p.add_run("Equity Share Approach")
            run2.bold = True
            p.add_run(f". The organization accounts for greenhouse gas emissions in proportion to its equity share of {equity_percentage}%, meaning {equity_percentage}% of total emissions from joint operations are attributed to the organization based on its ownership stake.")
        
        # Add additional boundary notes on the next line
        if additional_notes and additional_notes != 'NA':
            p = doc.add_paragraph()
            run = p.add_run("Additional Boundary Notes: ")
            run.bold = True
            p.add_run(additional_notes)
        
        doc.add_page_break()
    
    def _generate_chapter3(self, doc: Document, facilities: List[Dict], emissions: List[Dict],
                          reporting_period_start: str = None, reporting_period_end: str = None):
        """Chapter 3: Reporting Boundaries"""
        self._add_styled_heading(doc, "Chapter 3: Reporting Boundaries", level=1)
        
        # Check if this is a Scope 3 report
        is_scope3_report = getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3'
        
        # Introductory paragraph
        p = doc.add_paragraph()
        p.add_run("After establishing the organizational boundary, the organization identifies all emission sources associated with its operations. The organizational boundary is determined based on the organization's ownership interest or level of control over operational activities, in accordance with recognized GHG accounting standards.")
        
        p = doc.add_paragraph()
        if is_scope3_report:
            p.add_run("Once the boundary is defined, the organization systematically reviews its operational activities to identify all relevant greenhouse gas (GHG) emission sources. These sources are then classified according to internationally accepted GHG accounting categories, primarily Scope 1, Scope 2 and Scope 3 emissions. This classification helps define the scope of accounting and reporting within the GHG inventory.")
        else:
            p.add_run("Once the boundary is defined, the organization systematically reviews its operational activities to identify all relevant greenhouse gas (GHG) emission sources. These sources are then classified according to internationally accepted GHG accounting categories, primarily Direct (Scope 1) and Indirect (Scope 2) emissions. This classification helps define the scope of accounting and reporting within the GHG inventory.")
        
        # Scope 1 Emissions
        p = doc.add_paragraph()
        if is_scope3_report:
            run = p.add_run("Scope 1 Emissions")
        else:
            run = p.add_run("Direct GHG Emissions/Removals (Scope 1)")
        run.bold = True
        
        p = doc.add_paragraph()
        p.add_run("Scope 1 emissions refer to direct greenhouse gas emissions from sources that are owned or controlled by the organization. These emissions occur as a direct result of the organization's operational activities.")
        
        # p = doc.add_paragraph()
        # p.add_run("Typical examples of Scope 1 emissions include:")
        
        # scope1_examples = [
        #     "Fuel combustion in stationary sources, such as boilers, furnaces, generators, and industrial equipment.",
        #     "Fuel combustion in mobile sources, including company-owned vehicles and fleet operations using fuels such as diesel or petrol.",
        #     "Process emissions arising from industrial or chemical processes during manufacturing or production activities.",
        #     "Fugitive emissions, such as leakage of refrigerants from air conditioning systems, refrigeration units, or other equipment."
        # ]
        
        # for example in scope1_examples:
        #     doc.add_paragraph(example, style='List Bullet')
        
        p = doc.add_paragraph()
        p.add_run("Since these emission sources are directly controlled by the organization, the organization is responsible for measuring, managing, and reporting these emissions as part of its GHG inventory.")
        
        # Scope 2 Emissions
        p = doc.add_paragraph()
        if is_scope3_report:
            run = p.add_run("Scope 2 Emissions")
        else:
            run = p.add_run("Indirect GHG Emissions (Scope 2)")
        run.bold = True
        
        p = doc.add_paragraph()
        p.add_run("Scope 2 emissions are indirect greenhouse gas emissions associated with the consumption of purchased or acquired energy by the organization. Although these emissions physically occur at the facility where the energy is generated (such as a power plant), they are attributed to the organization because the energy is consumed in its operations.")
        
        p = doc.add_paragraph()
        p.add_run("Scope 2 emissions primarily include emissions from the generation of:")
        
        scope2_examples = [
            "Purchased electricity",
            "Purchased steam",
            "Purchased heating",
            "Purchased cooling"
        ]
        
        for example in scope2_examples:
            doc.add_paragraph(example, style='List Bullet')
        
        p = doc.add_paragraph()
        p.add_run("These emissions are calculated based on the amount of energy consumed by the organization and the corresponding emission factors associated with energy generation.")
        
        # Scope 3 Emissions - Only for scope_1_2_3 reports
        if getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3':
            p = doc.add_paragraph()
            run = p.add_run("Scope 3 Emissions")
            run.bold = True
            
            p = doc.add_paragraph()
            p.add_run("Scope 3 emissions are indirect greenhouse gas emissions that occur across the organization's value chain and are not included within Scope 1 or Scope 2 emissions. These emissions arise from activities associated with the organization's operations but occur from sources that are owned or controlled by external parties such as suppliers, service providers, customers, or end users.")
            
            p = doc.add_paragraph()
            p.add_run("Scope 3 emissions often represent a significant portion of an organization's total carbon footprint and include both upstream and downstream activities throughout the value chain.")
            
            p = doc.add_paragraph()
            p.add_run("Typical examples of Scope 3 emissions include:")
            
            # scope3_examples = [
            #     "Purchased goods and services acquired by the organization from suppliers.",
            #     "Capital goods such as machinery, equipment, and infrastructure purchased during the reporting period.",
            #     "Fuel- and energy-related activities not already included in Scope 1 or Scope 2 emissions.",
            #     "Transportation and distribution of goods through third-party logistics providers.",
            #     "Waste generated in operations and its treatment or disposal by external contractors.",
            #     "Business travel undertaken by employees through commercial transportation services.",
            #     "Employee commuting between homes and workplaces.",
            #     "Upstream and downstream leased assets not directly controlled by the organization.",
            #     "Processing, use, and end-of-life treatment of sold products.",
            #     "Investments and financed emissions, where applicable."
            # ]
            
            # for example in scope3_examples:
            #     doc.add_paragraph(example, style='List Bullet')
            
            p = doc.add_paragraph()
            p.add_run("Since Scope 3 emissions originate outside the organization's direct operational control, their calculation often requires the use of supplier data, activity-based data, spend-based methodologies, industry-average emission factors, or hybrid calculation approaches.")
        
        # For each facility
        for i, facility in enumerate(facilities, 1):
            facility_id = facility.get('id')
            facility_name = self._get_value_or_na(facility, 'name')
            facility_emissions = self._get_emissions_by_facility(emissions, facility_id)
            
            # Filter by reporting period (same as Chapter 4)
            if reporting_period_start and reporting_period_end:
                facility_emissions = self._filter_emissions_by_period(
                    facility_emissions, reporting_period_start, reporting_period_end
                )
            
            self._add_styled_heading(doc, f"3.{i} {facility_name}", level=2)
            
            # Check if facility has emissions
            if not facility_emissions:
                p = doc.add_paragraph()
                run = p.add_run("No emission reported for this facility.")
                run.italic = True
                continue
            
            # 3.x.1 List of Emissions
            self._add_styled_heading(doc, f"3.{i}.1 List of Emissions", level=3)
            
            # Get categorized Scope 1 emissions with extended data
            scope1_by_category_ext = self._get_emission_processes_by_category_extended(facility_emissions)
            scope2_processes_ext = self._get_scope2_processes_extended(facility_emissions)
            
            # Get Scope 3 processes (only for scope_1_2_3 report type)
            scope3_by_category_ext = {}
            if getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3':
                scope3_by_category_ext = self._get_scope3_processes_by_category_extended(facility_emissions)
            
            # Check for data presence
            has_scope1 = any(scope1_by_category_ext.get(cat) for cat in ['stationary_combustion', 'mobile_combustion', 'fugitive_emissions', 'other'])
            has_scope3 = bool(scope3_by_category_ext) and getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3'
            
            # Create 4-column table for emissions list
            self._add_emissions_list_table_4col(doc, scope1_by_category_ext, scope2_processes_ext, scope3_by_category_ext, has_scope1, has_scope3, getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3')
            
            # 3.x.2 Process Overview
            self._add_styled_heading(doc, f"3.{i}.2 Process Overview", level=3)
            self._add_process_overview_table(doc, facility_emissions)
        
        doc.add_page_break()
    
    def _generate_chapter4(self, doc: Document, organization: Dict, facilities: List[Dict], 
                          emissions: List[Dict], reporting_period_start: str, reporting_period_end: str,
                          include_previous_years: bool = True):
        """Chapter 4: QUANTIFIED GHG INVENTORY OF EMISSIONS AND REMOVALS"""
        self._add_styled_heading(doc, "Chapter 4: QUANTIFIED GHG INVENTORY OF EMISSIONS AND REMOVALS", level=1)
        
        # Introductory paragraph for Chapter 4 (BEFORE Section 4.1)
        p = doc.add_paragraph()
        p.add_run("This chapter includes quantified data results by emission or removal category, descriptions of methodologies and activity data used, references and/or explanations of emission and removal factors, uncertainties and their impact on results (disaggregated by category), and a description of planned actions for reducing uncertainty in future inventories.")
        
        doc.add_paragraph()
        
        # Check if organization uses equity share approach
        use_equity_share = organization.get('org_boundaries_approach') == 'equity_share'
        
        # Build facility equity share map
        facility_equity_map = {}
        for f in facilities:
            equity_pct = f.get('equity_share_percentage', 100.0) or 100.0
            facility_equity_map[f.get('id')] = equity_pct
        
        # 4.1 Methodology
        self._add_styled_heading(doc, "4.1 Methodology", level=2)
        
        # Check if this is a Scope 3 report
        is_scope3_report = getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3'
        
        if is_scope3_report:
            # Use tabular methodology format for Scope 1,2,3 reports - 4 columns
            # Scope and Category columns should not be empty - repeat values for merged cell effect
            headers = ['Scope', 'Category', 'Subcategory/\nMethodology', 'Formula']
            data = []
            
            # Scope 1 methodologies - all cells filled
            data.append(['Scope 1', 'Stationary Combustion/\nMobile Combustion', '-', 'Emissions = Quantity of Fuel Consumed × Calorific Value × Emission Factor (Heat Basis) × Density (if applicable)'])
            data.append(['Scope 1', 'Stationary Combustion/\nMobile Combustion', '-', 'Emissions = Quantity of Fuel Consumed × Emission Factor (Quantity Basis)'])
            data.append(['Scope 1', 'Fugitive Emissions', '-', 'Emissions = Quantity of Gas Consumed × GWP'])
            
            # Scope 2 methodology
            data.append(['Scope 2', 'Purchased Electricity/ Heat or Steam', '-', 'Emissions = Quantity of Energy Consumed × Emission Factor (Quantity Basis)'])
            
            # Scope 3 methodologies - Each subcategory/methodology on separate row with × instead of *
            scope3_data = [
                # C1
                ['Scope 3', 'C1 - Purchased Goods and Services', 'Spend Based', 'Emissions = Amount Spent × Emission Factor / (Inflation Rate × Purchase Power Value)'],
                ['Scope 3', 'C1 - Purchased Goods and Services', 'Average Data Based', 'Emissions = Quantity Used × Emission Factor'],
                ['Scope 3', 'C1 - Purchased Goods and Services', 'Supplier Based', 'Emissions = Quantity Used × Emission Factor'],
                # C2
                ['Scope 3', 'C2 - Capital Goods', 'Spend Based', 'Emissions = Amount Spent × Emission Factor / (Inflation Rate × Purchase Power Value)'],
                ['Scope 3', 'C2 - Capital Goods', 'Average Data Based', 'Emissions = Quantity Used × Emission Factor'],
                ['Scope 3', 'C2 - Capital Goods', 'Supplier Based', 'Emissions = Quantity Used × Emission Factor'],
                # C3
                ['Scope 3', 'C3 - Fuel and Energy Related Activities Not Included in Scope 1 or Scope 2', 'Average Data Based', 'Emissions = Quantity Used × (WTT Emission Factor + T&D Loss)'],
                ['Scope 3', 'C3 - Fuel and Energy Related Activities Not Included in Scope 1 or Scope 2', 'Supplier Based', 'Emissions = Quantity Used × Emission Factor'],
                # C4
                ['Scope 3', 'C4 - Upstream Transportation and Distribution', 'Spend Based', 'Emissions = Amount Spent × Emission Factor / (Inflation Rate × Purchase Power Value)'],
                ['Scope 3', 'C4 - Upstream Transportation and Distribution', 'Average Data Based', 'Emissions = Emission Factor × Distance travelled × Quantity of Goods travelled'],
                ['Scope 3', 'C4 - Upstream Transportation and Distribution', 'Supplier Based', 'Emissions = Quantity Used × Emission Factor'],
                # C5
                ['Scope 3', 'C5 - Waste Generated in Operations', 'Average Data Based', 'Emissions = Quantity Used × Emission Factor'],
                ['Scope 3', 'C5 - Waste Generated in Operations', 'Supplier Based', 'Emissions = Quantity Used × Emission Factor'],
                # C6
                ['Scope 3', 'C6 - Business Travel', 'Average Data Based – Hotel Stay', 'Emissions = Emission Factor × No. of room taken × No. of nights stayed'],
                ['Scope 3', 'C6 - Business Travel', 'Average Data Based – Air Travel, Water Travel, Taxi Travel, Bus Travel, Train Travel', 'Emissions = Emission Factor × No. of passengers Travelled × Distance Travelled'],
                ['Scope 3', 'C6 - Business Travel', 'Average Data Based – Car Travel, Bike Travel', 'Emissions = Distance Travelled × Emission Factor'],
                ['Scope 3', 'C6 - Business Travel', 'Supplier Based', 'Emissions = Quantity Used × Emission Factor'],
                # C7
                ['Scope 3', 'C7 - Employee Commuting', 'Average Data Based – Air Travel, Water Travel, Taxi Travel, Bus Travel, Train Travel', 'Emissions = Emission Factor × No. of passengers Travelled × Distance Travelled'],
                ['Scope 3', 'C7 - Employee Commuting', 'Average Data Based – Car Travel, Bike Travel', 'Emissions = Distance Travelled × Emission Factor'],
                ['Scope 3', 'C7 - Employee Commuting', 'Average Data Based – Work From Home', 'Emissions = Emission Factor × Working Days × Working Hours per day'],
                ['Scope 3', 'C7 - Employee Commuting', 'Supplier Based', 'Emissions = Quantity Used × Emission Factor'],
                # C8, C10, C13, C14 (without C11)
                ['Scope 3', 'C8 - Upstream Leased Assets, C10 - Processing of Sold Products, C13 - Downstream Leased Assets, C14 - Franchises', 'Average Data Based – Stationary Combustion', 'Emissions = Quantity of Fuel Consumed × Emission Factor'],
                ['Scope 3', 'C8 - Upstream Leased Assets, C10 - Processing of Sold Products, C13 - Downstream Leased Assets, C14 - Franchises', 'Average Data Based – Mobile Combustion', 'Emissions = Quantity of Fuel Consumed × Emission Factor'],
                ['Scope 3', 'C8 - Upstream Leased Assets, C10 - Processing of Sold Products, C13 - Downstream Leased Assets, C14 - Franchises', 'Average Data Based – Fugitive Emissions', 'Emissions = Quantity of Gas Consumed × GWP'],
                ['Scope 3', 'C8 - Upstream Leased Assets, C10 - Processing of Sold Products, C13 - Downstream Leased Assets, C14 - Franchises', 'Average Data Based – Energy', 'Emissions = Quantity Consumed × Emission Factor'],
                ['Scope 3', 'C8 - Upstream Leased Assets, C10 - Processing of Sold Products, C13 - Downstream Leased Assets, C14 - Franchises', 'Supplier Based', 'Emissions = Quantity Used × Emission Factor'],
                # C9
                ['Scope 3', 'C9 - Downstream Transportation and Distribution', 'Spend Based', 'Emissions = Amount Spent × Emission Factor / (Inflation Rate × Purchase Power Value)'],
                ['Scope 3', 'C9 - Downstream Transportation and Distribution', 'Average Data Based', 'Emissions = Emission Factor × Distance travelled × Quantity of Goods travelled'],
                ['Scope 3', 'C9 - Downstream Transportation and Distribution', 'Supplier Based', 'Emissions = Quantity Used × Emission Factor'],
                # C11 - Use of Sold Products (separate section)
                ['Scope 3', 'C11 - Use of Sold Products', 'Average Data Based – Energy Consuming Product Over Lifetime', 'Emissions = No. of units of products produced in reporting period × Lifetime Expected Usage of the product × Emission Factor/GWP × Quantity Used per usage'],
                ['Scope 3', 'C11 - Use of Sold Products', 'Average Data Based – Stationary Combustion – One Time Combustion', 'Emissions = Quantity of Fuel Consumed × Emission Factor'],
                ['Scope 3', 'C11 - Use of Sold Products', 'Average Data Based – Mobile Combustion – One Time Combustion', 'Emissions = Quantity of Fuel Consumed × Emission Factor'],
                ['Scope 3', 'C11 - Use of Sold Products', 'Average Data Based – Fugitive Emissions – One Time Combustion', 'Emissions = Quantity of Gas Consumed × GWP'],
                ['Scope 3', 'C11 - Use of Sold Products', 'Average Data Based – Energy – One Time Combustion', 'Emissions = Quantity of Fuel Consumed × Emission Factor'],
                ['Scope 3', 'C11 - Use of Sold Products', 'Supplier Based', 'Emissions = Quantity Used × Emission Factor'],
                # C12
                ['Scope 3', 'C12 - End-of-Life Treatment of Sold Products', 'Average Data Based', 'Emissions = Quantity Used × Emission Factor'],
                ['Scope 3', 'C12 - End-of-Life Treatment of Sold Products', 'Average Data Based - Energy', 'Emissions = Energy Used × (Emission Factor + WTT Emission Factor + T&D Loss Emission Factor)'],
                ['Scope 3', 'C12 - End-of-Life Treatment of Sold Products', 'Supplier Based', 'Emissions = Quantity Used × Emission Factor'],
                # C15
                ['Scope 3', 'C15 - Investments', 'Supplier Based', 'Emissions = Quantity Used × Emission Factor'],
            ]
            
            data.extend(scope3_data)
            
            # Biogenic Emissions
            data.append(['Biogenic', '-', '-', 'Emissions = Quantity Used × Emission Factor'])
            
            self._create_styled_table(doc, headers, data)
            
            doc.add_paragraph()
        else:
            # Use text-based methodology format for Scope 1,2 reports
            p = doc.add_paragraph()
            p.add_run("The greenhouse gas (GHG) emissions inventory has been developed using a bottom-up approach, where emissions are calculated based on activity-level data collected from individual emission sources within the organization.")
            
            # Scope 1 – Direct Emissions
            p = doc.add_paragraph()
            run = p.add_run("Scope 1 – Direct Emissions (Fuel Combustion)")
            run.bold = True
            
            p = doc.add_paragraph()
            p.add_run("Direct emissions from stationary or mobile fuel combustion sources are calculated using activity data such as fuel consumption. Where emission factors are provided on an energy basis, the fuel quantity is converted into energy using the calorific value and density of the fuel.")
            
            p = doc.add_paragraph()
            p.add_run("The calculation methodology is as follows:")
            
            p = doc.add_paragraph()
            run = p.add_run("Energy-Based Emission Factor Approach")
            run.bold = True
            
            p = doc.add_paragraph()
            p.add_run("Emissions = Quantity of Fuel Consumed × Calorific Value × Density (if applicable) × Emission Factor")
            
            p = doc.add_paragraph()
            p.add_run("Where:")
            
            where_points_1 = [
                "Quantity of Fuel Consumed refers to the measured amount of fuel used.",
                "Density is applied where fuels are measured by volume.",
                "Calorific Value converts the fuel quantity into energy content.",
                "Emission Factor represents the amount of GHG emitted per unit of energy."
            ]
            
            for point in where_points_1:
                doc.add_paragraph(point, style='List Bullet')
            
            # Scope 1, Scope 2, and Biogenic Emissions
            p = doc.add_paragraph()
            run = p.add_run("Scope 1, Scope 2, and Biogenic Emissions (Quantity-Based Factors)")
            run.bold = True
            
            p = doc.add_paragraph()
            p.add_run("For emission sources where emission factors are directly available on a quantity basis, emissions are calculated using a simpler approach:")
            
            p = doc.add_paragraph()
            p.add_run("Emissions = Activity Data × Emission Factor")
            
            p = doc.add_paragraph()
            p.add_run("Where:")
            
            where_points_2 = [
                "Activity Data represents the quantity of fuel, electricity, or other emission-generating activities.",
                "Emission Factor represents the amount of GHG emitted per unit of activity."
            ]
            
            for point in where_points_2:
                doc.add_paragraph(point, style='List Bullet')
            
            p = doc.add_paragraph()
            p.add_run("This methodology is typically applied for:")
            
            applied_for_points = [
                "Scope 1 emissions such as fuel combustion or refrigerant leakage.",
                "Scope 2 emissions arising from purchased electricity, steam, heating, or cooling.",
                "Biogenic emissions associated with biomass or biofuels."
            ]
            
            for point in applied_for_points:
                doc.add_paragraph(point, style='List Bullet')
        
        # Total Emissions Calculation
        p = doc.add_paragraph()
        run = p.add_run("Total Emissions Calculation")
        run.bold = True
        
        p = doc.add_paragraph()
        p.add_run("Greenhouse gas emissions are calculated individually for the major GHGs. These gases are then converted into a common unit of carbon dioxide equivalent (CO₂e) using their respective Global Warming Potentials (GWP).")
        
        p = doc.add_paragraph()
        p.add_run("The total emissions are calculated as follows:")
        
        p = doc.add_paragraph()
        p.add_run("tCO₂e = tCO₂ + tCH₄ × GWP(CH₄) + tN₂O × GWP(N₂O)")
        
        p = doc.add_paragraph()
        p.add_run("Where:")
        
        where_points_3 = [
            "tCO₂ = tonnes of carbon dioxide emitted",
            "tCH₄ = tonnes of methane emitted",
            "tN₂O = tonnes of nitrous oxide emitted",
            "GWP = Global Warming Potential relative to CO₂"
        ]
        
        for point in where_points_3:
            doc.add_paragraph(point, style='List Bullet')
        
        p = doc.add_paragraph()
        p.add_run("This conversion ensures that emissions from different gases are aggregated into a single standardized metric (tCO₂e) for reporting and comparison.")
        
        # Data Sources and Standards
        p = doc.add_paragraph()
        run = p.add_run("Data Sources and Standards")
        run.bold = True
        
        p = doc.add_paragraph()
        p.add_run("The emission calculations are based on internationally recognized standards and scientific references, ensuring methodological consistency and reliability. The following sources have been used:")
        
        # Use different sources list based on report type
        if is_scope3_report:
            sources_points = [
                "Emission Factors: Derived from the IPCC Guidelines for National Greenhouse Gas Inventories and applicable national emission factor databases.",
                "Global Warming Potentials (GWP): Adopted from the Intergovernmental Panel on Climate Change (IPCC) Sixth Assessment Report (AR6).",
                "Reference Databases and Publications: DEFRA (Department for Environment, Food & Rural Affairs), USEEIO (United States Environmentally-Extended Input-Output Model), TERI (The Energy and Resources Institute), CEA (Central Electricity Authority, India), Indian Railways emission datasets and published conversion factors, USEPA (United States Environmental Protection Agency).",
                "Methodological Framework: Calculations and reporting follow the principles outlined in GHG Protocol – Greenhouse Gases: Specification with guidance at the organization level for quantification and reporting of greenhouse gas emissions and removals."
            ]
        else:
            # For Scope 1,2 report: Remove Reference Databases, and simplify Activity Data
            sources_points = [
                "Emission Factors: Derived from the IPCC Guidelines for National Greenhouse Gas Inventories and applicable national emission factor databases.",
                "Global Warming Potentials (GWP): Adopted from the Intergovernmental Panel on Climate Change (IPCC) Sixth Assessment Report (AR6).",
                "Activity Data: For scope 1, scope 2 is collected from facility operational records, fuel purchase records, energy monitoring systems, and internal documentation.",
                "Methodological Framework: Calculations and reporting follow the principles outlined in ISO 14064-1:2018 – Greenhouse Gases: Specification with guidance at the organization level for quantification and reporting of greenhouse gas emissions and removals."
            ]
        
        for point in sources_points:
            doc.add_paragraph(point, style='List Bullet')
        
        doc.add_paragraph()
        
        # 4.2 Uncertainty Assessment
        self._add_styled_heading(doc, "4.2 Uncertainty Assessment", level=2)
        
        p = doc.add_paragraph()
        p.add_run("GHG inventory data are associated with varying degrees of uncertainty, and such actual uncertainties have both technical and policy implications. For Addressing Uncertainty, \"to ensure that a company's strategies and forward-looking actions are based on the most robust data set and most appropriate computational methods, it is important that this data set and method be based on four key factors (\"The Four C's\"). These are Comparability, Consistency, Certainty, and Confidence. Uncertainties in inventories are the result of three categories:")
        
        uncertainty_categories = [
            "Spurious errors, which may be due to incomplete, unclear, or faulty definitions of emission sources that result from human error or machine malfunction.",
            "Systematic errors, which may be due to the methods (or models) used to quantify emissions for the process under consideration; and",
            "Random errors, which may be due to natural variability of the process that produces the emissions."
        ]
        
        for category in uncertainty_categories:
            p = doc.add_paragraph(category, style='List Bullet')
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run("Uncertainty in quantification of GHG emissions can be on account of uncertainty in available activity data and input parameters used in calculation of emissions. Normally, it is beyond the scope of a company to address the uncertainties in equations that are used for calculating emissions.")
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run("A bottom-up approach has been used for compiling emission inventory. The emissions from individual sources are quantified initially. Emissions from all the sources have been added to obtain emission inventory for the entire operations. Following quality control steps have been adhered to in preparation of inventory to minimize uncertainty:")
        
        doc.add_paragraph()
        
        # Display Uncertainty Assessment selections from Organization Details
        uncertainty_selections = organization.get('uncertainty_assessment', [])
        
        # Map keys to full label text
        uncertainty_labels = {
            'activity_data_checked': 'The activity data has been checked from the respective sources to avoid transcription errors.',
            'inventory_calculations_checked': 'Emission inventory calculations have been checked for integrity of database and consistency of data between source categories.',
            'emission_factors_reliable': 'Emission factors have been used from reliable sources which minimizes uncertainty.',
            'instruments_calibrated': 'Instruments used for measurement and Lab analysis are calibrated regularly to reduce measurement uncertainty.'
        }
        
        if uncertainty_selections and len(uncertainty_selections) > 0:
            for selection in uncertainty_selections:
                # Use the full label if available, otherwise use the key as-is
                label_text = uncertainty_labels.get(selection, selection)
                p = doc.add_paragraph(label_text, style='List Bullet')
        else:
            p = doc.add_paragraph()
            p.add_run("NA")
        
        doc.add_paragraph()
        
        # Track organization totals
        org_totals = {
            'scope1': 0.0,
            'scope2': 0.0,
            'scope3': 0.0,
            'biogenic': 0.0,
            'removals': 0.0,
            'by_category': defaultdict(float),
            'by_fuel': defaultdict(float),
            'by_category_fuel': defaultdict(lambda: defaultdict(float)),
            'by_scope_category_fuel': defaultdict(lambda: defaultdict(lambda: defaultdict(float))),
            'by_facility': {},
            'by_facility_sinks': {}
        }
        
        period_display = f"{self._format_month_full(reporting_period_start)} - {self._format_month_full(reporting_period_end)}"
        
        # Track if any facility has prorated data (for footnote)
        any_facility_has_prorated_data = False
        
        # For each facility
        for i, facility in enumerate(facilities, 1):
            facility_id = facility.get('id')
            facility_name = self._get_value_or_na(facility, 'name')
            facility_emissions = self._get_emissions_by_facility(emissions, facility_id)
            
            # Filter by reporting period
            facility_emissions = self._filter_emissions_by_period(
                facility_emissions, reporting_period_start, reporting_period_end
            )
            
            # Deduplicate emissions to prevent double counting (yearly vs monthly)
            facility_emissions = self._deduplicate_emissions(facility_emissions)
            
            # Apply proration to yearly emissions based on reporting period overlap
            facility_emissions, has_prorated_data = self._apply_proration_to_emissions(
                facility_emissions, reporting_period_start, reporting_period_end
            )
            if has_prorated_data:
                any_facility_has_prorated_data = True
            
            # Calculate raw totals (before equity adjustment) - now using prorated values
            raw_totals = self._calculate_facility_totals(facility_emissions, facility_id)
            
            # Get equity share percentage for this facility
            equity_pct = facility_equity_map.get(facility_id, 100.0)
            equity_factor = equity_pct / 100.0
            
            # Apply equity share adjustment if applicable
            if use_equity_share and equity_factor < 1.0:
                totals = {
                    'scope1': raw_totals['scope1'] * equity_factor,
                    'scope2': raw_totals['scope2'] * equity_factor,
                    'scope3': raw_totals.get('scope3', 0) * equity_factor,
                    'biogenic': raw_totals['biogenic'] * equity_factor,
                    'removals': raw_totals['removals'] * equity_factor,
                    'total': raw_totals['total'] * equity_factor,
                    'total_ghg': raw_totals['total_ghg'] * equity_factor,
                    'by_month': raw_totals['by_month'],
                    'by_category': {k: v * equity_factor for k, v in raw_totals['by_category'].items()},
                    'by_fuel': {k: v * equity_factor for k, v in raw_totals['by_fuel'].items()},
                    'by_category_fuel': raw_totals['by_category_fuel'],
                    'by_scope_category_fuel': raw_totals['by_scope_category_fuel'],
                    'scope1_co2': raw_totals.get('scope1_co2', 0) * equity_factor,
                    'scope1_ch4': raw_totals.get('scope1_ch4', 0) * equity_factor,
                    'scope1_n2o': raw_totals.get('scope1_n2o', 0) * equity_factor,
                    'scope2_co2': raw_totals.get('scope2_co2', 0) * equity_factor,
                    'scope2_ch4': raw_totals.get('scope2_ch4', 0) * equity_factor,
                    'scope2_n2o': raw_totals.get('scope2_n2o', 0) * equity_factor,
                    # Include scope1-specific breakdowns (missing before, causing KeyError)
                    'scope1_by_category': {k: v * equity_factor for k, v in raw_totals.get('scope1_by_category', {}).items()},
                    'scope1_by_fuel': {k: v * equity_factor for k, v in raw_totals.get('scope1_by_fuel', {}).items()},
                    'scope3_by_category': {k: v * equity_factor for k, v in raw_totals.get('scope3_by_category', {}).items()},
                }
            else:
                totals = raw_totals
            
            # Update organization totals (with equity-adjusted values)
            org_totals['scope1'] += totals['scope1']
            org_totals['scope2'] += totals['scope2']
            org_totals['scope3'] += totals.get('scope3', 0)
            org_totals['biogenic'] += totals['biogenic']
            org_totals['removals'] += totals['removals']
            org_totals['by_facility'][facility_name] = totals['total']
            org_totals['by_facility_sinks'][facility_name] = totals['removals']
            
            for cat, val in totals['by_category'].items():
                org_totals['by_category'][cat] += val
            for fuel, val in totals['by_fuel'].items():
                org_totals['by_fuel'][fuel] += val
            for cat, fuels in raw_totals['by_category_fuel'].items():
                for fuel, val in fuels.items():
                    adjusted_val = val * equity_factor if use_equity_share else val
                    org_totals['by_category_fuel'][cat][fuel] += adjusted_val
            for scope, cats in raw_totals['by_scope_category_fuel'].items():
                for cat, fuels in cats.items():
                    for fuel, val in fuels.items():
                        adjusted_val = val * equity_factor if use_equity_share else val
                        org_totals['by_scope_category_fuel'][scope][cat][fuel] += adjusted_val
            
            self._add_styled_heading(doc, f"4.{i+2} Facility - {facility_name}", level=2)
            
            # Check if facility has any emissions in the reporting period
            has_emissions = len(facility_emissions) > 0
            
            # Get ALL emissions for THIS FACILITY for historical data (before checking has_emissions)
            all_facility_emissions = self._get_emissions_by_facility(emissions, facility_id)
            
            # Check if facility has sinks in the reporting period
            facility_sinks = [s for s in (self.sinks_data or []) if s.get('facility_id') == facility_id]
            has_sinks = len(facility_sinks) > 0
            raw_facility_sink_total = sum(s.get('total_emissions_reduced', 0) for s in facility_sinks)
            # Apply equity share to sinks
            facility_sink_total = raw_facility_sink_total * equity_factor if use_equity_share else raw_facility_sink_total
            
            if not has_emissions:
                # No emissions reported for this facility in the current period
                p = doc.add_paragraph()
                if has_sinks:
                    if use_equity_share and equity_pct < 100:
                        run = p.add_run(f"No emission reported for this facility in the selected reporting period. However, carbon sinks/removals totaling {self._format_number(facility_sink_total)} tCO₂e (equity-adjusted at {equity_pct:.0f}%) have been reported.")
                    else:
                        run = p.add_run(f"No emission reported for this facility in the selected reporting period. However, carbon sinks/removals totaling {self._format_number(facility_sink_total)} tCO₂e have been reported.")
                else:
                    run = p.add_run("No emission reported for this facility in the selected reporting period.")
                run.italic = True
                doc.add_paragraph()
                
                # Show Carbon Sinks section if facility has sinks
                if has_sinks:
                    self._add_styled_heading(doc, f"4.{i+2}.1 Carbon Sinks / Removals", level=3)
                    sink_headers = ['Description', 'Period', 'Emissions Reduced (tCO₂e)']
                    sink_data = []
                    for s in facility_sinks:
                        desc = s.get('description') or '-'
                        month = s.get('reporting_month')
                        year = s.get('reporting_year') or ''
                        if month is not None and year:
                            months_short = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
                            period_str = f"{months_short[month]}'{year}"
                        elif s.get('start_date'):
                            period_str = s.get('start_date', '')[:7]
                        else:
                            period_str = year or '-'
                        # Apply equity share to individual sink values
                        sink_value = s.get('total_emissions_reduced', 0) * equity_factor if use_equity_share else s.get('total_emissions_reduced', 0)
                        sink_data.append([desc, period_str, self._format_number(sink_value)])
                    
                    if sink_data:
                        self._create_styled_table(doc, sink_headers, sink_data)
                    
                    # Add equity share statement for sinks if applicable
                    if use_equity_share:
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        run = p.add_run(f"The organization has chosen the Equity Share approach. For this facility, the organization accounts for {equity_pct:.0f}% equity share; therefore, {equity_pct:.0f}% of the carbon sinks/removals from this facility are attributed to the organization.")
                        run.bold = True
                    
                    p = doc.add_paragraph()
                    p.add_run(f"Total Removals/Sinks: {self._format_number(facility_sink_total)} tCO₂e")
                    doc.add_paragraph()
                
                # Still show historical data even if no current period emissions
                if include_previous_years:
                    prev_year_data = self._get_previous_year_data(all_facility_emissions, reporting_period_start, reporting_period_end)
                    
                    # Section number depends on whether sinks section was added
                    section_num = 2 if has_sinks else 1
                    self._add_styled_heading(doc, f"4.{i+2}.{section_num} Emissions of Previous Years", level=3)
                    
                    if prev_year_data:
                        self._add_previous_years_table(doc, prev_year_data, equity_factor)
                    else:
                        doc.add_paragraph("NA")
                    doc.add_paragraph()
                
                # Base Year Emissions - show even if no current emissions (Issue 3 fix)
                base_year_data = self._get_base_year_emissions_for_entity('facility', facility_id)
                if base_year_data:
                    section_num = (3 if has_sinks else 2) if include_previous_years else (2 if has_sinks else 1)
                    self._add_styled_heading(doc, f"4.{i+2}.{section_num} Base Year Emissions", level=3)
                    # Pass zero totals since no current emissions
                    zero_totals = {'total': 0, 'scope1': 0, 'scope2': 0, 'scope3': 0, 'biogenic': 0, 'removals': 0}
                    self._add_base_year_emissions_section(doc, base_year_data, zero_totals, facility_name, equity_factor, use_equity_share, reporting_period_start, reporting_period_end)
                    doc.add_paragraph()
                
                continue
            
            # 4.x.1 List of Emissions
            self._add_styled_heading(doc, f"4.{i+2}.1 List of Emissions", level=3)
            
            # Get categorized Scope 1 emissions with extended data (responsible person, record source)
            scope1_by_category_ext = self._get_emission_processes_by_category_extended(facility_emissions)
            scope2_processes_ext = self._get_scope2_processes_extended(facility_emissions)
            
            # Get Scope 3 processes (only for scope_1_2_3 report type)
            scope3_by_category_ext = {}
            if getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3':
                scope3_by_category_ext = self._get_scope3_processes_by_category_extended(facility_emissions)
            
            # Check for data presence
            has_scope1 = any(scope1_by_category_ext.get(cat) for cat in ['stationary_combustion', 'mobile_combustion', 'fugitive_emissions', 'other'])
            has_scope3 = bool(scope3_by_category_ext) and getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3'
            
            # Create 4-column table for emissions list with Person Responsible and Source of Information
            self._add_emissions_list_table_4col(doc, scope1_by_category_ext, scope2_processes_ext, scope3_by_category_ext, has_scope1, has_scope3, getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3')
            
            # 4.x.2 Summary of GHG Emissions (renumbered since Source of Emissions is removed)
            self._add_styled_heading(doc, f"4.{i+2}.2 Summary of GHG Emissions - {period_display}", level=3)
            
            self._add_emissions_summary_table(doc, facility_emissions, totals, use_equity_share, equity_pct)
            
            # Add Scope 3 calculation method breakdown for this facility (if applicable)
            is_scope3_report = getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3'
            if is_scope3_report and totals.get('scope3', 0) > 0:
                self._add_scope3_method_breakdown(doc, facility_emissions, facility_name)
            
            doc.add_paragraph()
            
            # 4.x.3 Emissions of Previous Years - Use FACILITY-SPECIFIC historical data (already fetched above)
            if include_previous_years:
                prev_year_data = self._get_previous_year_data(all_facility_emissions, reporting_period_start, reporting_period_end)
                
                # Always add the section heading
                self._add_styled_heading(doc, f"4.{i+2}.3 Emissions of Previous Years", level=3)
                
                if prev_year_data:
                    self._add_previous_years_table(doc, prev_year_data, equity_factor)
                else:
                    # Show NA when no previous year data available
                    doc.add_paragraph("NA")
                doc.add_paragraph()
            
            # 4.x.4 Base Year Emissions - Only show if base year data is available for this facility
            base_year_data = self._get_base_year_emissions_for_entity('facility', facility_id)
            if base_year_data:
                self._add_styled_heading(doc, f"4.{i+2}.4 Base Year Emissions", level=3)
                self._add_base_year_emissions_section(doc, base_year_data, totals, facility_name, equity_factor, use_equity_share, reporting_period_start, reporting_period_end)
                doc.add_paragraph()
            
            # 4.x.5 Carbon Sinks / Removals for this facility
            if totals['removals'] > 0:
                self._add_styled_heading(doc, f"4.{i+2}.5 Carbon Sinks / Removals", level=3)
                # Get individual sink records for this facility
                facility_sinks = [s for s in (self.sinks_data or []) if s.get('facility_id') == facility_id]
                if facility_sinks:
                    sink_headers = ['Description', 'Period', 'Emissions Reduced (tCO₂e)']
                    sink_data = []
                    for s in facility_sinks:
                        desc = s.get('description') or '-'
                        month = s.get('reporting_month')
                        year = s.get('reporting_year') or ''
                        if month is not None and year:
                            months_short = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
                            period_str = f"{months_short[month]}'{year}"
                        elif s.get('start_date'):
                            period_str = s.get('start_date', '')[:7]
                        else:
                            period_str = year or '-'
                        # Apply equity share to individual sink values
                        sink_value = s.get('total_emissions_reduced', 0) * equity_factor if use_equity_share else s.get('total_emissions_reduced', 0)
                        sink_data.append([desc, period_str, self._format_number(sink_value)])
                    self._create_styled_table(doc, sink_headers, sink_data)
                
                # Add equity share statement for sinks if applicable
                if use_equity_share:
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    run = p.add_run(f"The organization has chosen the Equity Share approach. For this facility, the organization accounts for {equity_pct:.0f}% equity share; therefore, {equity_pct:.0f}% of the carbon sinks/removals from this facility are attributed to the organization.")
                    run.bold = True
                
                p = doc.add_paragraph()
                p.add_run(f"Total Removals/Sinks: {self._format_number(totals['removals'])} tCO₂e")
                doc.add_paragraph()
            
            # 4.x.6 Analysis
            next_section = 6 if totals['removals'] > 0 else 5
            self._add_styled_heading(doc, f"4.{i+2}.{next_section} Analysis", level=3)
            self._add_facility_analysis(doc, facility_name, totals)
            
            doc.add_paragraph()
            
            # 4.x.7 Carbon Intensity - Only show if production data is provided
            production_data = self.facility_production.get(facility_id)
            has_valid_production = (production_data and 
                                   production_data.get('quantity') and 
                                   float(production_data.get('quantity', 0)) > 0 and 
                                   production_data.get('unit'))
            
            if has_valid_production:
                carbon_intensity_section = next_section + 1
                self._add_styled_heading(doc, f"4.{i+2}.{carbon_intensity_section} Carbon Intensity", level=3)
                
                production_qty = float(production_data['quantity'])
                production_unit = production_data['unit']
                
                # Calculate net emissions for this facility (total emissions - sinks)
                net_emissions = totals['total'] - totals['removals']
                
                # Calculate carbon intensity
                carbon_intensity = net_emissions / production_qty
                carbon_intensity_unit = f"tCO₂e/{production_unit}"
                
                p = doc.add_paragraph()
                run = p.add_run("Carbon Intensity Formula:")
                run.bold = True
                
                p = doc.add_paragraph()
                p.add_run("Carbon Intensity = Net Emissions / Production Quantity")
                
                p = doc.add_paragraph()
                p.add_run(f"Carbon Intensity = {self._format_number(net_emissions)} tCO₂e / {self._format_number(production_qty)} {production_unit}")
                
                p = doc.add_paragraph()
                run = p.add_run(f"Carbon Intensity = {self._format_number(carbon_intensity)} {carbon_intensity_unit}")
                run.bold = True
                
                doc.add_paragraph()
                
                p = doc.add_paragraph()
                p.add_run(f"The carbon intensity of {facility_name} is {self._format_number(carbon_intensity)} {carbon_intensity_unit}. "
                          f"This metric represents the greenhouse gas emissions associated with each unit of output, providing a normalized measure of environmental performance. "
                          f"Lower carbon intensity values indicate more efficient operations from an emissions perspective, and tracking this metric over time helps identify opportunities for improvement and benchmark against industry standards.")
                
                doc.add_paragraph()
        
        # Organization-level sections - only include if all facilities are selected (is_complete_organization)
        include_org_sections = getattr(self, 'is_complete_organization', True)
        
        if include_org_sections:
            # Organization Emissions Section
            self._add_styled_heading(doc, f"4.{len(facilities)+3} Organization Emissions", level=2)
            self._add_organization_emissions_table(doc, org_totals)
            
            # Add Scope 3 calculation method breakdown at org level (if applicable)
            is_scope3_report = getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3'
            if is_scope3_report and org_totals.get('scope3', 0) > 0:
                self._add_scope3_method_breakdown(doc, emissions, organization.get('name', 'Organization'))
            
            doc.add_paragraph()
            
            # Organization Base Year Emissions - Only show if base year data is available for the organization
            org_id = organization.get('id')
            org_base_year_data = self._get_base_year_emissions_for_entity('organization', org_id)
            if org_base_year_data:
                self._add_styled_heading(doc, f"4.{len(facilities)+4} Organization Base Year Emissions", level=2)
                self._add_base_year_emissions_section(doc, org_base_year_data, org_totals, organization.get('name', 'Organization'), 1.0, False, reporting_period_start, reporting_period_end)
                doc.add_paragraph()
            
            # Organization Analysis
            analysis_section_num = len(facilities) + 5 if org_base_year_data else len(facilities) + 4
            self._add_styled_heading(doc, f"4.{analysis_section_num} Organization Analysis", level=2)
            self._add_organization_analysis(doc, organization, org_totals, facilities)
        
        doc.add_page_break()
    
    def _add_scope3_method_breakdown(self, doc: Document, emissions: List[Dict], entity_name: str):
        """Add inline text showing Scope 3 calculation method breakdown for a facility/organization.
        
        Format: "Supplier Basis: 45%, Activity Basis: 30%, Spend Basis: 25%"
        """
        # Filter scope 3 emissions
        scope3_emissions = [e for e in emissions if (e.get('scope') or '').lower() == 'scope3']
        
        if not scope3_emissions:
            return
        
        # Analyze emissions by methodology
        methodology_totals = {
            'activity_basis': 0.0,
            'spend_basis': 0.0,
            'supplier_basis': 0.0,
            'other': 0.0
        }
        
        for em in scope3_emissions:
            method = (em.get('calculation_method_scope3') or 'other').lower()
            co2e = float(em.get('total_emissions') or em.get('co2e_emissions') or 0)
            
            # Normalize method name
            if 'activity' in method:
                method_key = 'activity_basis'
            elif 'spend' in method:
                method_key = 'spend_basis'
            elif 'supplier' in method:
                method_key = 'supplier_basis'
            else:
                method_key = 'other'
            
            methodology_totals[method_key] += co2e
        
        # Calculate grand total
        grand_total = sum(methodology_totals.values())
        
        if grand_total > 0:
            # Build inline text for non-zero methods
            method_labels = {
                'supplier_basis': 'Supplier Basis',
                'activity_basis': 'Activity Basis',
                'spend_basis': 'Spend Basis',
                'other': 'Other'
            }
            
            breakdown_parts = []
            for method_key, label in method_labels.items():
                if methodology_totals[method_key] > 0:
                    pct = (methodology_totals[method_key] / grand_total) * 100
                    breakdown_parts.append(f"{label}: {pct:.1f}%")
            
            if breakdown_parts:
                p = doc.add_paragraph()
                run = p.add_run("Scope 3 Calculation Method Breakdown: ")
                run.bold = True
                p.add_run(", ".join(breakdown_parts))
    
    def _add_emissions_summary_table(self, doc: Document, facility_emissions: List[Dict], totals: Dict, 
                                      use_equity_share: bool = False, equity_pct: float = 100.0):
        """Add emissions summary table for a facility - sorted hierarchically: Scope → Category → Fuel/Activity → Reporting Period
        
        Includes proration markers (*) in Reporting Period column for yearly data that has been prorated.
        """
        # Check if this is a Scope 3 report
        is_scope3_report = getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3'
        
        # Changed "Month" to "Reporting Period"
        if is_scope3_report:
            headers = ['Scope', 'Category', 'Fuel/ Activity Name', 'Reporting Period', 'tCO2e', 'tCO2', 'tCH4', 'tN2O']
        else:
            headers = ['Scope', 'Category', 'Fuel', 'Reporting Period', 'tCO2e', 'tCO2', 'tCH4', 'tN2O']
        data = []
        
        # Track unique entries to prevent duplicates
        seen_entries = set()
        
        # Track if any emissions are prorated
        has_prorated_data = False
        
        # Helper function to get scope sort order
        def get_scope_order(scope):
            scope_lower = (scope or '').lower()
            if 'scope1' in scope_lower or 'scope 1' in scope_lower or scope == '1':
                return 1
            elif 'scope2' in scope_lower or 'scope 2' in scope_lower or scope == '2':
                return 2
            elif 'scope3' in scope_lower or 'scope 3' in scope_lower or scope == '3':
                return 3
            elif 'biogenic' in scope_lower:
                return 4
            return 9
        
        # Helper function to parse date for sorting
        def get_date_key(em):
            period = em.get('reporting_period') or ''
            month_str = period.split(' to ')[0].strip() if ' to ' in period else period.strip()
            try:
                if '-' in month_str:
                    parts = month_str.split('-')
                    if len(parts) >= 2 and len(parts[0]) == 4:  # YYYY-MM format
                        return (int(parts[0]), int(parts[1]))
                return (9999, 99)  # Put unparseable dates at end
            except Exception:
                return (9999, 99)
        
        # Sort emissions hierarchically: 1) Scope, 2) Category, 3) Fuel, 4) Date
        def sort_key(em):
            scope_order = get_scope_order(em.get('scope', ''))
            category = (self._get_category_from_emission(em) or '').lower()
            fuel = (self._get_fuel_from_emission(em) or '').lower()
            date_key = get_date_key(em)
            return (scope_order, category, fuel, date_key[0], date_key[1])
        
        sorted_emissions = sorted(facility_emissions, key=sort_key)
        
        for em in sorted_emissions:
            # Create unique key to prevent duplicates
            period = em.get('reporting_period') or ''
            month_str = period.split(' to ')[0].strip() if ' to ' in period else period.strip()
            category = self._get_category_from_emission(em)
            fuel = self._get_fuel_from_emission(em)
            scope = em.get('scope') or ''
            
            unique_key = f"{scope}|{category}|{fuel}|{month_str}|{em.get('id', '')}"
            if unique_key in seen_entries:
                continue
            seen_entries.add(unique_key)
            
            # Check if this emission is prorated
            is_prorated = em.get('_is_prorated', False)
            if is_prorated:
                has_prorated_data = True
            
            # Format scope display
            scope_lower = scope.lower() if scope else ''
            if 'scope1' in scope_lower or 'scope 1' in scope_lower or scope == '1':
                scope_display = 'Scope 1 (Direct)'
            elif 'scope2' in scope_lower or 'scope 2' in scope_lower or scope == '2':
                scope_display = 'Scope 2 (Indirect)'
            elif 'biogenic' in scope_lower:
                scope_display = 'Biogenic'
            else:
                scope_display = scope or 'Unknown'
            
            # Format reporting period display - add star (*) here if prorated
            reporting_period_display = self._format_month(month_str)
            if is_prorated:
                reporting_period_display = f"{reporting_period_display} *"
            
            # Emission values - no stars here anymore
            tco2e_display = self._format_number(em.get('total_emissions', 0) or em.get('co2e_emissions', 0))
            
            # For Scope 3, use "-" for tCO2, tCH4, tN2O columns since they're not calculated individually
            is_scope3 = 'scope3' in scope_lower or 'scope 3' in scope_lower or scope == '3'
            
            if is_scope3:
                data.append([
                    scope_display,
                    category,
                    fuel,
                    reporting_period_display,
                    tco2e_display,
                    '-',
                    '-',
                    '-'
                ])
            else:
                co2_val = self._format_number(em.get('co2_emissions', 0))
                ch4_val = self._format_number(em.get('ch4_emissions', 0))
                n2o_val = self._format_number(em.get('n2o_emissions', 0))
                
                data.append([
                    scope_display,
                    category,
                    fuel,
                    reporting_period_display,
                    tco2e_display,
                    co2_val,
                    ch4_val,
                    n2o_val
                ])
        
        # Create table WITHOUT totals (totals will be added separately)
        self._create_styled_table(doc, headers, data)
        
        # Add proration footnote if any data was prorated
        if has_prorated_data:
            p = doc.add_paragraph()
            run = p.add_run("* - The emissions reported here are proportionally calculated based on the actual data corresponding to the selected reporting period.")
            run.italic = True
            run.font.size = Pt(10)
        
        # Add Equity Share statement BEFORE Summary Totals (for all equity share facilities)
        if use_equity_share:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(f"The organization has chosen the Equity Share approach. For this facility, the organization accounts for {equity_pct:.0f}% equity share; therefore, {equity_pct:.0f}% of the GHG emissions from this facility are attributed to the organization.")
            run.bold = True
        
        # Add totals OUTSIDE the table
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        run = p.add_run("Summary Totals (all values in tCO₂e):")
        run.bold = True
        run.font.size = Pt(11)
        
        # Check if this is a Scope 3 report
        is_scope3_report = getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3'
        scope3_total = totals.get('scope3', 0)
        
        if is_scope3_report:
            total_abc = totals['scope1'] + totals['scope2'] + scope3_total
            net_ghg = total_abc - totals['removals']
            totals_text = [
                f"Total Scope 1 Emissions (A): {self._format_number(totals['scope1'])} tCO₂e",
                f"Total Scope 2 Emissions (B): {self._format_number(totals['scope2'])} tCO₂e",
                f"Total Scope 3 Emissions (C): {self._format_number(scope3_total)} tCO₂e",
                f"Total Emissions (A + B + C): {self._format_number(total_abc)} tCO₂e",
                f"Total Removals/Sinks (D): {self._format_number(totals['removals'])} tCO₂e",
                f"Net GHG Emissions (A + B + C - D): {self._format_number(net_ghg)} tCO₂e",
                f"Total Biogenic Emissions: {self._format_number(totals['biogenic'])} tCO₂e"
            ]
        else:
            totals_text = [
                f"Total Direct Emissions (A): {self._format_number(totals['scope1'])} tCO₂e",
                f"Total Indirect Emissions (B): {self._format_number(totals['scope2'])} tCO₂e",
                f"Total Emissions (A + B): {self._format_number(totals['total'])} tCO₂e",
                f"Total Removals/Sinks (C): {self._format_number(totals['removals'])} tCO₂e",
                f"Net GHG Emissions (A + B - C): {self._format_number(totals['total_ghg'])} tCO₂e",
                f"Total Biogenic Emissions: {self._format_number(totals['biogenic'])} tCO₂e"
            ]
        
        for text in totals_text:
            p = doc.add_paragraph()
            p.add_run(text)
    
    def _add_previous_years_table(self, doc: Document, prev_year_data: List[Dict], equity_factor: float = 1.0):
        """
        Add previous years emissions table with new format.
        
        Columns: Scope, Category, Fuel/Activity, Reporting Period, tCO2e
        Adds * to Reporting Period for prorated records.
        """
        if not prev_year_data:
            doc.add_paragraph("No previous year data available.")
            return
        
        headers = ['Scope', 'Category', 'Fuel/Activity', 'Reporting Period', 'tCO2e']
        data = []
        has_prorated = False
        
        for record in prev_year_data:
            scope = record.get('scope', '')
            category = record.get('category', '')
            fuel = record.get('fuel', '')
            reporting_period = record.get('reporting_period', '')
            tco2e = record.get('tco2e', 0)
            is_prorated = record.get('is_prorated', False)
            
            # Apply equity factor
            adjusted_tco2e = tco2e * equity_factor
            
            # Add * to reporting period if prorated
            period_display = reporting_period
            if is_prorated:
                period_display = f"{reporting_period} *"
                has_prorated = True
            
            data.append([
                scope,
                category,
                fuel if fuel else 'NA',
                period_display,
                self._format_number(adjusted_tco2e)
            ])
        
        if data:
            self._create_styled_table(doc, headers, data)
            
            # Add proration footnote if any data was prorated
            if has_prorated:
                doc.add_paragraph()
                p = doc.add_paragraph()
                run = p.add_run("* - The emissions reported here are proportionally calculated based on the portion falling before the selected reporting period.")
                run.italic = True
                run.font.size = Pt(10)
        else:
            doc.add_paragraph("No previous year data available.")
    
    def _add_facility_analysis(self, doc: Document, facility_name: str, totals: Dict):

        # Check if this is a Scope 3 report
        is_scope3_report = getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3'

        """Add analysis text for a facility"""
        total_emissions = totals['total']
        scope1 = totals['scope1']
        scope2 = totals['scope2']
        scope3 = totals.get('scope3', 0) if is_scope3_report else 0
        
        if is_scope3_report:
            total_emissions = scope1 + scope2 + scope3
        else:
            total_emissions = scope1 + scope2 
        # Total emissions statement
        p = doc.add_paragraph()
        p.add_run(f"Total emissions from {facility_name} amount to ")
        run = p.add_run(f"{self._format_number(total_emissions)} tCO2e")
        run.bold = True
        p.add_run(".")
        
        # Scope contribution
        if total_emissions > 0:
            scope1_pct = (scope1 / total_emissions) * 100
            scope2_pct = (scope2 / total_emissions) * 100
            p = doc.add_paragraph()
            if is_scope3_report:
                scope3_pct = (scope3 / total_emissions) * 100
                p.add_run(f" Scope 1 emissions account for {scope1_pct:.1f}%, Scope 2 emissions account for {scope2_pct:.1f}%, and Scope 3 emissions account for {scope3_pct:.1f}% of total emissions.")
            else:        
                p.add_run(f"Scope 1 (Direct) emissions contribute {scope1_pct:.1f}% ({self._format_number(scope1)} tCO2e) of total emissions, while Scope 2 (Indirect) emissions contribute {scope2_pct:.1f}% ({self._format_number(scope2)} tCO2e).")
            
            # Category dominance - use scope1-specific breakdown (use .get() for safety)
            scope1_by_category = totals.get('scope1_by_category', {})
            if scope1_by_category:
                top_category = max(scope1_by_category.items(), key=lambda x: x[1])
                cat_pct = (top_category[1] / scope1) * 100 if scope1 > 0 else 0
                p = doc.add_paragraph()
                p.add_run("Among Scope 1 categories, ")
                run = p.add_run(f"{top_category[0]}")
                run.bold = True
                p.add_run(f" is the dominant source, contributing {cat_pct:.1f}% of direct emissions.")
            
            # Fuel dominance - use scope1-specific breakdown (use .get() for safety)
            scope1_by_fuel = totals.get('scope1_by_fuel', {})
            if scope1_by_fuel:
                top_fuel = max(scope1_by_fuel.items(), key=lambda x: x[1])
                fuel_pct = (top_fuel[1] / scope1) * 100 if scope1 > 0 else 0
                p = doc.add_paragraph()
                p.add_run("In terms of fuel consumption, ")
                run = p.add_run(f"{top_fuel[0]}")
                run.bold = True
                p.add_run(f" is the primary contributor, accounting for {fuel_pct:.1f}% of Scope 1 emissions.")
        
        doc.add_paragraph()
        
        # Determine if any charts will be shown
        has_scope_chart = scope1 > 0 or scope2 > 0
        has_category_chart = bool(totals['by_category'])
        has_fuel_chart = bool(totals['by_fuel'])
        has_monthly_chart = bool(totals['by_month'])
        has_any_chart = has_scope_chart or has_category_chart or has_fuel_chart or has_monthly_chart
        

        
        # Add charts (reduced size) - Only add header text if at least one chart is successfully added
        charts_added = False
        try:
            # Scope comparison chart - show if any one of scope1 or scope2 has values
            if scope1 > 0 or scope2 > 0 or (scope3 and scope3 > 0):
                # For Scope 1,2,3 reports, include scope3 in the comparison
                chart_buf = self._create_scope_comparison_chart(scope1, scope2, scope3)
                if not charts_added:
                    p = doc.add_paragraph()
                    p.add_run("The following figures illustrate the emission distribution:")
                    charts_added = True
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(chart_buf, width=Inches(4.5 if is_scope3_report else 4))
                # Update caption based on report type
                caption = "Figure: Scope 1 vs Scope 2 vs Scope 3 Emissions Comparison" if is_scope3_report else "Figure: Scope 1 vs Scope 2 Emissions Comparison"
                self._add_figure_caption(doc, caption)
            
            # Category chart
            if totals['by_category']:
                chart_buf = self._create_category_chart(dict(totals['by_category']))
                if not charts_added:
                    p = doc.add_paragraph()
                    p.add_run("The following figures illustrate the emission distribution:")
                    charts_added = True
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(chart_buf, width=Inches(3.4))  # Reduced from 4 by 15%
                self._add_figure_caption(doc, "Figure: Category-wise Emission Distribution")
            
            # Fuel chart (now a bar chart like scope comparison)
            if totals['by_fuel']:
                chart_buf = self._create_fuel_chart(dict(totals['by_fuel']))
                if not charts_added:
                    p = doc.add_paragraph()
                    p.add_run("The following figures illustrate the emission distribution:")
                    charts_added = True
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(chart_buf, width=Inches(4))  # Same width as scope comparison
                self._add_figure_caption(doc, "Figure: Fuel-wise Emission Distribution")
            
            # Monthly trend - Only for Scope 1,2 reports (not for Scope 1,2,3)
            # is_scope3_report is already defined above
            if totals['by_month'] and not is_scope3_report:
                chart_buf = self._create_monthly_trend_chart(dict(totals['by_month']))
                if not charts_added:
                    p = doc.add_paragraph()
                    p.add_run("The following figures illustrate the emission distribution:")
                    charts_added = True
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(chart_buf, width=Inches(4.5))
                self._add_figure_caption(doc, "Figure: Monthly Emission Trend")
                
        except Exception as e:
            print(f"Error adding facility charts: {e}")
    
    def _add_organization_emissions_table(self, doc: Document, org_totals: Dict):
        """Add organization-level emissions summary table"""
        headers = ['Category', 'Fuel', 'Total Emissions (tCO₂e)']
        data = []
        bold_rows = []  # Track which rows should be bold
        
        scope_cat_fuel = org_totals.get('by_scope_category_fuel', {})
        
        # Direct/Scope 1 Emissions (bold header row)
        scope1_data = scope_cat_fuel.get('scope1', {})
        bold_rows.append(len(data))
        data.append(['Scope 1 Emissions', '', ''])
        
        if scope1_data:
            for cat in sorted(scope1_data.keys()):
                fuels_in_cat = scope1_data[cat]
                fuels_str = ", ".join(self._deduplicate_list(sorted(fuels_in_cat.keys()), case_insensitive=True))
                cat_total = sum(fuels_in_cat.values())
                data.append([cat, fuels_str, self._format_number(cat_total)])
        else:
            data.append(['No emission reported', '-', '0.00'])
        
        # Indirect/Scope 2 Emissions (bold header row)
        scope2_data = scope_cat_fuel.get('scope2', {})
        bold_rows.append(len(data))
        data.append(['Scope 2 Emissions', '', ''])
        
        if scope2_data:
            for cat in sorted(scope2_data.keys()):
                fuels_in_cat = scope2_data[cat]
                fuels_str = ", ".join(self._deduplicate_list(sorted(fuels_in_cat.keys()), case_insensitive=True))
                cat_total = sum(fuels_in_cat.values())
                data.append([cat, fuels_str, self._format_number(cat_total)])
        else:
            data.append(['No emission reported', '-', self._format_number(org_totals['scope2'])])
        
        # Biogenic Emissions (bold header row if there's data)
        biogenic_data = scope_cat_fuel.get('biogenic', {})
        biogenic = org_totals.get('biogenic', 0)
        if biogenic > 0:
            bold_rows.append(len(data))
            data.append(['Biogenic Emissions', '', ''])
            if biogenic_data:
                for cat in sorted(biogenic_data.keys()):
                    fuels_in_cat = biogenic_data[cat]
                    fuels_str = ", ".join(self._deduplicate_list(sorted(fuels_in_cat.keys()), case_insensitive=True))
                    cat_total = sum(fuels_in_cat.values())
                    data.append([cat, fuels_str, self._format_number(cat_total)])
            else:
                data.append(['Biogenic sources', '-', self._format_number(biogenic)])
        
        # Sinks/Removals (bold header row if there's data)
        removals = org_totals.get('removals', 0)
        if removals > 0:
            bold_rows.append(len(data))
            data.append(['Sinks/Removals', '', ''])
            data.append(['Carbon sinks and removals', '-', self._format_number(removals)])
        
        # Create table with bold header rows
        self._create_styled_table(doc, headers, data, bold_rows=bold_rows)
        
        # Add totals OUTSIDE the table
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        run = p.add_run("Organization Summary Totals (all values in tCO₂e):")
        run.bold = True
        run.font.size = Pt(11)
        
        # Check if this is a Scope 3 report
        is_scope3_report = getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3'
        scope3_total = org_totals.get('scope3', 0)
        
        if is_scope3_report:
            total_emissions = org_totals['scope1'] + org_totals['scope2'] + scope3_total
            net_emissions = total_emissions - removals
            
            totals_text = [
                f"Total Scope1 Emissions (A): {self._format_number(org_totals['scope1'])} tCO₂e",
                f"Total Scope2 Emissions (B): {self._format_number(org_totals['scope2'])} tCO₂e",
                f"Total Scope 3 Emissions ( C ): {self._format_number(scope3_total)} tCO₂e",
                f"Total Emissions (A + B + C): {self._format_number(total_emissions)} tCO₂e",
                f"Total Removals/Sinks (D): {self._format_number(removals)} tCO₂e",
                f"Total Biogenic: {self._format_number(org_totals.get('biogenic', 0))} tCO₂e",
                f"Net GHG Emissions (A + B + C - D): {self._format_number(net_emissions)} tCO₂e"
            ]
        else:
            total_emissions = org_totals['scope1'] + org_totals['scope2']
            net_emissions = total_emissions - removals
            
            totals_text = [
                f"Total Direct Emissions (A): {self._format_number(org_totals['scope1'])} tCO₂e",
                f"Total Indirect Emissions (B): {self._format_number(org_totals['scope2'])} tCO₂e",
                f"Total Emissions (A + B): {self._format_number(total_emissions)} tCO₂e",
                f"Total Removals/Sinks (C): {self._format_number(removals)} tCO₂e",
                f"Total Biogenic: {self._format_number(org_totals.get('biogenic', 0))} tCO₂e",
                f"Net GHG Emissions (A + B - C): {self._format_number(net_emissions)} tCO₂e"
            ]
        
        for text in totals_text:
            p = doc.add_paragraph()
            p.add_run(text)
    
    def _add_organization_analysis(self, doc: Document, organization: Dict, org_totals: Dict, facilities: List[Dict]):
        """Add organization-level analysis"""
        org_name = self._get_value_or_na(organization, 'name')
        is_scope3_report = getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3'
        
        scope3_total = org_totals.get('scope3', 0) if is_scope3_report else 0
        total = org_totals['scope1'] + org_totals['scope2'] + scope3_total
        removals = org_totals.get('removals', 0)
        net_emissions = total - removals
        
        p = doc.add_paragraph()
        p.add_run("The total GHG emissions for ")
        run = p.add_run(f"{org_name}")
        run.bold = True
        p.add_run(" amount to ")
        run = p.add_run(f"{self._format_number(total)} tCO2e")
        run.bold = True
        p.add_run(f" across {len(facilities)} selected facilities.")
        
        # Add Carbon Sinks information if present
        if removals > 0:
            p = doc.add_paragraph()
            p.add_run("The organization has reported carbon sinks/removals totaling ")
            run = p.add_run(f"{self._format_number(removals)} tCO2e")
            run.bold = True
            p.add_run(". After accounting for these removals, the ")
            run = p.add_run("net GHG emissions")
            run.bold = True
            p.add_run(" for the organization amount to ")
            run = p.add_run(f"{self._format_number(net_emissions)} tCO2e")
            run.bold = True
            p.add_run(".")
            
            # Sinks by facility breakdown if available
            if org_totals.get('by_facility_sinks'):
                facilities_with_sinks = {k: v for k, v in org_totals['by_facility_sinks'].items() if v > 0}
                if facilities_with_sinks:
                    p = doc.add_paragraph()
                    p.add_run("Carbon sinks contribution by facility:")
                    for fac_name, sink_total in sorted(facilities_with_sinks.items(), key=lambda x: -x[1]):
                        sink_pct = (sink_total / removals) * 100 if removals > 0 else 0
                        p = doc.add_paragraph()
                        p.add_run(f"• {fac_name}: {self._format_number(sink_total)} tCO2e ({sink_pct:.1f}%)")
        
        # Scope distribution
        if total > 0:
            scope1_pct = (org_totals['scope1'] / total) * 100
            scope2_pct = (org_totals['scope2'] / total) * 100
            
            p = doc.add_paragraph()
            if is_scope3_report and scope3_total > 0:
                scope3_pct = (scope3_total / total) * 100
                p.add_run(f"At the organizational level, Scope 1 emissions account for {scope1_pct:.1f}%, Scope 2 emissions account for {scope2_pct:.1f}%, and Scope 3 emissions account for {scope3_pct:.1f}% of total emissions.")
            else:
                p.add_run(f"At the organizational level, Scope 1 emissions account for {scope1_pct:.1f}% of total emissions, while Scope 2 emissions account for {scope2_pct:.1f}%.")
        
        # Facility-wise comparison
        if org_totals['by_facility']:
            p = doc.add_paragraph()
            p.add_run("Facility-wise emission contribution:")
            
            for fac_name, fac_total in sorted(org_totals['by_facility'].items(), key=lambda x: -x[1]):
                fac_pct = (fac_total / total) * 100 if total > 0 else 0
                p = doc.add_paragraph()
                p.add_run(f"• {fac_name}: {self._format_number(fac_total)} tCO2e ({fac_pct:.1f}%)")
        
        # Mathematical validation
        p = doc.add_paragraph()
        p.add_run("\n")
        run = p.add_run("Mathematical Validation:")
        run.bold = True
        
        p = doc.add_paragraph()
        facility_sum = sum(org_totals['by_facility'].values())
        p.add_run(f"Sum of facility totals: {self._format_number(facility_sum)} tCO2e")
        
        p = doc.add_paragraph()
        if is_scope3_report:
            p.add_run(f"Organization total (A+B+C): {self._format_number(total)} tCO2e")
        else:
            p.add_run(f"Organization total (A+B): {self._format_number(total)} tCO2e")
        
        # Percentage validation
        if total > 0:
            pct_sum = sum((v/total)*100 for v in org_totals['by_facility'].values())
            p = doc.add_paragraph()
            p.add_run(f"Sum of facility percentages: {pct_sum:.1f}%")
        
        # Facility comparison chart
        try:
            if org_totals['by_facility']:
                chart_buf = self._create_facility_comparison_chart(org_totals['by_facility'])
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(chart_buf, width=Inches(4.5))
                self._add_figure_caption(doc, "Figure: Facility-wise Emission Comparison")
        except Exception as e:
            print(f"Error adding organization chart: {e}")
        
        # Category-wise analysis chart for Scope 1,2,3 reports
        if is_scope3_report:
            try:
                # Build category data from by_scope_category_fuel
                category_data = {}
                if org_totals.get('by_scope_category_fuel'):
                    for scope, categories in org_totals['by_scope_category_fuel'].items():
                        for cat, fuels in categories.items():
                            if cat not in category_data:
                                category_data[cat] = 0.0
                            category_data[cat] += sum(fuels.values())
                
                # Alternative: use by_category if by_scope_category_fuel is empty
                if not category_data and org_totals.get('by_category'):
                    category_data = dict(org_totals['by_category'])
                
                if category_data:
                    # Filter out zero/negligible values and sort by emissions
                    category_data = {k: v for k, v in category_data.items() if v > 0.001}
                    
                    if category_data:
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        run = p.add_run("Category-wise Emission Analysis:")
                        run.bold = True
                        
                        # Add text summary
                        sorted_categories = sorted(category_data.items(), key=lambda x: -x[1])
                        total_cat = sum(category_data.values())
                        
                        for cat_name, cat_total in sorted_categories[:10]:  # Top 10 categories
                            cat_pct = (cat_total / total_cat) * 100 if total_cat > 0 else 0
                            p = doc.add_paragraph()
                            p.add_run(f"• {cat_name}: {self._format_number(cat_total)} tCO2e ({cat_pct:.1f}%)")
                        
                        # Create category chart
                        chart_buf = self._create_category_analysis_chart(category_data)
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(chart_buf, width=Inches(5.5))
                        self._add_figure_caption(doc, "Figure: Category-wise Emission Distribution")
            except Exception as e:
                print(f"Error adding category analysis chart: {e}")
                import traceback
                traceback.print_exc()
    
    def _create_category_analysis_chart(self, category_data: Dict[str, float]) -> io.BytesIO:
        """Create a horizontal bar chart for category-wise emission analysis"""
        # Sort by emissions descending and take top 15
        sorted_data = sorted(category_data.items(), key=lambda x: -x[1])[:15]
        
        if not sorted_data:
            # Return empty chart
            fig, ax = plt.subplots(figsize=(8, 4))
            ax.text(0.5, 0.5, 'No category data available', ha='center', va='center')
            buf = io.BytesIO()
            plt.savefig(buf, format='png', dpi=120, bbox_inches='tight')
            buf.seek(0)
            plt.close(fig)
            return buf
        
        categories = [item[0] for item in sorted_data]
        values = [item[1] for item in sorted_data]
        
        # Truncate long category names
        categories = [cat[:40] + '...' if len(cat) > 40 else cat for cat in categories]
        
        # Reverse for horizontal bar chart (so highest is at top)
        categories = categories[::-1]
        values = values[::-1]
        
        fig, ax = plt.subplots(figsize=(10, max(6, len(categories) * 0.4)))
        
        # Create color gradient
        colors = plt.cm.Blues(np.linspace(0.4, 0.9, len(categories)))
        
        bars = ax.barh(categories, values, color=colors, edgecolor='black', linewidth=0.5)
        
        # Add value labels
        for bar, val in zip(bars, values):
            ax.text(bar.get_width() + max(values) * 0.01, bar.get_y() + bar.get_height()/2,
                    f'{val:,.2f}', ha='left', va='center', fontsize=8)
        
        ax.set_xlabel('tCO2e', fontsize=10)
        ax.set_title('Category-wise Emission Distribution', fontsize=12, fontweight='bold')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        
        # Adjust x-axis to make room for labels
        ax.set_xlim(0, max(values) * 1.15)
        
        plt.tight_layout()
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=120, bbox_inches='tight', facecolor='white')
        buf.seek(0)
        plt.close(fig)
        
        return buf
    
    def _generate_chapter5(self, doc: Document, organization: Dict):
        """Chapter 5: GHG REDUCTION INITIATIVE AND INTERNAL PERFORMANCE TRACKING"""
        self._add_styled_heading(doc, "Chapter 5: GHG REDUCTION INITIATIVE AND INTERNAL PERFORMANCE TRACKING", level=1)
        
        # Check if organization has any reduction initiatives or performance tracking data
        initiatives = self._get_value_or_na(organization, 'ghg_reduction_initiatives')
        tracking = self._get_value_or_na(organization, 'internal_performance_tracking')
        
        has_initiatives_data = initiatives and initiatives != 'NA' and initiatives.strip()
        has_tracking_data = tracking and tracking != 'NA' and tracking.strip()
        
        # Add introductory text only for Scope 1,2,3 reports AND only if organization has data
        is_scope3_report = getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3'
        
        if is_scope3_report and (has_initiatives_data or has_tracking_data):
            # Add Chapter 5 introduction
            intro_paragraphs = [
                "Greenhouse Gas (GHG) reduction initiatives and internal performance tracking form an essential part of an organization's overall sustainability and climate management strategy. While the preparation of a GHG inventory helps quantify and understand emissions sources, reduction initiatives focus on actively minimizing the organization's environmental impact through targeted operational, technological, and strategic improvements.",
                
                "GHG reduction initiatives may include measures such as improving energy efficiency, optimizing resource consumption, adopting cleaner technologies, transitioning to renewable energy sources, enhancing waste management practices, improving logistics and transportation efficiency, and promoting sustainable operational practices across facilities and business functions. These initiatives help organizations reduce operational emissions, improve environmental performance, and contribute toward broader global climate goals.",
                
                "Internal performance tracking plays a critical role in ensuring that emission reduction efforts are measurable, transparent, and effective over time. By regularly monitoring emissions data, energy consumption patterns, reduction targets, and sustainability performance indicators, organizations can evaluate progress, identify areas for improvement, and make informed decisions for future climate action planning.",
                
                "Establishing a structured framework for performance tracking also enables organizations to:"
            ]
            
            for para_text in intro_paragraphs:
                doc.add_paragraph(para_text)
            
            # Add bullet points
            tracking_benefits = [
                "Measure progress against emission reduction goals and sustainability commitments",
                "Improve accountability across departments and operational units",
                "Identify inefficiencies and opportunities for operational optimization",
                "Support regulatory compliance and reporting requirements",
                "Enhance data accuracy, consistency, and transparency",
                "Strengthen stakeholder confidence through measurable climate action",
                "Facilitate long-term sustainability planning and risk management"
            ]
            
            for benefit in tracking_benefits:
                doc.add_paragraph(benefit, style='List Bullet')
            
            doc.add_paragraph()
            
            closing_para = "Effective GHG reduction and performance management not only support environmental stewardship but also contribute to operational resilience, cost optimization, improved resource efficiency, and alignment with evolving stakeholder and regulatory expectations. As climate-related risks and sustainability considerations continue to gain importance globally, organizations increasingly recognize the need for continuous monitoring, proactive emissions management, and transparent reporting practices as part of responsible business operations."
            doc.add_paragraph(closing_para)
            
            doc.add_paragraph()
        
        # GHG Reduction Initiatives
        self._add_styled_heading(doc, "5.1 GHG Reduction Initiatives", level=2)
        
        if has_initiatives_data:
            doc.add_paragraph(initiatives)
        else:
            doc.add_paragraph("NA")
        
        # Internal Performance Tracking
        self._add_styled_heading(doc, "5.2 Internal Performance Tracking", level=2)
        
        if has_tracking_data:
            doc.add_paragraph(tracking)
        else:
            doc.add_paragraph("NA")
        
        doc.add_page_break()
    
    def _generate_chapter6(self, doc: Document, organization: Dict):
        """Chapter 6: Conclusion"""
        self._add_styled_heading(doc, "Chapter 6: Conclusion", level=1)
        
        org_name = self._get_value_or_na(organization, 'name')
        is_scope3_report = getattr(self, 'report_type', 'scope_1_2') == 'scope_1_2_3'
        
        if is_scope3_report:
            # New conclusion for Scope 1,2,3 reports
            p = doc.add_paragraph()
            run = p.add_run(f"{org_name}")
            run.bold = True
            p.add_run(" recognizes the growing importance of climate accountability and sustainable business practices in addressing global environmental challenges. This GHG Inventory Report provides a comprehensive assessment of the organization's greenhouse gas emissions and has been prepared in accordance with the principles and requirements of the GHG Protocol, ensuring transparency, consistency, accuracy, and reliability in emissions accounting and reporting.")
            
            doc.add_paragraph()
            
            doc.add_paragraph("The inventory reflects the organization's ongoing commitment to understanding, measuring, and managing its environmental impact across operational activities and facilities. By establishing a robust emissions baseline, the report enables the organization to identify key emission sources, evaluate reduction opportunities, and strengthen climate-related decision-making processes.")
            
            doc.add_paragraph()
            
            doc.add_paragraph("As part of its sustainability and climate action commitments, the organization aims to:")
            
            commitments = [
                "Continuously monitor, measure, and report greenhouse gas emissions in alignment with internationally recognized standards and best practices",
                "Enhance the accuracy, completeness, and reliability of emissions data through improved data collection, validation, and verification mechanisms",
                "Identify operational efficiencies and implement emission reduction initiatives across facilities and business activities",
                "Promote energy efficiency, resource optimization, and the adoption of cleaner and lower-carbon technologies wherever feasible",
                "Develop and pursue science-based and measurable emission reduction targets aligned with long-term sustainability objectives",
                "Strengthen internal awareness and stakeholder engagement to encourage collaborative climate action and responsible environmental practices",
                "Regularly review and improve GHG management frameworks to align with evolving regulatory requirements, industry expectations, and global climate goals"
            ]
            
            for commitment in commitments:
                doc.add_paragraph(commitment, style='List Bullet')
            
            doc.add_paragraph()
            
            doc.add_paragraph("This report serves not only as a record of current emissions performance but also as a strategic foundation for future sustainability initiatives, climate risk management, and environmental stewardship.")
            
            doc.add_paragraph()
            
            p = doc.add_paragraph()
            run = p.add_run(f"{org_name}")
            run.bold = True
            p.add_run(" remains committed to continuously improving its greenhouse gas management practices and contributing meaningfully toward the global transition to a more sustainable and low-carbon future.")
        else:
            # Original conclusion for Scope 1,2 reports
            p = doc.add_paragraph()
            p.add_run("This GHG Inventory Report presents a comprehensive assessment of the greenhouse gas emissions for ")
            run = p.add_run(f"{org_name}")
            run.bold = True
            p.add_run(". The inventory has been prepared in accordance with the principles and requirements of ")
            run = p.add_run("ISO 14064-1:2018")
            run.bold = True
            p.add_run(".")
            
            doc.add_paragraph()
            
            p = doc.add_paragraph()
            p.add_run("The organization is committed to:")
            
            commitments = [
                "Continuously monitoring and reporting GHG emissions in accordance with international standards",
                "Implementing measures to reduce emissions across all operational facilities",
                "Improving data collection and verification processes",
                "Setting science-based targets for emission reduction",
                "Engaging stakeholders in climate action initiatives"
            ]
            
            for commitment in commitments:
                doc.add_paragraph(commitment, style='List Bullet')
            
            doc.add_paragraph()
            
            p = doc.add_paragraph()
            p.add_run("This report serves as a foundation for informed decision-making regarding climate action and sustainability strategies. The organization remains dedicated to environmental stewardship and will continue to enhance its GHG management practices in alignment with global climate goals.")
    
    # ==================== MAIN GENERATION METHOD ====================
    
    def generate_report(self, organization: Dict, facilities: List[Dict], emissions: List[Dict],
                       reporting_period_start: str, reporting_period_end: str,
                       include_previous_years: bool = True,
                       sinks_total: float = 0.0, sinks_data: List[Dict] = None,
                       facility_production: Dict = None,
                       report_type: str = "scope_1_2",
                       is_complete_organization: bool = True) -> io.BytesIO:
        """Generate the complete GHG Inventory Report
        
        Args:
            report_type: "scope_1_2" for Scope 1,2 report or "scope_1_2_3" for Scope 1,2,3 report
            is_complete_organization: Whether all org facilities are included (for org-level sections)
        """
        
        # Store report type for use in chapter generation
        self.report_type = report_type
        
        # Store flag for whether to include org-level sections
        self.is_complete_organization = is_complete_organization
        
        # Store sinks data for use in calculations
        self.sinks_total = sinks_total
        self.sinks_data = sinks_data or []
        self.facility_production = facility_production or {}
        
        # Create new document
        doc = Document()
        
        # Set document-wide formatting
        self._set_document_font(doc)
        self._add_page_border(doc)
        
        # Generate all chapters
        self._generate_cover_page(doc, organization, reporting_period_start, reporting_period_end)
        self._generate_chapter1(doc, organization, facilities)
        self._generate_chapter2(doc, organization)
        self._generate_chapter3(doc, facilities, emissions, reporting_period_start, reporting_period_end)
        self._generate_chapter4(doc, organization, facilities, emissions, 
                               reporting_period_start, reporting_period_end, include_previous_years)
        self._generate_chapter5(doc, organization)
        self._generate_chapter6(doc, organization)
        
        # Add footer to all sections
        self._add_footer(doc)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
